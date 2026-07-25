from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
from PIL import Image, ImageDraw, ImageFont

from build_project_documents import (
    AMBER,
    CALLOUT,
    DOCS,
    GREEN,
    BLUE_GRAY,
    bullet,
    callout,
    h1,
    h2,
    h3,
    header_footer,
    matrix,
    number,
    para,
    process_table,
    save,
    style_document,
    title_block,
)


SRS_OUTPUT = DOCS / "ResidenceCore_SRS_Detailed.docx"
BLUEPRINT_OUTPUT = DOCS / "ResidenceCore_BA_Blueprint_Detailed.docx"
MANUAL_OUTPUT = DOCS / "ResidenceCore_User_Manual_Detailed.docx"
ASSETS = DOCS / "assets" / "ba"
SCREENSHOTS = DOCS / "assets" / "screenshots"


def _font(size: int, bold: bool = False):
    candidates = [
        "C:/Windows/Fonts/arialbd.ttf" if bold else "C:/Windows/Fonts/arial.ttf",
        "C:/Windows/Fonts/calibrib.ttf" if bold else "C:/Windows/Fonts/calibri.ttf",
    ]
    for candidate in candidates:
        try:
            return ImageFont.truetype(candidate, size)
        except OSError:
            continue
    return ImageFont.load_default()


def _wrap_text(draw: ImageDraw.ImageDraw, text: str, font, max_width: int) -> list[str]:
    lines: list[str] = []
    for raw in text.split("\n"):
        words = raw.split()
        current = ""
        for word in words:
            candidate = f"{current} {word}".strip()
            if draw.textbbox((0, 0), candidate, font=font)[2] <= max_width:
                current = candidate
            else:
                if current:
                    lines.append(current)
                current = word
        if current:
            lines.append(current)
    return lines


def _box(draw: ImageDraw.ImageDraw, xy, title: str, body: str = "", fill="#F4F6F9", outline="#2E74B5") -> None:
    x1, y1, x2, y2 = xy
    draw.rounded_rectangle(xy, radius=18, fill=fill, outline=outline, width=3)
    title_font = _font(22, True)
    body_font = _font(17)
    draw.text((x1 + 18, y1 + 16), title, fill="#0B2545", font=title_font)
    if body:
        y = y1 + 50
        for line in _wrap_text(draw, body, body_font, x2 - x1 - 36):
            draw.text((x1 + 18, y), line, fill="#334155", font=body_font)
            y += 23


def _arrow(draw: ImageDraw.ImageDraw, start: tuple[int, int], end: tuple[int, int], color="#1F4D78") -> None:
    draw.line([start, end], fill=color, width=4)
    x1, y1 = start
    x2, y2 = end
    if abs(x2 - x1) >= abs(y2 - y1):
        direction = 1 if x2 >= x1 else -1
        points = [(x2, y2), (x2 - 16 * direction, y2 - 9), (x2 - 16 * direction, y2 + 9)]
    else:
        direction = 1 if y2 >= y1 else -1
        points = [(x2, y2), (x2 - 9, y2 - 16 * direction), (x2 + 9, y2 - 16 * direction)]
    draw.polygon(points, fill=color)


def _diagram_canvas(title: str) -> tuple[Image.Image, ImageDraw.ImageDraw]:
    image = Image.new("RGB", (1600, 920), "#FFFFFF")
    draw = ImageDraw.Draw(image)
    draw.rectangle((0, 0, 1600, 86), fill="#0B2545")
    draw.text((48, 25), title, fill="#FFFFFF", font=_font(30, True))
    return image, draw


def build_diagrams() -> dict[str, Path]:
    ASSETS.mkdir(parents=True, exist_ok=True)
    diagrams: dict[str, Path] = {}

    image, draw = _diagram_canvas("ResidenceCore - Main Manager Workflow")
    boxes = [
        ((70, 160, 360, 310), "Login", "Manager đăng nhập\nvà đổi mật khẩu\nnếu cần"),
        ((470, 160, 760, 310), "Dashboard", "Xem tổng quan\nvà cảnh báo nhanh"),
        ((870, 160, 1160, 310), "Học viên", "Danh sách, hồ sơ,\ntrạng thái cư trú"),
        ((1270, 160, 1530, 310), "Tổ chức", "Nhiệm kỳ, Tổ/Ban,\nchức vụ, bổ nhiệm"),
        ((270, 520, 560, 690), "Tài chính", "Kỳ thu, khoản thu,\nthanh toán"),
        ((680, 520, 970, 690), "Sinh hoạt", "Lịch ngày,\nviệc hôm nay"),
        ((1090, 520, 1380, 690), "Công tác", "Mẫu công tác,\nphân công, trạng thái"),
    ]
    for xy, title, body in boxes:
        _box(draw, xy, title, body)
    for start, end in [
        ((360, 235), (470, 235)),
        ((760, 235), (870, 235)),
        ((1160, 235), (1270, 235)),
        ((1400, 310), (410, 520)),
        ((560, 605), (680, 605)),
        ((970, 605), (1090, 605)),
    ]:
        _arrow(draw, start, end)
    path = ASSETS / "manager_workflow.png"
    image.save(path)
    diagrams["manager_workflow"] = path

    image, draw = _diagram_canvas("ResidenceCore - Data Flow Context")
    _box(draw, (610, 350, 990, 560), "ResidenceCore", "API + business services\nRBAC + validation", fill="#E8EEF5")
    _box(draw, (80, 170, 380, 320), "Manager", "Quản lý dữ liệu\nvà nghiệp vụ")
    _box(draw, (80, 610, 380, 760), "Resident", "Xem portal\nvà công tác cá nhân")
    _box(draw, (1220, 160, 1520, 320), "Database", "Users, residents,\nfinance, duties,\norganization, store")
    _box(draw, (1220, 610, 1520, 760), "Documents", "SRS, Blueprint,\nUser Manual,\nDeployment Standard")
    for start, end in [
        ((380, 245), (610, 420)),
        ((380, 685), (610, 500)),
        ((990, 420), (1220, 245)),
        ((990, 500), (1220, 685)),
    ]:
        _arrow(draw, start, end)
    draw.text((430, 260), "CRUD / UAT", fill="#1F4D78", font=_font(18, True))
    draw.text((430, 650), "Portal request", fill="#1F4D78", font=_font(18, True))
    draw.text((1040, 280), "Read / write", fill="#1F4D78", font=_font(18, True))
    draw.text((1040, 650), "Traceability", fill="#1F4D78", font=_font(18, True))
    path = ASSETS / "data_flow_context.png"
    image.save(path)
    diagrams["data_flow_context"] = path

    image, draw = _diagram_canvas("ResidenceCore - Resident Store Shift Workflow")
    boxes = [
        ((70, 150, 360, 300), "Resident", "Mở /resident/store"),
        ((470, 150, 760, 300), "Chọn ngày + ca", "Không dùng token"),
        ((870, 150, 1160, 300), "Kiểm tra phân công", "Đúng resident,\nledger, ngày, ca"),
        ((1270, 150, 1530, 300), "Vào phiên", "Có quyền theo ca"),
        ((390, 540, 690, 710), "Phiên hiện tại", "Cho thêm/sửa/xóa\nbán/nhập theo rule"),
        ((910, 540, 1210, 710), "Không hiện tại", "Chỉ xem/chốt sổ,\nkhông ghi giao dịch"),
    ]
    for xy, title, body in boxes:
        _box(draw, xy, title, body)
    for start, end in [
        ((360, 225), (470, 225)),
        ((760, 225), (870, 225)),
        ((1160, 225), (1270, 225)),
        ((1400, 300), (540, 540)),
        ((1400, 300), (1060, 540)),
    ]:
        _arrow(draw, start, end)
    label_font = _font(20, True)
    draw.rounded_rectangle((590, 456, 865, 500), radius=12, fill="#FFFFFF", outline="#FFFFFF")
    draw.text((610, 465), "Nếu là phiên đang trực", fill="#1F4D78", font=label_font)
    draw.rounded_rectangle((1005, 456, 1345, 500), radius=12, fill="#FFFFFF", outline="#FFFFFF")
    draw.text((1030, 465), "Nếu không phải phiên hiện tại", fill="#1F4D78", font=label_font)
    path = ASSETS / "resident_store_shift_workflow.png"
    image.save(path)
    diagrams["resident_store_shift"] = path

    image, draw = _diagram_canvas("ResidenceCore - BA Delivery Process")
    swimlanes = [
        ("BA/PO", 125, "#FFF7E6"),
        ("Dev", 325, "#E8EEF5"),
        ("QA/UAT", 525, "#EAF7EF"),
        ("Ops", 725, "#F4F6F9"),
    ]
    for label, y, fill in swimlanes:
        draw.rectangle((60, y, 1540, y + 150), fill=fill, outline="#CBD5E1", width=2)
        draw.text((80, y + 55), label, fill="#0B2545", font=_font(22, True))
    steps = [
        (240, 160, "Chốt scope\nSRS"),
        (470, 360, "Build\nmodule"),
        (700, 360, "Check/Test\nBuild"),
        (930, 560, "UAT\nscript"),
        (1160, 760, "Deploy\nDemo"),
        (1390, 160, "Nghiệm thu\nBacklog"),
    ]
    for x, y, label in steps:
        _box(draw, (x, y, x + 170, y + 88), label.split("\n")[0], "\n".join(label.split("\n")[1:]), fill="#FFFFFF")
    for start, end in [
        ((410, 204), (470, 404)),
        ((640, 404), (700, 404)),
        ((870, 404), (930, 604)),
        ((1100, 604), (1160, 804)),
        ((1330, 804), (1390, 204)),
    ]:
        _arrow(draw, start, end)
    path = ASSETS / "ba_delivery_process.png"
    image.save(path)
    diagrams["ba_delivery_process"] = path
    return diagrams


def figure(doc: Document, image_path: Path, caption: str, width: float = 6.35) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.add_run().add_picture(str(image_path), width=Inches(width))
    caption_paragraph = doc.add_paragraph()
    caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = caption_paragraph.add_run(caption)
    run.italic = True
    run.font.size = Pt(9)
    doc.add_paragraph()


def screenshot_backlog(doc: Document) -> None:
    matrix(
        doc,
        ["Mã ảnh", "Màn hình", "Route", "Nội dung cần chụp"],
        [
            ["SS-01", "Login", "/", "Form đăng nhập và trạng thái đổi mật khẩu lần đầu."],
            ["SS-02", "Dashboard", "/dashboard", "Các chỉ số tổng quan sau khi manager đăng nhập."],
            ["SS-03", "Học viên", "/members", "Danh sách học viên dạng thẻ/list, filter và trạng thái."],
            ["SS-04", "Tổ chức", "/organization", "Nhiệm kỳ, Tổ/Ban, chức vụ và bổ nhiệm."],
            ["SS-05", "Tài chính", "/finance", "Kỳ thu, khoản phải thu, giao dịch thanh toán."],
            ["SS-06", "Sinh hoạt", "/daily-routine", "Lịch ngày và công tác hôm nay."],
            ["SS-07", "Công tác", "/duties", "Danh sách phân công và form tạo/cập nhật."],
            ["SS-08", "Portal học viên", "/resident/today", "Màn hình Hôm nay của resident demo."],
            ["SS-09", "Resident Store", "/resident/store", "Chọn ngày/ca và trạng thái quyền trong phiên."],
            ["SS-10", "Store Cashflow", "/store-cashflow", "Sổ phát sinh, group ngày, trạng thái chốt/xác nhận."],
        ],
        [0.8, 1.55, 1.35, 2.8],
    )


def actual_screenshot_figures(doc: Document) -> None:
    screenshots = [
        ("SS-02-dashboard-viewport.png", "Ảnh 1. Dashboard manager sau khi đăng nhập."),
        ("SS-03-members-viewport.png", "Ảnh 2. Màn hình Học viên dạng danh sách/thẻ."),
        ("SS-04-organization-viewport.png", "Ảnh 3. Màn hình Tổ chức lưu xá."),
        ("SS-05-finance-viewport.png", "Ảnh 4. Màn hình Tài chính lưu xá."),
        ("SS-06-daily-routine-viewport.png", "Ảnh 5. Màn hình Sinh hoạt hằng ngày."),
        ("SS-07-duties-viewport.png", "Ảnh 6. Màn hình Công tác / Trực nhật."),
    ]
    available = [(SCREENSHOTS / filename, caption) for filename, caption in screenshots if (SCREENSHOTS / filename).exists()]
    if not available:
        callout(
            doc,
            "Chưa có ảnh chụp thực tế",
            "Khi app demo đang chạy, chụp các route trong bảng screenshot backlog rồi lưu vào docs/assets/screenshots để nhúng vào manual.",
            fill=AMBER,
        )
        return
    for image_path, caption in available:
        figure(doc, image_path, caption, width=6.35)


def numbered_step_table(doc: Document, rows: list[list[str]]) -> None:
    matrix(doc, ["Bước", "Người thao tác", "Thao tác chi tiết", "Kết quả mong đợi"], rows, [0.75, 1.25, 2.95, 1.55])


def intro_document_control(doc: Document, document_name: str, purpose: str) -> None:
    h1(doc, "0. Kiểm soát tài liệu")
    matrix(
        doc,
        ["Thuộc tính", "Nội dung"],
        [
            ["Tên tài liệu", document_name],
            ["Dự án", "ResidenceCore - Hệ thống quản lý lưu xá"],
            ["Phiên bản", "v1.0 - BA baseline cho demo/triển khai"],
            ["Người đọc chính", "Project Owner, BA, Developer, QA, DevOps, người vận hành lưu xá"],
            ["Mục đích", purpose],
            ["Ngôn ngữ", "Tiếng Việt, ưu tiên dễ hiểu, có đủ tiêu chí nghiệm thu"],
            ["Ghi chú", "Tài liệu bám theo trạng thái mã nguồn hiện tại và main flow P0 đã được ưu tiên."],
        ],
        [1.75, 4.75],
        fill=BLUE_GRAY,
    )


def build_srs() -> Path:
    diagrams = build_diagrams()
    doc = Document()
    style_document(doc)
    header_footer(doc, "SRS", "Software Requirements Specification for ResidenceCore")
    title_block(
        doc,
        "Software Requirements Specification (SRS)",
        "Tài liệu đặc tả yêu cầu phần mềm chi tiết theo chuẩn BA triển khai dự án",
        "SRS / Functional & Non-functional Requirements",
        "Baseline dùng để chốt phạm vi, phát triển, kiểm thử UAT và triển khai demo",
    )
    intro_document_control(
        doc,
        "ResidenceCore_SRS_Detailed.docx",
        "Mô tả chi tiết yêu cầu nghiệp vụ, yêu cầu chức năng, phi chức năng, dữ liệu, phân quyền, kiểm thử và tiêu chí nghiệm thu của ResidenceCore.",
    )

    callout(
        doc,
        "Tóm tắt điều hành",
        "ResidenceCore là hệ thống quản lý lưu xá tập trung cho manager và học viên. Phạm vi ưu tiên hiện tại là main flow P0: đăng nhập, dashboard, học viên, tổ chức lưu xá, tài chính, sinh hoạt hằng ngày, công tác/trực nhật và portal học viên. Module cửa hàng đã có nền nhưng không bắt buộc trong demo chính nếu chưa chốt dữ liệu 2 ca.",
        fill=GREEN,
    )

    h1(doc, "1. Mục tiêu và phạm vi")
    h2(doc, "1.1 Mục tiêu hệ thống")
    bullet(doc, "Số hóa dữ liệu học viên lưu xá, tổ chức nội bộ, tài chính, sinh hoạt và công tác.")
    bullet(doc, "Giúp manager đi từ tổng quan đến xử lý công việc hằng ngày trong một giao diện thống nhất.")
    bullet(doc, "Cho học viên xem thông tin cá nhân, công tác, tài chính, thông báo và hoạt động qua portal.")
    bullet(doc, "Chuẩn bị nền triển khai demo, UAT và mở rộng các module nâng cao sau khi main flow ổn định.")
    h2(doc, "1.2 Phạm vi P0")
    matrix(
        doc,
        ["Mã", "Phạm vi P0", "Mục tiêu nghiệm thu"],
        [
            ["P0-01", "Đăng nhập và phiên làm việc", "Manager và resident đăng nhập được, phân quyền đúng."],
            ["P0-02", "Dashboard", "Manager xem được tổng quan dữ liệu lưu xá."],
            ["P0-03", "Quản lý học viên", "Danh sách/hồ sơ học viên hiển thị đúng và dễ thao tác."],
            ["P0-04", "Tổ chức lưu xá", "Hiển thị nhiệm kỳ, Tổ/Ban, chức vụ, bổ nhiệm."],
            ["P0-05", "Tài chính lưu xá", "Xem kỳ thu, khoản phải thu, ghi nhận thu chi cơ bản."],
            ["P0-06", "Sinh hoạt hằng ngày", "Xem lịch/công tác theo ngày."],
            ["P0-07", "Công tác / trực nhật", "Tạo, xem, cập nhật phân công công tác."],
            ["P0-08", "Portal học viên", "Học viên xem hôm nay, hồ sơ, công tác, tài chính, thông báo, hoạt động."],
        ],
        [0.75, 2.75, 3.0],
    )
    h2(doc, "1.3 Ngoài phạm vi P0")
    bullet(doc, "Các module học tập nâng cao, kỹ năng, phụng vụ, báo cáo nâng cao, phụ huynh, phân công thông minh.")
    bullet(doc, "Cửa hàng lưu xá trong demo chính nếu chưa hoàn tất dữ liệu resident portal và kịch bản 1 ngày / 2 ca.")
    bullet(doc, "Tích hợp thanh toán online, chữ ký số, thông báo đa kênh hoặc đồng bộ hệ thống bên ngoài.")

    h1(doc, "2. Stakeholders và vai trò")
    matrix(
        doc,
        ["Vai trò", "Mục tiêu", "Quyền chính trong hệ thống", "Tài liệu cần đọc"],
        [
            ["Project Owner", "Chốt phạm vi, nghiệm thu business value.", "Xem toàn cảnh, duyệt UAT.", "SRS, Blueprint, UAT checklist."],
            ["Manager lưu xá", "Vận hành dữ liệu và công việc hằng ngày.", "Quản lý học viên, tổ chức, tài chính, công tác.", "User Manual, SRS flow."],
            ["Học viên", "Xem dữ liệu cá nhân và công tác.", "Portal resident, dữ liệu của chính mình.", "User Manual phần resident."],
            ["BA", "Đặc tả nghiệp vụ và tiêu chí nghiệm thu.", "Không bắt buộc có quyền app.", "SRS, Blueprint."],
            ["Developer", "Hiện thực yêu cầu và sửa lỗi.", "Quyền kỹ thuật theo môi trường.", "SRS, Blueprint, Deployment Standard."],
            ["QA/UAT", "Kiểm thử hệ thống và nghiệm thu.", "Tài khoản test manager/resident.", "SRS, User Manual, Test scenarios."],
            ["DevOps/Operator", "Deploy, backup, vận hành.", "Môi trường, env, database.", "Deployment Standard."],
        ],
        [1.25, 1.8, 2.05, 1.4],
    )

    h1(doc, "3. Thuật ngữ nghiệp vụ")
    matrix(
        doc,
        ["Thuật ngữ", "Ý nghĩa trong ResidenceCore"],
        [
            ["Lưu xá", "Đơn vị quản lý nơi học viên cư trú, sinh hoạt, đóng phí và tham gia công tác."],
            ["Học viên / Resident", "Người ở trong lưu xá, có hồ sơ và có thể có tài khoản portal."],
            ["Manager", "Người quản lý hệ thống và vận hành nghiệp vụ lưu xá."],
            ["Simple mode", "Chế độ menu gọn, chỉ hiển thị main flow P0 để demo/vận hành đơn giản."],
            ["Detailed mode", "Chế độ menu đầy đủ hơn cho quản trị hoặc các module mở rộng."],
            ["Nhiệm kỳ", "Khoảng thời gian áp dụng cơ cấu tổ chức, chức vụ, bổ nhiệm."],
            ["Tổ/Ban", "Đơn vị tổ chức nội bộ của lưu xá."],
            ["Công tác", "Việc được phân công cho học viên hoặc đơn vị, có ngày/ca/trạng thái."],
            ["Kỳ thu", "Đợt thu tài chính theo tháng/kỳ/nội dung."],
            ["Store shift", "Ca trực cửa hàng theo ngày/ca, chỉ áp dụng khi demo module cửa hàng."],
        ],
        [1.75, 4.75],
    )

    h1(doc, "4. Mô hình phân quyền")
    matrix(
        doc,
        ["Vai trò", "Được phép", "Không được phép trong P0"],
        [
            ["Manager", "Quản lý dữ liệu main flow, xem dashboard, tạo/cập nhật học viên, tài chính, công tác, tổ chức.", "Không thao tác dữ liệu hệ thống ngoài quyền môi trường."],
            ["Resident", "Xem portal cá nhân, công tác, tài chính, thông báo, hoạt động; thao tác cửa hàng nếu đúng ca được phân công.", "Không xem dữ liệu học viên khác, không quản trị hệ thống."],
            ["Appointed resident", "Có thể có menu/phạm vi theo chức vụ nếu được bổ nhiệm.", "Không tự cấp quyền cho bản thân."],
        ],
        [1.35, 3.0, 2.15],
    )
    callout(doc, "Nguyên tắc bảo mật", "Mọi dữ liệu resident trên portal phải được lọc theo user đang đăng nhập và hồ sơ resident đã liên kết. Không dùng dữ liệu thật cho demo public nếu chưa được phép.", fill=AMBER)

    h1(doc, "5. Yêu cầu chức năng chi tiết")
    modules = [
        (
            "FR-01",
            "Đăng nhập và tài khoản",
            [
                ["FR-01.1", "Người dùng đăng nhập bằng username/password.", "Must", "Đăng nhập đúng thông tin thì vào được app."],
                ["FR-01.2", "Hệ thống hỗ trợ mustChangePassword cho lần đăng nhập đầu.", "Must", "Người dùng đổi mật khẩu xong đăng nhập lại được."],
                ["FR-01.3", "Hệ thống phân route/menu theo role manager/resident.", "Must", "Resident không thấy menu manager."],
                ["FR-01.4", "Seed demo manager dùng admin / Admin@123.", "Must", "Login UI và seed thống nhất."],
            ],
        ),
        (
            "FR-02",
            "Dashboard",
            [
                ["FR-02.1", "Manager xem chỉ số tổng quan học viên/phòng/công tác/tài chính.", "Must", "Dashboard render không lỗi và có dữ liệu demo."],
                ["FR-02.2", "Dashboard có lối đi nhanh sang module chính.", "Should", "Bấm mở được các màn hình P0."],
                ["FR-02.3", "Dashboard không phụ thuộc module P2.", "Must", "Ẩn/không crash khi module nâng cao chưa sẵn sàng."],
            ],
        ),
        (
            "FR-03",
            "Quản lý học viên",
            [
                ["FR-03.1", "Hiển thị danh sách học viên dạng thẻ/list.", "Must", "Danh sách rõ ràng, font và màu hài hòa."],
                ["FR-03.2", "Xem hồ sơ chi tiết học viên.", "Must", "Thông tin cá nhân, trạng thái, phòng/tổ hiển thị."],
                ["FR-03.3", "Tìm kiếm/lọc danh sách học viên.", "Should", "Nhập từ khóa lọc được danh sách."],
                ["FR-03.4", "Liên kết tài khoản portal với hồ sơ học viên.", "Must", "Resident login xem đúng hồ sơ."],
                ["FR-03.5", "Không phân công/phòng cho học viên inactive/left nếu rule không cho phép.", "Must", "Backend chặn và UI báo rõ."],
            ],
        ),
        (
            "FR-04",
            "Phòng ở",
            [
                ["FR-04.1", "Route phòng ở giữ cho Detailed mode.", "Should", "Detailed mode mở được /rooms."],
                ["FR-04.2", "Simple mode không hiển thị menu Phòng ở.", "Must", "Sidebar Simple không có Phòng ở."],
                ["FR-04.3", "Kiểm tra sức chứa phòng khi gán học viên.", "Must", "Không vượt capacity."],
            ],
        ),
        (
            "FR-05",
            "Tổ chức lưu xá",
            [
                ["FR-05.1", "Quản lý nhiệm kỳ active.", "Must", "Có ít nhất một nhiệm kỳ dùng cho demo."],
                ["FR-05.2", "Quản lý Tổ/Ban và chức vụ.", "Must", "Hiển thị được cấu trúc tổ chức."],
                ["FR-05.3", "Bổ nhiệm học viên vào chức vụ/đơn vị.", "Must", "Bổ nhiệm hiển thị trong cơ cấu."],
                ["FR-05.4", "Kết thúc bổ nhiệm thay vì xóa lịch sử.", "Should", "Lưu được quá trình tham gia."],
            ],
        ),
        (
            "FR-06",
            "Tài chính lưu xá",
            [
                ["FR-06.1", "Tạo/xem kỳ thu và khoản phải thu.", "Must", "Có kỳ thu demo và danh sách khoản phải thu."],
                ["FR-06.2", "Ghi nhận thanh toán cho học viên.", "Must", "Số đã thu/còn nợ cập nhật đúng."],
                ["FR-06.3", "Chặn amount <= 0.", "Must", "Không tạo giao dịch sai số tiền."],
                ["FR-06.4", "Không tạo trùng khoản theo tháng/kỳ nếu rule chặn trùng.", "Should", "Backend trả lỗi rõ."],
                ["FR-06.5", "Nhận dòng tổng hợp từ cửa hàng sau khi manager xác nhận.", "Should", "Finance chỉ nhận dữ liệu đã confirmed."],
            ],
        ),
        (
            "FR-07",
            "Sinh hoạt hằng ngày",
            [
                ["FR-07.1", "Xem lịch theo ngày/tuần/tháng.", "Must", "Màn hình render được lịch."],
                ["FR-07.2", "Hiển thị công tác trong ngày.", "Must", "Ngày demo có dữ liệu công tác."],
                ["FR-07.3", "Dùng shared date/time picker.", "Must", "Không dùng input thời gian rời rạc."],
            ],
        ),
        (
            "FR-08",
            "Công tác / trực nhật",
            [
                ["FR-08.1", "Quản lý mẫu công tác.", "Must", "Có thể tạo/xem mẫu công tác."],
                ["FR-08.2", "Phân công học viên hoặc đơn vị.", "Must", "Phân công lưu đúng ngày/ca/người."],
                ["FR-08.3", "Cập nhật trạng thái công tác.", "Must", "Hoàn thành/vắng/hủy hiển thị đúng."],
                ["FR-08.4", "Resident xem công tác của chính mình.", "Must", "Không xem công tác của người khác."],
                ["FR-08.5", "Trực cửa hàng chọn ngày/ca, không dùng token.", "Should", "Đúng ngày+ca của resident thì vào phiên."],
            ],
        ),
        (
            "FR-09",
            "Portal học viên",
            [
                ["FR-09.1", "Resident xem màn hình Hôm nay.", "Must", "Có công tác/lịch/thông báo liên quan."],
                ["FR-09.2", "Resident xem hồ sơ cá nhân.", "Must", "Dữ liệu đúng hồ sơ liên kết."],
                ["FR-09.3", "Resident xem tài chính cá nhân.", "Must", "Khoản phải thu/đã thu hiển thị đúng."],
                ["FR-09.4", "Resident xem thông báo/hoạt động.", "Should", "Danh sách render không lỗi."],
                ["FR-09.5", "Resident store chỉ cho ghi giao dịch khi là phiên hiện tại.", "Should", "Không phải phiên hiện tại chỉ xem/chốt theo rule."],
            ],
        ),
        (
            "FR-10",
            "Cửa hàng lưu xá",
            [
                ["FR-10.1", "Quản lý sản phẩm, nhóm hàng, đơn vị tính.", "Should", "Tạo sản phẩm không bắt buộc giá bán."],
                ["FR-10.2", "Nhập kho đa nguồn.", "Should", "Mua hàng tạo chi; nguồn khác chỉ tăng tồn."],
                ["FR-10.3", "Bán hàng theo sản phẩm.", "Should", "Bán giảm tồn, tạo thu, chặn tồn âm."],
                ["FR-10.4", "Báo cáo tồn kho.", "Should", "Tổng tồn/giá vốn/doanh thu dự kiến khớp chi tiết."],
                ["FR-10.5", "Chốt ngày, xác nhận, đẩy sổ chung.", "Should", "Finance nhận tổng sau confirmed, không trùng batch."],
            ],
        ),
    ]
    for code, title, rows in modules:
        h2(doc, f"{code}. {title}")
        matrix(doc, ["Mã yêu cầu", "Mô tả", "Mức ưu tiên", "Tiêu chí nghiệm thu"], rows, [1.0, 2.7, 0.95, 1.85])

    h1(doc, "6. Yêu cầu phi chức năng")
    matrix(
        doc,
        ["Mã", "Nhóm", "Yêu cầu", "Tiêu chí kiểm tra"],
        [
            ["NFR-01", "Usability", "Giao diện Simple mode gọn, dễ demo, không đưa module chưa sẵn sàng vào menu chính.", "Sidebar đúng danh sách P0."],
            ["NFR-02", "Visual style", "Header các trang đồng nhất, title/subtitle căn giữa, form premium nhẹ, không rối.", "Review UI các trang P0."],
            ["NFR-03", "Reliability", "Main route P0 không crash runtime.", "Smoke browser pass."],
            ["NFR-04", "Security", "Phân quyền theo role, portal lọc theo resident hiện tại.", "Test negative role/resident."],
            ["NFR-05", "Data integrity", "Giao dịch tiền/tồn kho không tạo số âm hoặc trùng bất hợp lý.", "Unit/integration test pass."],
            ["NFR-06", "Maintainability", "Logic backend có service/router rõ, migration đặt trong /drizzle.", "Code review."],
            ["NFR-07", "Compatibility", "Ứng dụng chạy trên Chrome/Edge mới.", "Smoke test browser."],
            ["NFR-08", "Performance", "Demo route P0 tải trong ngưỡng chấp nhận được với dữ liệu seed.", "Manual smoke, theo dõi chunk nếu cần."],
            ["NFR-09", "Auditability", "Checklist, tài liệu và kịch bản test được cập nhật sau thay đổi lớn.", "RESIDENCECORE_CHECKLIST.md có record."],
        ],
        [0.85, 1.1, 3.1, 1.45],
    )

    h1(doc, "7. Yêu cầu dữ liệu")
    matrix(
        doc,
        ["Nhóm dữ liệu", "Entity chính", "Yêu cầu chất lượng dữ liệu"],
        [
            ["Người dùng", "users, roles", "Username duy nhất, role rõ, password hash hợp lệ, mustChangePassword đúng rule."],
            ["Học viên", "residents, rooms", "Hồ sơ có trạng thái, liên kết phòng/tổ nếu có, không dùng dữ liệu cá nhân thật trong demo public."],
            ["Tổ chức", "organization terms, units, roles, assignments", "Có nhiệm kỳ active, bổ nhiệm có ngày bắt đầu/kết thúc."],
            ["Công tác", "duties, duty templates, assignments", "Có ngày/ca/người phụ trách/trạng thái; portal chỉ trả dữ liệu của resident."],
            ["Tài chính", "finance periods, receivables, transactions", "Amount > 0, không trùng ngoài rule, số đã thu/còn nợ tính đúng."],
            ["Cửa hàng", "products, stock movements, ledger entries, closings", "Tồn không âm, chốt ngày không cho sửa phát sinh đã khóa."],
            ["Thông báo/hoạt động", "notifications, activities", "Ngày hiệu lực rõ, nội dung demo không nhạy cảm."],
        ],
        [1.25, 2.25, 3.0],
    )

    h1(doc, "8. Luồng nghiệp vụ chính")
    h2(doc, "8.1 Workflow tổng quan manager")
    figure(
        doc,
        diagrams["manager_workflow"],
        "Hình 1. Workflow manager từ đăng nhập đến các module main flow P0.",
    )
    h2(doc, "8.2 Data flow context")
    figure(
        doc,
        diagrams["data_flow_context"],
        "Hình 2. Data flow context giữa người dùng, ResidenceCore, database và bộ tài liệu triển khai.",
    )
    h2(doc, "8.3 Process tổng thể")
    process_table(
        doc,
        [
            ["1", "Manager", "Đăng nhập bằng tài khoản demo hoặc tài khoản thật.", "Vào Dashboard."],
            ["2", "Manager", "Xem Dashboard và kiểm tra dữ liệu tổng quan.", "Biết trạng thái lưu xá trong ngày."],
            ["3", "Manager", "Mở Học viên, xem danh sách và hồ sơ.", "Xác nhận dữ liệu cư trú."],
            ["4", "Manager", "Mở Tổ chức, xem nhiệm kỳ và bổ nhiệm.", "Xác nhận cơ cấu lưu xá."],
            ["5", "Manager", "Mở Tài chính, xem kỳ thu và khoản phải thu.", "Xác nhận tình trạng thu/còn nợ."],
            ["6", "Manager", "Mở Sinh hoạt/Công tác, tạo hoặc theo dõi phân công.", "Ngày sinh hoạt có người phụ trách."],
            ["7", "Resident", "Đăng nhập portal và xem dữ liệu cá nhân.", "Resident thấy đúng hồ sơ/công tác/tài chính."],
            ["8", "QA/UAT", "Chạy checklist nghiệm thu.", "Pass/Fail được ghi nhận theo từng yêu cầu."],
        ],
    )
    h2(doc, "8.4 Workflow trực cửa hàng theo ngày/ca")
    figure(
        doc,
        diagrams["resident_store_shift"],
        "Hình 3. Workflow resident chọn ngày/ca để vào phiên trực cửa hàng, không dùng token.",
    )

    h1(doc, "9. Kịch bản kiểm thử UAT")
    matrix(
        doc,
        ["Mã test", "Kịch bản", "Dữ liệu cần có", "Kết quả mong đợi"],
        [
            ["UAT-01", "Manager đăng nhập và đổi mật khẩu lần đầu.", "admin / Admin@123.", "Đổi mật khẩu thành công, login lại được."],
            ["UAT-02", "Mở 7 route P0 manager.", "DB demo có seed.", "Không 404, không runtime crash."],
            ["UAT-03", "Xem danh sách học viên.", "8-15 học viên demo.", "Danh sách đẹp, có trạng thái, tìm kiếm được."],
            ["UAT-04", "Xem tổ chức lưu xá.", "1 nhiệm kỳ active, Tổ/Ban, chức vụ.", "Cơ cấu hiển thị rõ."],
            ["UAT-05", "Ghi nhận thanh toán học viên.", "Kỳ thu và khoản phải thu.", "Số đã thu/còn lại cập nhật đúng."],
            ["UAT-06", "Tạo/cập nhật công tác.", "Mẫu công tác, resident demo.", "Công tác xuất hiện trong manager và portal."],
            ["UAT-07", "Resident xem portal.", "Resident user liên kết hồ sơ.", "Hồ sơ/công tác/tài chính đúng resident."],
            ["UAT-08", "Negative role test.", "Resident login.", "Không truy cập route manager."],
            ["UAT-09", "Store shift access nếu demo cửa hàng.", "2 resident, 1 ngày, ca sáng/chiều.", "Đúng ngày+ca mới vào; non-current chỉ xem/chốt."],
        ],
        [0.9, 2.25, 1.7, 1.65],
    )

    h1(doc, "10. Traceability yêu cầu - module - kiểm thử")
    matrix(
        doc,
        ["Yêu cầu", "Module/UI", "API/Service liên quan", "Test/UAT"],
        [
            ["FR-01", "Login, layout, navigation", "authRouter, authService, RBAC", "UAT-01, UAT-08"],
            ["FR-02", "Dashboard", "dashboardRouter", "UAT-02"],
            ["FR-03/04", "Members, Rooms", "membersRouter, roomsRouter, memberService, roomService", "UAT-03"],
            ["FR-05", "Organization", "organizationRouter, organizationService", "UAT-04"],
            ["FR-06", "FinanceLite", "financeRouter", "UAT-05"],
            ["FR-07/08", "DailyRoutine, Duties", "dailyRoutineRouter, dutiesRouter, dailyRoutineService", "UAT-06"],
            ["FR-09", "Resident portal", "residentPortalRouter, residentPortalService", "UAT-07, UAT-08"],
            ["FR-10", "StoreLedger, ResidentStore", "storeLedgerRouter, storeDutyAccessService", "UAT-09"],
        ],
        [1.0, 1.7, 2.25, 1.55],
    )

    h1(doc, "11. Giả định, ràng buộc và rủi ro")
    matrix(
        doc,
        ["Loại", "Nội dung", "Cách xử lý"],
        [
            ["Giả định", "Demo dùng database seed, chưa phải dữ liệu thật.", "Tạo seed repeatable và smoke test sau seed."],
            ["Giả định", "Simple mode là luồng chính khi demo.", "Ẩn module P2/Store khỏi demo chính nếu chưa cần."],
            ["Ràng buộc", "Một số module cũ/orphan vẫn còn trong repo.", "Không đưa vào menu chính cho đến khi có quyết định nối/archive."],
            ["Rủi ro", "Thiếu dữ liệu resident portal làm portal trống.", "Seed 2 resident demo có user và công tác/tài chính."],
            ["Rủi ro", "Cửa hàng phức tạp làm lệch main flow.", "Chỉ demo cửa hàng sau khi test 1 ngày/2 ca pass."],
            ["Rủi ro", "Encoding/terminal hiển thị mojibake với file cũ.", "Kiểm tra DOCX bằng Python/Word, không dựa duy nhất vào terminal."],
        ],
        [1.1, 3.25, 2.15],
    )

    h1(doc, "12. Tiêu chí nghiệm thu tổng")
    bullet(doc, "Typecheck, test và build pass trước khi bàn giao demo.")
    bullet(doc, "Manager đi hết main flow P0 không gặp lỗi chặn.")
    bullet(doc, "Resident portal có dữ liệu thật theo resident demo, không chỉ là màn hình trống.")
    bullet(doc, "Simple mode chỉ hiển thị main flow đã chốt.")
    bullet(doc, "Bộ tài liệu SRS, Blueprint, User Manual, Deployment Standard lưu trong thư mục docs.")
    bullet(doc, "Checklist dự án cập nhật đúng trạng thái đã làm/chưa làm.")

    save(doc, SRS_OUTPUT, "ResidenceCore Detailed SRS", "Detailed software requirements specification")
    return SRS_OUTPUT


def build_ba_blueprint() -> Path:
    diagrams = build_diagrams()
    doc = Document()
    style_document(doc)
    header_footer(doc, "BA Blueprint", "Business and solution blueprint for ResidenceCore")
    title_block(
        doc,
        "BA Blueprint",
        "Blueprint nghiệp vụ - giải pháp - dữ liệu - triển khai theo chuẩn BA",
        "Business / Solution Blueprint",
        "Baseline dùng để align stakeholder, scope, kiến trúc module và roadmap triển khai",
    )
    intro_document_control(
        doc,
        "ResidenceCore_BA_Blueprint_Detailed.docx",
        "Mô tả bức tranh tổng thể sản phẩm, module, người dùng, dữ liệu, luồng nghiệp vụ, trạng thái triển khai và roadmap.",
    )

    h1(doc, "1. Product vision")
    para(doc, "ResidenceCore hướng tới một hệ thống vận hành lưu xá gọn, rõ, có thể triển khai theo từng lớp. Lớp đầu tiên phải giúp manager quản lý học viên, tổ chức, tài chính, sinh hoạt và công tác. Lớp tiếp theo mở rộng sang portal học viên, cửa hàng, báo cáo và các module chuyên sâu.")
    callout(doc, "Định hướng triển khai", "Không mở tất cả module cùng lúc. Demo và triển khai ban đầu ưu tiên main flow P0, dữ liệu seed đủ, UI premium nhưng không làm rối người dùng.", fill=GREEN)

    h1(doc, "2. Business capability map")
    matrix(
        doc,
        ["Capability", "Module", "Mức ưu tiên", "Giá trị nghiệp vụ"],
        [
            ["Quản trị truy cập", "Auth, Users, Roles", "P0", "Đăng nhập, phân quyền, vận hành an toàn."],
            ["Quản lý cư trú", "Members, Rooms", "P0/P1", "Nắm danh sách học viên, trạng thái, phòng/tổ."],
            ["Tổ chức nội bộ", "Organization", "P0", "Quản lý nhiệm kỳ, Tổ/Ban, chức vụ, bổ nhiệm."],
            ["Tài chính", "FinanceLite", "P0", "Theo dõi khoản phải thu, đã thu, còn nợ."],
            ["Sinh hoạt", "DailyRoutine", "P0", "Theo dõi lịch sinh hoạt và công tác trong ngày."],
            ["Công tác", "Duties", "P0", "Phân công, cập nhật, kiểm soát việc hằng ngày."],
            ["Portal học viên", "Resident pages", "P0", "Học viên tự xem dữ liệu cá nhân và công tác."],
            ["Cửa hàng", "StoreLedger, ResidentStore", "P1", "Quản lý sản phẩm, nhập/bán, ca trực, chốt sổ."],
            ["Mở rộng", "Activities, Discipline, Reports, Skills", "P2", "Mở rộng khi dữ liệu và quy trình đã ổn."],
        ],
        [1.7, 1.55, 0.9, 2.35],
    )

    h1(doc, "3. User journey tổng quan")
    h2(doc, "3.1 Manager journey")
    figure(
        doc,
        diagrams["manager_workflow"],
        "Hình 1. Manager workflow theo main flow P0.",
    )
    process_table(doc, [
        ["1", "Manager", "Đăng nhập.", "Vào Dashboard."],
        ["2", "Manager", "Xem tổng quan.", "Biết tình hình lưu xá."],
        ["3", "Manager", "Kiểm tra học viên và tổ chức.", "Dữ liệu nền đã rõ."],
        ["4", "Manager", "Theo dõi tài chính.", "Biết khoản thu và trạng thái thanh toán."],
        ["5", "Manager", "Theo dõi sinh hoạt/công tác.", "Biết ai làm gì hôm nay."],
        ["6", "Manager", "Chạy báo cáo/chốt nếu module sẵn sàng.", "Có dữ liệu bàn giao hoặc UAT."],
    ])
    h2(doc, "3.2 Resident journey")
    process_table(doc, [
        ["1", "Resident", "Đăng nhập portal.", "Vào Hôm nay."],
        ["2", "Resident", "Xem hồ sơ, công tác, tài chính.", "Hiểu trách nhiệm và khoản cần đóng."],
        ["3", "Resident", "Xem thông báo/hoạt động.", "Theo dõi thông tin lưu xá."],
        ["4", "Resident", "Nếu trực cửa hàng, chọn ngày/ca.", "Vào đúng phiên được giao."],
    ])
    h2(doc, "3.3 Resident store shift workflow")
    figure(
        doc,
        diagrams["resident_store_shift"],
        "Hình 2. Workflow resident store: chọn ngày/ca, kiểm tra phân công, phân quyền theo phiên hiện tại.",
    )

    h1(doc, "4. Module blueprint")
    matrix(
        doc,
        ["Module", "Input", "Processing", "Output"],
        [
            ["Auth", "Username/password, role", "Xác thực, tạo session, kiểm tra mustChangePassword", "User context, menu theo role"],
            ["Members", "Hồ sơ học viên, trạng thái, phòng/tổ", "Lọc, tìm kiếm, hiển thị, cập nhật hồ sơ", "Danh sách/hồ sơ học viên"],
            ["Organization", "Nhiệm kỳ, đơn vị, chức vụ, bổ nhiệm", "Liên kết học viên với vai trò tổ chức", "Cơ cấu lưu xá"],
            ["Finance", "Kỳ thu, khoản phải thu, giao dịch", "Tính đã thu/còn lại, chặn amount sai", "Sổ tài chính, trạng thái thanh toán"],
            ["DailyRoutine", "Ngày, lịch, công tác", "Gom dữ liệu sinh hoạt theo thời gian", "Lịch ngày/tuần/tháng"],
            ["Duties", "Mẫu công tác, phân công, trạng thái", "Tạo/cập nhật công tác, lọc theo role", "Danh sách công tác manager/resident"],
            ["Resident Portal", "User resident, dữ liệu liên kết", "Lọc dữ liệu theo resident", "Hôm nay, hồ sơ, công tác, tài chính"],
            ["Store", "Sản phẩm, nhập/bán, ca trực, ledger", "Tính tồn, thu/chi, chốt ngày, post finance", "Tồn kho, sổ cửa hàng, dòng tổng hợp"],
        ],
        [1.3, 1.75, 2.2, 1.25],
    )

    h1(doc, "5. Data blueprint")
    figure(
        doc,
        diagrams["data_flow_context"],
        "Hình 3. Data flow context của ResidenceCore.",
    )
    matrix(
        doc,
        ["Domain", "Dữ liệu lõi", "Quan hệ quan trọng", "Rule cần giữ"],
        [
            ["Identity", "users, roles", "User có role; resident user liên kết hồ sơ", "Role quyết định menu và quyền."],
            ["Residence", "residents, rooms", "Resident có thể gắn phòng/tổ", "Không vượt sức chứa, không dùng inactive sai rule."],
            ["Organization", "terms, units, roles, assignments", "Assignment thuộc nhiệm kỳ/đơn vị/chức vụ", "Giữ lịch sử bổ nhiệm."],
            ["Finance", "periods, receivables, transactions", "Receivable thuộc resident/kỳ; transaction ghi nhận thu", "Amount > 0; tránh trùng sai."],
            ["Duties", "templates, assignments", "Assignment gắn resident/unit/date/shift", "Portal chỉ xem việc của mình."],
            ["Store", "products, movements, ledger, closings, handovers", "Movement ảnh hưởng tồn; ledger ảnh hưởng sổ; closing post finance", "Không sửa phát sinh sau confirmed."],
        ],
        [1.2, 1.9, 2.0, 1.4],
    )

    h1(doc, "6. Navigation blueprint")
    matrix(
        doc,
        ["Chế độ", "Menu hiển thị", "Mục đích"],
        [
            ["Simple Manager", "Dashboard, Học viên, Tổ chức, Tài chính, Sinh hoạt hằng ngày, Công tác", "Demo/vận hành gọn, tránh module chưa chốt."],
            ["Detailed Manager", "Thêm Phòng ở, Cửa hàng, Users/Roles và các module chuẩn bị sau", "Quản trị hoặc demo mở rộng."],
            ["Resident", "Hôm nay, Hồ sơ, Công tác, Cửa hàng, Tài chính, Thông báo, Hoạt động", "Portal cá nhân của học viên."],
        ],
        [1.45, 3.4, 1.65],
    )

    h1(doc, "7. Deployment blueprint")
    figure(
        doc,
        diagrams["ba_delivery_process"],
        "Hình 4. Process triển khai theo vai trò BA/PO, Dev, QA/UAT và Ops.",
    )
    process_table(doc, [
        ["1", "Developer", "Chốt source code và chạy check/test/build.", "Artifact đủ điều kiện deploy."],
        ["2", "Operator", "Chuẩn bị database và env.", "Môi trường sẵn sàng."],
        ["3", "Developer/Operator", "Chạy migration và seed manager/demo.", "DB có dữ liệu P0."],
        ["4", "Operator", "Start server.", "Ứng dụng online."],
        ["5", "QA", "Smoke test browser.", "Route P0 pass."],
        ["6", "PO/BA", "UAT theo kịch bản.", "Chốt pass/fail và backlog."],
    ])

    h1(doc, "8. Roadmap triển khai")
    matrix(
        doc,
        ["Giai đoạn", "Nội dung", "Điều kiện hoàn tất"],
        [
            ["Phase 1 - Demo P0", "Main flow manager + portal resident, seed demo, bộ tài liệu BA.", "SRS/Blueprint/User Manual/Deployment Standard có trong docs; smoke pass."],
            ["Phase 2 - Resident portal chuẩn", "Dữ liệu resident demo, e2e portal, công tác/tài chính/thông báo ổn.", "Resident login test pass."],
            ["Phase 3 - Store controlled demo", "Kịch bản 1 ngày/2 ca, nhập/bán/bàn giao/chốt/đẩy finance.", "Test cửa hàng pass."],
            ["Phase 4 - Production readiness", "Backup, monitoring, env secret, UAT chính thức, rollback plan.", "Checklist release pass."],
            ["Phase 5 - Modules mở rộng", "Reports, discipline, skills, parents, smart assignment.", "Chốt scope từng module."],
        ],
        [1.6, 3.0, 1.9],
    )

    h1(doc, "9. Open items cần chốt")
    bullet(doc, "Seed demo repeatable thành command chuẩn, không chỉ script SQL rời.")
    bullet(doc, "Tạo resident demo đủ dữ liệu để test portal và cửa hàng nếu cần.")
    bullet(doc, "Bổ sung e2e smoke login + mở route P0.")
    bullet(doc, "Quyết định module cửa hàng có nằm trong demo chính hay demo phụ.")
    bullet(doc, "Chuẩn hóa env.example và guide deploy khi đưa lên demo public.")

    save(doc, BLUEPRINT_OUTPUT, "ResidenceCore BA Blueprint Detailed", "Business and solution blueprint")
    return BLUEPRINT_OUTPUT


def build_user_manual() -> Path:
    diagrams = build_diagrams()
    doc = Document()
    style_document(doc)
    header_footer(doc, "User Manual", "Detailed user manual for ResidenceCore")
    title_block(
        doc,
        "User Guide / User Manual",
        "Tài liệu hướng dẫn sử dụng chi tiết cho manager, học viên và người vận hành demo",
        "User Manual",
        "Dùng cho đào tạo, demo, UAT và bàn giao vận hành ban đầu",
    )
    intro_document_control(
        doc,
        "ResidenceCore_User_Manual_Detailed.docx",
        "Hướng dẫn người dùng thao tác từng màn hình chính của ResidenceCore, kèm lưu ý, lỗi thường gặp và checklist trước demo.",
    )

    callout(doc, "Cách đọc nhanh", "Manager đọc mục 1-9. Học viên đọc mục 10. Người demo đọc thêm mục 11-13. Khi gặp lỗi, xem mục 14 trước khi báo kỹ thuật.", fill=GREEN)

    h1(doc, "1. Bắt đầu sử dụng")
    h2(doc, "1.1 Quy trình đăng nhập")
    numbered_step_table(doc, [
        ["1.1", "Người dùng", "Mở link ứng dụng ResidenceCore.", "Màn hình đăng nhập hiển thị."],
        ["1.2", "Người dùng", "Nhập username và password.", "Nút đăng nhập sẵn sàng."],
        ["1.3", "Người dùng", "Bấm Đăng nhập.", "Hệ thống kiểm tra tài khoản và role."],
        ["1.4", "Người dùng", "Nếu được yêu cầu đổi mật khẩu lần đầu, nhập mật khẩu hiện tại và mật khẩu mới.", "Mật khẩu mới được lưu."],
        ["1.5", "Người dùng", "Đăng nhập lại sau khi đổi mật khẩu.", "Vào Dashboard hoặc portal đúng vai trò."],
    ])
    number(doc, "Mở link ứng dụng ResidenceCore.")
    number(doc, "Tại màn hình đăng nhập, nhập username và password.")
    number(doc, "Nếu là demo manager, dùng admin / Admin@123.")
    number(doc, "Nếu được yêu cầu đổi mật khẩu lần đầu, nhập mật khẩu hiện tại và mật khẩu mới.")
    number(doc, "Sau khi đổi mật khẩu, đăng nhập lại bằng mật khẩu mới.")
    callout(doc, "Lưu ý bảo mật", "Không gửi mật khẩu thật trong group chat hoặc tài liệu public. Tài khoản demo chỉ dùng cho môi trường demo.", fill=AMBER)

    h1(doc, "2. Hiểu menu và chế độ hiển thị")
    h2(doc, "2.1 Screenshot cần chụp cho User Manual")
    screenshot_backlog(doc)
    h2(doc, "2.2 Ảnh chụp chức năng hiện có")
    actual_screenshot_figures(doc)
    h2(doc, "2.3 Workflow tổng quan")
    figure(
        doc,
        diagrams["manager_workflow"],
        "Hình 1. Workflow manager dùng để đối chiếu khi đọc hướng dẫn từng màn hình.",
    )
    matrix(
        doc,
        ["Chế độ", "Khi nào dùng", "Menu chính"],
        [
            ["Simple mode", "Demo chính hoặc người dùng mới.", "Dashboard, Học viên, Tổ chức, Tài chính, Sinh hoạt, Công tác."],
            ["Detailed mode", "Quản trị hoặc kiểm thử module mở rộng.", "Thêm Phòng ở, Cửa hàng, Users/Roles và các mục sau."],
            ["Resident portal", "Học viên đăng nhập.", "Hôm nay, Hồ sơ, Công tác, Cửa hàng, Tài chính, Thông báo, Hoạt động."],
        ],
        [1.35, 2.1, 3.05],
    )

    h1(doc, "3. Dashboard")
    h2(doc, "3.1 Các bước thao tác Dashboard")
    numbered_step_table(doc, [
        ["3.1", "Manager", "Đăng nhập bằng tài khoản manager.", "Vào Dashboard."],
        ["3.2", "Manager", "Quan sát các chỉ số tổng quan ở đầu trang.", "Biết nhanh tình hình lưu xá."],
        ["3.3", "Manager", "Kiểm tra khối công tác hoặc cảnh báo trong ngày.", "Biết việc cần xử lý trước."],
        ["3.4", "Manager", "Bấm vào lối đi nhanh hoặc menu liên quan.", "Chuyển sang module cần xử lý."],
    ])
    para(doc, "Dashboard là màn hình đầu tiên sau khi manager đăng nhập. Mục tiêu là trả lời nhanh: hôm nay lưu xá đang có gì cần chú ý.")
    bullet(doc, "Xem chỉ số tổng quan học viên, phòng ở, công tác và tài chính.")
    bullet(doc, "Dùng các khối thông tin để đi nhanh sang module liên quan.")
    bullet(doc, "Nếu Dashboard trống, kiểm tra seed demo hoặc quyền người dùng.")

    h1(doc, "4. Học viên")
    h2(doc, "4.1 Xem danh sách")
    numbered_step_table(doc, [
        ["4.1", "Manager", "Chọn menu Học viên.", "Danh sách học viên hiển thị."],
        ["4.2", "Manager", "Chọn dạng thẻ hoặc danh sách.", "Giao diện đổi đúng kiểu xem."],
        ["4.3", "Manager", "Nhập từ khóa tìm kiếm hoặc chọn bộ lọc.", "Danh sách được lọc đúng."],
        ["4.4", "Manager", "Bấm vào một học viên.", "Mở hồ sơ chi tiết."],
        ["4.5", "Manager", "Kiểm tra trạng thái, phòng/tổ, tài khoản portal.", "Hồ sơ đủ điều kiện dùng cho flow liên quan."],
    ])
    number(doc, "Chọn menu Học viên.")
    number(doc, "Chọn dạng thẻ hoặc danh sách nếu màn hình hỗ trợ.")
    number(doc, "Dùng tìm kiếm/lọc để tìm học viên theo tên, trạng thái hoặc thông tin liên quan.")
    number(doc, "Bấm vào học viên để xem chi tiết.")
    h2(doc, "4.2 Kiểm tra hồ sơ")
    bullet(doc, "Thông tin cá nhân: họ tên, mã, trạng thái, liên hệ.")
    bullet(doc, "Thông tin cư trú: phòng/tổ nếu có.")
    bullet(doc, "Thông tin portal: tài khoản học viên đã liên kết chưa.")
    bullet(doc, "Nếu chuẩn bị demo portal, luôn chọn một học viên có tài khoản resident active.")
    h2(doc, "4.3 Lỗi thường gặp")
    matrix(
        doc,
        ["Tình huống", "Nguyên nhân thường gặp", "Cách xử lý"],
        [
            ["Không thấy học viên", "DB demo chưa seed hoặc filter đang bật.", "Xóa filter, kiểm tra seed demo."],
            ["Học viên không vào portal", "Chưa có user resident hoặc chưa liên kết hồ sơ.", "Tạo/link user resident."],
            ["Thông tin hiển thị thiếu", "Seed demo tối giản.", "Bổ sung dữ liệu demo cần trình bày."],
        ],
        [1.55, 2.4, 2.55],
    )

    h1(doc, "5. Tổ chức lưu xá")
    h2(doc, "5.1 Các bước xem và kiểm tra tổ chức")
    numbered_step_table(doc, [
        ["5.1", "Manager", "Chọn menu Tổ chức lưu xá.", "Màn hình tổ chức hiển thị."],
        ["5.2", "Manager", "Kiểm tra nhiệm kỳ active.", "Biết cơ cấu hiện hành."],
        ["5.3", "Manager", "Xem Tổ/Ban và chức vụ.", "Cấu trúc tổ chức rõ ràng."],
        ["5.4", "Manager", "Mở danh sách bổ nhiệm.", "Biết học viên nào đang giữ vai trò nào."],
        ["5.5", "Manager", "Kết thúc hiệu lực khi vai trò không còn áp dụng.", "Giữ được lịch sử bổ nhiệm."],
    ])
    number(doc, "Chọn menu Tổ chức lưu xá.")
    number(doc, "Xem nhiệm kỳ hiện tại.")
    number(doc, "Xem các Tổ/Ban và chức vụ.")
    number(doc, "Xem danh sách bổ nhiệm để biết học viên đang giữ vai trò nào.")
    number(doc, "Khi kết thúc vai trò, ưu tiên kết thúc hiệu lực thay vì xóa lịch sử.")
    callout(doc, "Khi demo", "Nên chuẩn bị một nhiệm kỳ active, ít nhất hai đơn vị và vài chức vụ có học viên được bổ nhiệm. Màn hình sẽ có chiều sâu hơn nhiều.", fill=CALLOUT)

    h1(doc, "6. Tài chính lưu xá")
    h2(doc, "6.1 Xem kỳ thu")
    numbered_step_table(doc, [
        ["6.1", "Manager", "Chọn menu Tài chính lưu xá.", "Màn hình tài chính hiển thị."],
        ["6.2", "Manager", "Chọn kỳ thu hoặc khoảng thời gian cần xem.", "Dữ liệu tài chính được lọc đúng."],
        ["6.3", "Manager", "Chọn học viên/khoản phải thu cần ghi nhận.", "Form thanh toán mở đúng đối tượng."],
        ["6.4", "Manager", "Nhập số tiền, ngày thu, nội dung và ghi chú.", "Dữ liệu hợp lệ trước khi lưu."],
        ["6.5", "Manager", "Bấm Lưu giao dịch.", "Số đã thu/còn lại cập nhật đúng."],
    ])
    number(doc, "Chọn menu Tài chính lưu xá.")
    number(doc, "Chọn kỳ thu cần xem.")
    number(doc, "Đọc các chỉ số tổng: phải thu, đã thu, còn lại nếu màn hình hiển thị.")
    h2(doc, "6.2 Ghi nhận thanh toán")
    number(doc, "Chọn học viên hoặc khoản phải thu.")
    number(doc, "Nhập số tiền thu.")
    number(doc, "Kiểm tra ngày thu, nội dung và ghi chú.")
    number(doc, "Lưu giao dịch.")
    bullet(doc, "Hệ thống cần chặn số tiền <= 0.")
    bullet(doc, "Không nên nhập giao dịch trùng nếu cùng kỳ/cùng nội dung đã có rule chặn.")
    h2(doc, "6.3 Đọc số liệu")
    matrix(
        doc,
        ["Chỉ số", "Ý nghĩa"],
        [
            ["Phải thu", "Tổng số tiền dự kiến thu trong kỳ hoặc theo lọc hiện tại."],
            ["Đã thu", "Tổng số tiền đã ghi nhận thanh toán."],
            ["Còn lại", "Phải thu trừ đã thu."],
            ["Giao dịch", "Các dòng thu/chi đã lưu, dùng để đối chiếu."],
        ],
        [1.65, 4.85],
    )

    h1(doc, "7. Sinh hoạt hằng ngày")
    h2(doc, "7.1 Các bước xem lịch sinh hoạt")
    numbered_step_table(doc, [
        ["7.1", "Manager", "Chọn menu Sinh hoạt hằng ngày.", "Màn hình lịch hiển thị."],
        ["7.2", "Manager", "Chọn ngày bằng date picker.", "Dữ liệu chuyển theo ngày."],
        ["7.3", "Manager", "Xem các khung giờ và công tác trong ngày.", "Biết lịch và người phụ trách."],
        ["7.4", "Manager", "Nếu dữ liệu trống, kiểm tra lại ngày hoặc seed demo.", "Xác định nguyên nhân trống dữ liệu."],
    ])
    number(doc, "Chọn menu Sinh hoạt hằng ngày.")
    number(doc, "Chọn ngày cần xem bằng date picker.")
    number(doc, "Xem các khung giờ, việc trong ngày, hoặc công tác liên quan.")
    number(doc, "Nếu không có dữ liệu, kiểm tra ngày đang chọn và seed demo.")
    bullet(doc, "Màn hình này dùng tốt nhất khi demo câu chuyện: hôm nay lưu xá có lịch gì, ai phụ trách, trạng thái ra sao.")

    h1(doc, "8. Công tác / Trực nhật")
    h2(doc, "8.1 Tạo phân công")
    numbered_step_table(doc, [
        ["8.1", "Manager", "Chọn menu Công tác / Trực nhật.", "Danh sách phân công hiển thị."],
        ["8.2", "Manager", "Chọn tạo phân công.", "Form tạo phân công mở."],
        ["8.3", "Manager", "Chọn mẫu công tác, ngày, ca/thời gian.", "Thông tin công tác hợp lệ."],
        ["8.4", "Manager", "Chọn học viên hoặc đơn vị phụ trách.", "Người/đơn vị được gán vào công tác."],
        ["8.5", "Manager", "Lưu phân công.", "Công tác xuất hiện trong manager và portal nếu liên quan."],
    ])
    number(doc, "Chọn menu Công tác / Trực nhật.")
    number(doc, "Chọn tạo phân công hoặc mở form tương ứng.")
    number(doc, "Chọn mẫu công tác, ngày, ca/thời gian và học viên hoặc đơn vị phụ trách.")
    number(doc, "Kiểm tra thông tin và lưu.")
    h2(doc, "8.2 Cập nhật trạng thái")
    bullet(doc, "Hoàn thành: dùng khi công tác đã thực hiện xong.")
    bullet(doc, "Vắng/không làm: dùng khi người được giao không thực hiện.")
    bullet(doc, "Hủy: dùng khi công tác không còn cần thực hiện.")
    h2(doc, "8.3 Trực cửa hàng")
    figure(
        doc,
        diagrams["resident_store_shift"],
        "Hình 2. Workflow học viên chọn ngày/ca để vào trực cửa hàng.",
    )
    bullet(doc, "Học viên vào cửa hàng bằng cách chọn ngày và ca, không cần token.")
    bullet(doc, "Nếu ngày+ca đúng phân công của học viên, hệ thống cho vào phiên của ngày đó.")
    bullet(doc, "Nếu không phải phiên hiện tại, học viên chỉ được xem hoặc chốt sổ theo rule, không thêm/xóa/sửa giao dịch.")

    h1(doc, "9. Thông báo và hoạt động")
    para(doc, "Thông báo và hoạt động giúp lưu xá truyền đạt thông tin nội bộ. Trong demo P0, phần này dùng ở mức vừa đủ để portal học viên không trống.")
    bullet(doc, "Manager có thể dùng thông báo mẫu để giải thích luồng truyền thông nội bộ.")
    bullet(doc, "Resident xem thông báo/hoạt động liên quan trong portal.")
    bullet(doc, "Không đưa nội dung nhạy cảm hoặc dữ liệu thật vào demo public.")

    h1(doc, "10. Portal học viên")
    h2(doc, "10.1 Hôm nay")
    numbered_step_table(doc, [
        ["10.1", "Resident", "Đăng nhập tài khoản học viên.", "Vào portal resident."],
        ["10.2", "Resident", "Mở Hôm nay.", "Xem thông tin trong ngày."],
        ["10.3", "Resident", "Mở Hồ sơ.", "Xem dữ liệu cá nhân đã liên kết."],
        ["10.4", "Resident", "Mở Công tác.", "Xem việc được giao."],
        ["10.5", "Resident", "Mở Tài chính.", "Xem khoản cần đóng và lịch sử thanh toán."],
        ["10.6", "Resident", "Nếu trực cửa hàng, mở Cửa hàng và chọn ngày/ca.", "Vào đúng phiên nếu được phân công."],
    ])
    bullet(doc, "Xem công tác, lịch hoặc thông báo trong ngày.")
    bullet(doc, "Nếu không thấy dữ liệu, kiểm tra học viên có công tác/ngày hiện tại không.")
    h2(doc, "10.2 Hồ sơ")
    bullet(doc, "Xem thông tin cá nhân đã lưu trong hệ thống.")
    bullet(doc, "Nếu thông tin sai, học viên báo manager cập nhật; resident không tự sửa dữ liệu quản trị.")
    h2(doc, "10.3 Công tác")
    bullet(doc, "Xem việc được giao.")
    bullet(doc, "Cập nhật trạng thái nếu hệ thống cho phép theo rule hiện tại.")
    h2(doc, "10.4 Tài chính")
    bullet(doc, "Xem khoản cần đóng và lịch sử thanh toán.")
    bullet(doc, "Nếu đã đóng nhưng chưa thấy cập nhật, báo manager kiểm tra giao dịch.")
    h2(doc, "10.5 Cửa hàng")
    bullet(doc, "Chỉ dùng khi học viên được phân công trực cửa hàng.")
    bullet(doc, "Chọn đúng ngày và ca để vào phiên.")
    bullet(doc, "Không phải ca của mình thì không được vào; không phải phiên hiện tại thì không được thêm/xóa/sửa giao dịch.")

    h1(doc, "11. Quản lý cửa hàng cho manager")
    para(doc, "Phần này là hướng dẫn mở rộng. Không bắt buộc dùng trong demo main flow nếu mục tiêu là trình bày hệ thống lưu xá cơ bản.")
    matrix(
        doc,
        ["Màn hình", "Dùng để làm gì", "Điều cần nhớ"],
        [
            ["Dữ liệu sản phẩm", "Tạo nhóm hàng, đơn vị tính, sản phẩm, xem tồn và giá.", "Sản phẩm có thể chưa có giá bán ban đầu."],
            ["Mua hàng / Nhập kho", "Tạo phiếu nhập kho đa nguồn.", "Chỉ nguồn mua hàng tự tạo khoản chi."],
            ["Bán hàng", "Tạo phiếu bán theo sản phẩm.", "Chặn bán vượt tồn."],
            ["Tổng hợp thu chi", "Xem sổ phát sinh, báo cáo dòng tiền, chốt ngày.", "Chỉ confirmed mới đẩy sang Finance."],
        ],
        [1.55, 2.55, 2.4],
    )

    h1(doc, "12. Checklist trước khi demo")
    matrix(
        doc,
        ["Mục", "Cách kiểm tra", "Đạt khi"],
        [
            ["Login manager", "Dùng admin / Admin@123.", "Vào được Dashboard."],
            ["Menu Simple", "Nhìn sidebar.", "Không có Phòng ở/Store trong main flow nếu chưa demo."],
            ["Dashboard", "Mở /dashboard.", "Không lỗi, có số liệu."],
            ["Học viên", "Mở /members.", "Có danh sách đẹp, dữ liệu rõ."],
            ["Tổ chức", "Mở /organization.", "Có nhiệm kỳ, Tổ/Ban, chức vụ."],
            ["Tài chính", "Mở /finance.", "Có kỳ thu/khoản thu demo."],
            ["Sinh hoạt", "Mở /daily-routine.", "Có lịch/công tác hoặc trạng thái rõ."],
            ["Công tác", "Mở /duties.", "Có mẫu/phân công demo."],
            ["Resident portal", "Login resident demo.", "Thấy dữ liệu cá nhân, không trống."],
        ],
        [1.55, 2.55, 2.4],
    )

    h1(doc, "13. Kịch bản demo đề xuất")
    h2(doc, "13.1 Process demo theo BA")
    figure(
        doc,
        diagrams["ba_delivery_process"],
        "Hình 3. Process chuẩn để demo, UAT và ghi nhận backlog sau triển khai.",
    )
    number(doc, "Giới thiệu mục tiêu ResidenceCore: quản lý lưu xá tập trung.")
    number(doc, "Đăng nhập manager và xem Dashboard.")
    number(doc, "Mở Học viên, xem một hồ sơ đẹp có dữ liệu rõ.")
    number(doc, "Mở Tổ chức để xem nhiệm kỳ, Tổ/Ban, chức vụ.")
    number(doc, "Mở Tài chính để xem kỳ thu và trạng thái thanh toán.")
    number(doc, "Mở Sinh hoạt hằng ngày và Công tác để xem việc trong ngày.")
    number(doc, "Đăng nhập resident demo để xem portal học viên.")
    number(doc, "Nếu cần, mở Detailed mode và giới thiệu module cửa hàng như phần mở rộng.")

    h1(doc, "14. Lỗi thường gặp và xử lý nhanh")
    matrix(
        doc,
        ["Lỗi/Tình huống", "Nguyên nhân thường gặp", "Cách xử lý"],
        [
            ["Không đăng nhập được", "Sai mật khẩu hoặc seed manager chưa đồng bộ.", "Reset/seed lại manager demo."],
            ["Trang trống", "Chưa seed dữ liệu nghiệp vụ.", "Chạy seed demo hoặc chọn ngày có dữ liệu."],
            ["Resident không thấy hồ sơ", "User chưa liên kết resident.", "Kiểm tra liên kết tài khoản-hồ sơ."],
            ["Menu thiếu", "Đang ở Simple mode hoặc role không đúng.", "Đổi mode/kiểm tra role."],
            ["Không tạo được giao dịch", "Amount sai, ngày đã chốt, hoặc quyền không đủ.", "Kiểm tra thông báo lỗi và trạng thái ngày."],
            ["Cửa hàng không cho thêm giao dịch", "Không phải phiên hiện tại hoặc không đúng ngày/ca.", "Chọn đúng ngày/ca hoặc dùng manager review."],
        ],
        [1.8, 2.35, 2.35],
    )

    h1(doc, "15. Quy tắc sử dụng sau triển khai")
    bullet(doc, "Không dùng tài khoản demo cho dữ liệu production thật.")
    bullet(doc, "Không xóa lịch sử nghiệp vụ nếu có cách kết thúc/đóng trạng thái.")
    bullet(doc, "Luôn kiểm tra ngày đang chọn trước khi nhập công tác, tài chính hoặc cửa hàng.")
    bullet(doc, "Báo lỗi kèm tài khoản, route, thời điểm, thao tác vừa làm và ảnh màn hình nếu có.")
    bullet(doc, "Sau mỗi đợt thay đổi nghiệp vụ, cập nhật lại User Manual và checklist.")

    save(doc, MANUAL_OUTPUT, "ResidenceCore Detailed User Manual", "Detailed user guide and user manual")
    return MANUAL_OUTPUT


def main() -> None:
    for path in (build_srs(), build_ba_blueprint(), build_user_manual()):
        print(path)


if __name__ == "__main__":
    main()
