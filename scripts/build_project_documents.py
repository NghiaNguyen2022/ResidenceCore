from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs"

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "0B2545"
MUTED = "64748B"
LIGHT_GRAY = "F2F4F7"
BLUE_GRAY = "E8EEF5"
CALLOUT = "F4F6F9"
AMBER = "FFF7E6"
GREEN = "EAF7EF"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for key, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{key}"))
        if node is None:
            node = OxmlElement(f"w:{key}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_in: float) -> None:
    cell.width = Inches(width_in)
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(width_in * 1440)))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths: list[float]) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), "9360")
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(int(width * 1440)))
        grid.append(col)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            set_cell_width(cell, widths[index])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_font(run, size=None, bold=None, color=None, name="Calibri") -> None:
    run.font.name = name
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:hAnsi"), name)
    if size:
        run.font.size = Pt(size)
    if bold is not None:
        run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def spacing(paragraph, before=0, after=6, line=1.10) -> None:
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = line


def style_document(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ):
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.10


def header_footer(doc: Document, title: str, footer_text: str) -> None:
    header = doc.sections[0].header.paragraphs[0]
    header.text = f"ResidenceCore | {title}"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    spacing(header, after=0)
    for run in header.runs:
        set_font(run, size=9, color=MUTED)

    footer = doc.sections[0].footer.paragraphs[0]
    footer.text = footer_text
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    spacing(footer, after=0)
    for run in footer.runs:
        set_font(run, size=9, color=MUTED)


def title_block(doc: Document, title: str, subtitle: str, kind: str, status: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    spacing(p, after=3)
    r = p.add_run("ResidenceCore")
    set_font(r, size=26, bold=True, color=INK)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    spacing(p, after=12)
    r = p.add_run(title)
    set_font(r, size=16, bold=True, color=BLUE)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    spacing(p, after=18)
    r = p.add_run(subtitle)
    set_font(r, size=12, color=MUTED)

    table = doc.add_table(rows=4, cols=2)
    set_table_geometry(table, [1.875, 4.625])
    rows = [
        ("Loại tài liệu", kind),
        ("Phiên bản", "v0.1 - Demo preparation"),
        ("Ngày cập nhật", date.today().strftime("%d/%m/%Y")),
        ("Trạng thái", status),
    ]
    for idx, (label, value) in enumerate(rows):
        table.cell(idx, 0).text = label
        table.cell(idx, 1).text = value
        set_cell_shading(table.cell(idx, 0), BLUE_GRAY)
        for cell in table.rows[idx].cells:
            for paragraph in cell.paragraphs:
                spacing(paragraph, after=2)
                for run in paragraph.runs:
                    set_font(run, size=10.5, bold=(cell == table.cell(idx, 0)), color=INK)
    doc.add_paragraph()


def h1(doc: Document, text: str) -> None:
    doc.add_heading(text, level=1)


def h2(doc: Document, text: str) -> None:
    doc.add_heading(text, level=2)


def h3(doc: Document, text: str) -> None:
    doc.add_heading(text, level=3)


def para(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    spacing(p)
    r = p.add_run(text)
    set_font(r, color=INK)


def bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.first_line_indent = Inches(-0.25)
    spacing(p, after=4, line=1.167)
    r = p.add_run(text)
    set_font(r, color=INK)


def number(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.first_line_indent = Inches(-0.25)
    spacing(p, after=4, line=1.167)
    r = p.add_run(text)
    set_font(r, color=INK)


def callout(doc: Document, title: str, body: str, fill: str = CALLOUT) -> None:
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [6.5])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    spacing(p, after=3)
    r = p.add_run(title)
    set_font(r, size=11, bold=True, color=DARK_BLUE)
    p = cell.add_paragraph()
    spacing(p, after=0)
    r = p.add_run(body)
    set_font(r, size=10.5, color=INK)
    doc.add_paragraph()


def matrix(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float], fill: str = LIGHT_GRAY) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_geometry(table, widths)
    for index, header in enumerate(headers):
        cell = table.cell(0, index)
        cell.text = header
        set_cell_shading(cell, fill)
        for p in cell.paragraphs:
            spacing(p, after=0)
            for run in p.runs:
                set_font(run, size=9.5, bold=True, color=INK)
    set_repeat_table_header(table.rows[0])

    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            cells[index].text = value
            for p in cells[index].paragraphs:
                spacing(p, after=0, line=1.08)
                for run in p.runs:
                    set_font(run, size=9.2, color=INK)
    doc.add_paragraph()


def process_table(doc: Document, rows: list[list[str]]) -> None:
    matrix(doc, ["Bước", "Vai trò", "Thao tác", "Kết quả"], rows, [0.55, 1.35, 2.55, 2.05])


def save(doc: Document, path: Path, title: str, subject: str) -> None:
    doc.core_properties.title = title
    doc.core_properties.subject = subject
    doc.core_properties.author = "Codex"
    doc.core_properties.comments = "Generated for ResidenceCore demo preparation."
    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(path)


def build_blueprint() -> Path:
    path = DOCS / "ResidenceCore_Blueprint.docx"
    doc = Document()
    style_document(doc)
    header_footer(doc, "Blueprint", "Blueprint for demo deployment and product alignment")
    title_block(
        doc,
        "Blueprint",
        "Bản thiết kế tổng quan nghiệp vụ, module và luồng triển khai demo",
        "Product / Business Blueprint",
        "Dùng để chốt phạm vi demo, dữ liệu mẫu và thứ tự triển khai",
    )

    callout(
        doc,
        "Mục tiêu của blueprint",
        "Tài liệu này giúp đội phát triển và người quản lý nhìn chung một bức tranh: ResidenceCore đang có những khối nghiệp vụ nào, luồng chính chạy ra sao, dữ liệu nào cần chuẩn bị và phần nào nên để sau demo.",
    )

    h1(doc, "1. Tầm nhìn sản phẩm")
    para(doc, "ResidenceCore là hệ thống quản lý lưu xá tập trung, giúp quản lý học viên, tổ chức nội bộ, tài chính, sinh hoạt hằng ngày, công tác và portal học viên trong một giao diện thống nhất.")
    bullet(doc, "Ưu tiên hiện tại: demo được main flow quản lý trước, chưa cần mở toàn bộ module nâng cao.")
    bullet(doc, "Simple mode dùng cho demo chính: gọn, dễ hiểu, ít điểm gây nhiễu.")
    bullet(doc, "Detailed mode giữ các module phụ hoặc module đang mở rộng.")

    h1(doc, "2. Phạm vi module")
    matrix(
        doc,
        ["Module", "Mục đích", "Trạng thái demo", "Ghi chú"],
        [
            ["Dashboard", "Nhìn nhanh tình hình lưu xá.", "P0 - sẵn sàng", "Dùng làm màn hình mở đầu demo."],
            ["Học viên", "Quản lý hồ sơ, trạng thái, phòng, tài khoản portal.", "P0 - sẵn sàng", "Cần dữ liệu học viên mẫu đầy đủ."],
            ["Tổ chức", "Quản lý nhiệm kỳ, Tổ/Ban, chức vụ, bổ nhiệm.", "P0 - sẵn sàng", "Cần một nhiệm kỳ active."],
            ["Tài chính", "Kỳ thu, khoản phải thu, thu chi.", "P0 - sẵn sàng", "Nên seed một kỳ thu dễ giải thích."],
            ["Sinh hoạt", "Lịch sinh hoạt và lịch công tác theo ngày.", "P0 - sẵn sàng", "Cần công tác hôm nay để màn hình có sức sống."],
            ["Công tác", "Mẫu công tác, phân công, cập nhật trạng thái.", "P0 - sẵn sàng", "Có cả phía manager và resident."],
            ["Portal học viên", "Học viên xem thông tin cá nhân, công tác, tài chính, thông báo.", "P0 - có nền", "Cần tài khoản resident demo chuẩn."],
            ["Cửa hàng", "Hàng hóa, ca trực, nhập bán, chốt sổ.", "P1 cho demo chung", "Chỉ demo riêng khi kịch bản 2 ca đã đủ."],
            ["Học tập/Kỹ năng/Phụng vụ/Báo cáo", "Mở rộng quản lý đào tạo và báo cáo.", "P2", "Giữ ẩn hoặc disabled trong Simple mode."],
        ],
        [1.35, 2.1, 1.35, 1.7],
    )

    h1(doc, "3. Kiến trúc nghiệp vụ ở mức cao")
    matrix(
        doc,
        ["Lớp", "Thành phần", "Vai trò"],
        [
            ["Người dùng", "Manager, Resident", "Quyết định menu, quyền thao tác và dữ liệu nhìn thấy."],
            ["Giao diện", "React pages, Simple/Detailed navigation, shared style components", "Trình bày workflow theo vai trò."],
            ["API", "tRPC routers: auth, dashboard, members, organization, finance, dailyRoutine, duties, residentPortal", "Cầu nối nghiệp vụ giữa UI và database."],
            ["Dữ liệu", "MySQL/Drizzle schema, migrations, seed demo", "Lưu hồ sơ, phân công, tài chính, tổ chức, phiên đăng nhập."],
            ["Kiểm soát", "RBAC, session, mustChangePassword, display mode", "Giữ demo an toàn và đúng phạm vi."],
        ],
        [1.15, 2.55, 2.8],
    )

    h1(doc, "4. Main flow blueprint")
    h2(doc, "4.1 Manager flow")
    process_table(doc, [
        ["1", "Manager", "Đăng nhập và đổi mật khẩu lần đầu nếu được yêu cầu.", "Vào Dashboard với quyền manager."],
        ["2", "Manager", "Xem Dashboard.", "Nắm tổng quan học viên, phòng, công tác, cảnh báo."],
        ["3", "Manager", "Mở Học viên.", "Xem danh sách, mở hồ sơ, kiểm tra tài khoản portal."],
        ["4", "Manager", "Mở Tổ chức.", "Xem nhiệm kỳ, Tổ/Ban, chức vụ và bổ nhiệm."],
        ["5", "Manager", "Mở Tài chính.", "Xem kỳ thu, khoản phải thu, ghi nhận thu chi."],
        ["6", "Manager", "Mở Sinh hoạt và Công tác.", "Theo dõi lịch ngày và phân công việc."],
    ])

    h2(doc, "4.2 Resident flow")
    process_table(doc, [
        ["1", "Resident", "Đăng nhập tài khoản đã liên kết học viên.", "Vào portal học viên."],
        ["2", "Resident", "Mở Hôm nay.", "Xem lịch, công tác, thông báo cần chú ý."],
        ["3", "Resident", "Mở Công tác.", "Xem và cập nhật công tác cá nhân nếu được phép."],
        ["4", "Resident", "Mở Tài chính.", "Xem khoản cần đóng và lịch sử thanh toán."],
        ["5", "Resident", "Mở Thông báo / Hoạt động.", "Theo dõi tin nội bộ và sinh hoạt chung."],
    ])

    h1(doc, "5. Dữ liệu demo cần có")
    matrix(
        doc,
        ["Nhóm dữ liệu", "Tối thiểu cần seed", "Lý do"],
        [
            ["Tài khoản", "1 manager, 2 resident", "Đủ demo cả quản lý và portal học viên."],
            ["Học viên", "8-15 học viên với trạng thái và phòng khác nhau", "Danh sách nhìn thật, có chiều sâu."],
            ["Phòng/Tổ", "3-5 phòng, 2 Tổ, 2 Ban", "Dashboard và tổ chức có dữ liệu."],
            ["Tổ chức", "1 nhiệm kỳ active, vài chức vụ đang hiệu lực", "Demo cơ cấu lưu xá rõ ràng."],
            ["Tài chính", "1 kỳ thu, vài khoản đã thu/còn nợ", "Demo được trạng thái tài chính."],
            ["Sinh hoạt/Công tác", "Mẫu lịch ngày và công tác hôm nay", "Màn hình sinh hoạt không trống."],
            ["Thông báo", "2-3 thông báo mẫu", "Portal học viên có nội dung."],
        ],
        [1.45, 2.45, 2.6],
    )

    h1(doc, "6. Quyết định trước deploy demo")
    bullet(doc, "Simple mode chỉ hiển thị main flow P0.")
    bullet(doc, "Không đưa module P2 vào demo chính nếu chưa có dữ liệu và test.")
    bullet(doc, "Seed demo phải repeatable, chạy được sau migration trên DB mới.")
    bullet(doc, "Tài khoản demo manager thống nhất: admin / Admin@123.")
    bullet(doc, "Cửa hàng nên là demo phụ, không kéo vào luồng chính nếu chưa chốt kịch bản 2 ca.")

    h1(doc, "7. Roadmap sau demo")
    matrix(
        doc,
        ["Ưu tiên", "Việc cần làm", "Kết quả mong muốn"],
        [
            ["P0", "Seed demo repeatable và smoke test từ DB sạch.", "Deploy demo ổn định."],
            ["P0", "User guide dễ hiểu cho manager/resident.", "Người demo có tài liệu đi theo."],
            ["P1", "E2E smoke test tự động.", "Giảm rủi ro regressions."],
            ["P1", "Polish responsive/mobile.", "Demo tốt trên nhiều màn hình."],
            ["P2", "Nối module học tập/kỹ năng/phụng vụ/báo cáo.", "Mở rộng sau khi main flow ổn."],
        ],
        [0.9, 3.0, 2.6],
    )

    save(doc, path, "ResidenceCore Blueprint", "Product and business blueprint")
    return path


def build_user_guide() -> Path:
    path = DOCS / "ResidenceCore_User_Guide.docx"
    doc = Document()
    style_document(doc)
    header_footer(doc, "User Guide", "User manual for demo users")
    title_block(
        doc,
        "User Guide / User Manual",
        "Hướng dẫn sử dụng dễ hiểu cho quản lý lưu xá và học viên",
        "User Guide",
        "Dành cho người dùng demo, viết theo từng việc cần làm",
    )

    callout(
        doc,
        "Cách dùng tài liệu này",
        "Bạn không cần đọc hết từ đầu đến cuối. Hãy mở đúng mục mình cần: đăng nhập, xem dashboard, quản lý học viên, tài chính, sinh hoạt, công tác hoặc portal học viên.",
        fill=GREEN,
    )

    h1(doc, "1. Đăng nhập")
    number(doc, "Mở ứng dụng ResidenceCore.")
    number(doc, "Chọn Đăng nhập.")
    number(doc, "Nhập tài khoản demo manager: admin.")
    number(doc, "Nhập mật khẩu demo: Admin@123.")
    number(doc, "Nếu hệ thống yêu cầu đổi mật khẩu lần đầu, nhập mật khẩu hiện tại và mật khẩu mới, sau đó đăng nhập lại.")
    callout(doc, "Lưu ý", "Nếu không đăng nhập được, hãy kiểm tra lại dữ liệu seed demo. Tài khoản demo chuẩn hiện tại là admin / Admin@123.", fill=AMBER)

    h1(doc, "2. Màn hình Dashboard")
    para(doc, "Dashboard là nơi xem nhanh tình hình lưu xá trong ngày.")
    bullet(doc, "Tổng số học viên đang lưu trú.")
    bullet(doc, "Tình trạng phòng ở.")
    bullet(doc, "Công tác hoặc việc cần chú ý trong ngày.")
    bullet(doc, "Các số liệu giúp quản lý đi nhanh sang màn hình cần xử lý.")

    h1(doc, "3. Quản lý học viên")
    h2(doc, "Xem danh sách học viên")
    number(doc, "Vào menu Học viên.")
    number(doc, "Chọn dạng Thẻ hoặc List tùy cách xem.")
    number(doc, "Dùng ô tìm kiếm hoặc bộ lọc nếu cần.")
    number(doc, "Bấm Xem để mở hồ sơ chi tiết.")
    h2(doc, "Khi xem hồ sơ học viên")
    bullet(doc, "Kiểm tra thông tin cá nhân và trạng thái lưu trú.")
    bullet(doc, "Kiểm tra phòng ở, liên hệ gia đình, tài khoản portal.")
    bullet(doc, "Nếu hồ sơ còn thiếu thông tin, cập nhật trước khi dùng cho các phân công khác.")

    h1(doc, "4. Quản lý tổ chức lưu xá")
    number(doc, "Vào menu Tổ chức lưu xá.")
    number(doc, "Xem nhiệm kỳ hiện tại.")
    number(doc, "Xem các Tổ, Ban và chức vụ đang có.")
    number(doc, "Khi cần, dùng Bổ nhiệm để gán học viên vào chức vụ hoặc đơn vị.")
    number(doc, "Nếu một bổ nhiệm không còn hiệu lực, dùng thao tác kết thúc thay vì xóa mất lịch sử.")

    h1(doc, "5. Quản lý tài chính")
    h2(doc, "Xem kỳ thu")
    number(doc, "Vào menu Tài chính lưu xá.")
    number(doc, "Chọn kỳ thu đang cần theo dõi.")
    number(doc, "Xem danh sách khoản phải thu của học viên.")
    h2(doc, "Ghi nhận thanh toán")
    number(doc, "Chọn học viên hoặc khoản phải thu cần ghi nhận.")
    number(doc, "Nhập số tiền thu.")
    number(doc, "Kiểm tra số tiền không vượt số còn lại.")
    number(doc, "Lưu giao dịch.")
    callout(doc, "Nguyên tắc dễ nhớ", "Khoản thu của học viên đi theo kỳ thu. Thu chi ngoài học viên dùng cho các khoản vận hành chung, không trộn lẫn khi không cần.", fill=CALLOUT)

    h1(doc, "6. Sinh hoạt hằng ngày")
    number(doc, "Vào menu Sinh hoạt hằng ngày.")
    number(doc, "Chọn Hôm nay, Ngày, Tuần hoặc Tháng.")
    number(doc, "Xem các khung giờ sinh hoạt và công tác trong ngày.")
    number(doc, "Nếu công tác chưa có người, chọn Phân công.")
    bullet(doc, "Màn hình này dùng để trả lời câu hỏi: hôm nay lưu xá có việc gì, ai làm, lúc nào.")

    h1(doc, "7. Công tác / Trực nhật")
    h2(doc, "Tạo phân công")
    number(doc, "Vào menu Công tác / Trực nhật.")
    number(doc, "Chọn Tạo phân công.")
    number(doc, "Chọn mẫu công tác, ngày, thời gian và người hoặc đơn vị phụ trách.")
    number(doc, "Kiểm tra lại thông tin rồi lưu.")
    h2(doc, "Cập nhật trạng thái")
    bullet(doc, "Hoàn thành: dùng khi học viên đã làm xong.")
    bullet(doc, "Vắng / Không làm: dùng khi học viên không thực hiện.")
    bullet(doc, "Hủy: dùng khi công tác không còn cần thực hiện.")

    h1(doc, "8. Portal học viên")
    h2(doc, "Học viên xem gì?")
    bullet(doc, "Hôm nay: việc cần làm và thông tin trong ngày.")
    bullet(doc, "Hồ sơ: thông tin cá nhân đang lưu trong hệ thống.")
    bullet(doc, "Công tác: các công việc được giao.")
    bullet(doc, "Tài chính: khoản cần đóng và lịch sử thanh toán.")
    bullet(doc, "Thông báo: tin nhắn nội bộ từ lưu xá.")
    bullet(doc, "Hoạt động: sự kiện hoặc sinh hoạt chung.")
    h2(doc, "Khi học viên không thấy dữ liệu")
    bullet(doc, "Kiểm tra tài khoản học viên đã liên kết đúng hồ sơ resident chưa.")
    bullet(doc, "Kiểm tra học viên còn trạng thái đang lưu trú hay không.")
    bullet(doc, "Kiểm tra dữ liệu demo đã seed cho ngày hiện tại chưa.")

    h1(doc, "9. Cửa hàng lưu xá trong demo")
    para(doc, "Phần cửa hàng có thể demo riêng sau khi dữ liệu ca trực đã chuẩn. Trong demo main flow, không bắt buộc phải đi sâu vào cửa hàng.")
    bullet(doc, "Nếu học viên được phân công trực cửa hàng, học viên chọn đúng ngày và ca để vào phiên.")
    bullet(doc, "Nếu không phải phiên hiện tại, học viên chỉ được xem hoặc chốt sổ theo quyền, không thêm/xóa/sửa giao dịch.")
    bullet(doc, "Quản lý dùng Store ledger để xem hàng hóa, nhập bán và tổng hợp thu chi.")

    h1(doc, "10. Checklist trước khi demo")
    matrix(
        doc,
        ["Việc cần kiểm tra", "Cách kiểm tra nhanh", "Đạt khi"],
        [
            ["Đăng nhập", "Dùng admin / Admin@123.", "Vào được Dashboard."],
            ["Menu Simple", "Nhìn sidebar sau login.", "Chỉ thấy main flow cần demo."],
            ["Học viên", "Mở danh sách học viên.", "Có dữ liệu, không trống."],
            ["Tổ chức", "Mở cơ cấu tổ chức.", "Có nhiệm kỳ, Tổ/Ban, chức vụ."],
            ["Tài chính", "Mở kỳ thu.", "Có khoản phải thu và trạng thái thanh toán."],
            ["Sinh hoạt", "Mở hôm nay.", "Có lịch hoặc công tác."],
            ["Portal học viên", "Đăng nhập resident demo.", "Thấy Hôm nay, Công tác, Tài chính."],
        ],
        [1.65, 2.65, 2.2],
    )

    h1(doc, "11. Câu hỏi thường gặp")
    h3(doc, "Tại sao không thấy menu Phòng ở trong Simple mode?")
    para(doc, "Vì demo hiện ưu tiên main flow. Phòng ở vẫn còn route trong Detailed mode, nhưng không đưa vào Simple để tránh người dùng đi lệch luồng.")
    h3(doc, "Tại sao một số menu có nhãn Sau?")
    para(doc, "Đó là các phần sẽ làm sau hoặc chưa phải trọng tâm demo. Người demo nên bỏ qua các mục này.")
    h3(doc, "Nếu trang bị trống thì sao?")
    para(doc, "Thường là do DB demo chưa seed đủ dữ liệu. Cần chạy seed demo repeatable trước khi demo.")
    h3(doc, "Nếu học viên không vào được portal?")
    para(doc, "Kiểm tra user học viên đã được tạo, có role resident và đã liên kết với hồ sơ học viên đang hoạt động.")

    h1(doc, "12. Gợi ý thứ tự trình bày")
    number(doc, "Giới thiệu Dashboard.")
    number(doc, "Mở Học viên và xem một hồ sơ.")
    number(doc, "Mở Tổ chức để xem Tổ/Ban và chức vụ.")
    number(doc, "Mở Tài chính để xem kỳ thu và ghi nhận thanh toán.")
    number(doc, "Mở Sinh hoạt và Công tác để xem việc trong ngày.")
    number(doc, "Đăng nhập học viên để xem portal.")

    save(doc, path, "ResidenceCore User Guide", "User manual for demo users")
    return path


def main() -> None:
    blueprint = build_blueprint()
    guide = build_user_guide()
    print(blueprint)
    print(guide)


if __name__ == "__main__":
    main()
