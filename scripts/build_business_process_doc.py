from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "docs" / "ResidenceCore_Business_Process_Document.docx"

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "0B2545"
MUTED = "64748B"
LIGHT_GRAY = "F2F4F7"
BLUE_GRAY = "E8EEF5"
CALLOUT = "F4F6F9"
AMBER = "FFF7E6"
GREEN = "EAF7EF"
RED = "FDECEC"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_in: float) -> None:
    cell.width = Inches(width_in)
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(width_in * 1440)))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths: list[float]) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), "9360")
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(int(width * 1440)))
        grid.append(col)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            set_cell_width(cell, widths[index])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_font(run, name="Calibri", size=None, bold=None, color=None) -> None:
    run.font.name = name
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:hAnsi"), name)
    if size:
        run.font.size = Pt(size)
    if bold is not None:
        run.font.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def set_paragraph_spacing(paragraph, before=0, after=6, line=1.10) -> None:
    pf = paragraph.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line


def style_document(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ):
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.10


def add_title(doc: Document) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, before=0, after=3)
    r = p.add_run("ResidenceCore")
    set_font(r, size=26, bold=True, color=INK)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, before=0, after=12)
    r = p.add_run("Business Process Document")
    set_font(r, size=16, bold=True, color=BLUE)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, before=0, after=18)
    r = p.add_run("Tài liệu nghiệp vụ chuẩn bị deploy demo")
    set_font(r, size=12, color=MUTED)

    meta = doc.add_table(rows=4, cols=2)
    set_table_geometry(meta, [1.875, 4.625])
    rows = [
        ("Phiên bản", "v0.1 - Demo readiness"),
        ("Ngày cập nhật", date.today().strftime("%d/%m/%Y")),
        ("Phạm vi", "Main flow quản lý lưu xá, portal học viên, module nghiệp vụ đang hiện trong demo"),
        ("Trạng thái", "Đủ nền để demo main flow; còn thiếu seed demo repeatable và smoke test từ DB sạch"),
    ]
    for i, (label, value) in enumerate(rows):
        meta.cell(i, 0).text = label
        meta.cell(i, 1).text = value
        set_cell_shading(meta.cell(i, 0), BLUE_GRAY)
        for cell in meta.rows[i].cells:
            for paragraph in cell.paragraphs:
                set_paragraph_spacing(paragraph, after=2, line=1.10)
                for run in paragraph.runs:
                    set_font(run, size=10.5, bold=(cell == meta.cell(i, 0)), color=INK)

    doc.add_paragraph()
    callout(
        doc,
        "Tóm tắt điều hành",
        "ResidenceCore hiện đã có main flow quản lý đủ để trình diễn: đăng nhập, dashboard, học viên, tổ chức, tài chính, sinh hoạt, công tác và portal học viên. Trước khi deploy demo sạch, cần đóng gói seed demo repeatable, đồng bộ tài khoản demo, kiểm tra DB mới tinh và giữ các module chưa chốt ở trạng thái ẩn hoặc disabled.",
        fill=CALLOUT,
    )


def add_header_footer(doc: Document) -> None:
    section = doc.sections[0]
    header = section.header.paragraphs[0]
    header.text = "ResidenceCore | Business Process Document"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_paragraph_spacing(header, after=0)
    for run in header.runs:
        set_font(run, size=9, color=MUTED)

    footer = section.footer.paragraphs[0]
    footer.text = "Internal demo preparation"
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(footer, after=0)
    for run in footer.runs:
        set_font(run, size=9, color=MUTED)


def h1(doc: Document, text: str):
    return doc.add_heading(text, level=1)


def h2(doc: Document, text: str):
    return doc.add_heading(text, level=2)


def h3(doc: Document, text: str):
    return doc.add_heading(text, level=3)


def para(doc: Document, text: str, bold_prefix: str | None = None) -> None:
    p = doc.add_paragraph()
    set_paragraph_spacing(p)
    if bold_prefix and text.startswith(bold_prefix):
        r = p.add_run(bold_prefix)
        set_font(r, bold=True, color=INK)
        r = p.add_run(text[len(bold_prefix):])
        set_font(r, color=INK)
    else:
        r = p.add_run(text)
        set_font(r, color=INK)


def bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.first_line_indent = Inches(-0.25)
    set_paragraph_spacing(p, after=4, line=1.167)
    r = p.add_run(text)
    set_font(r, color=INK)


def numbered(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.first_line_indent = Inches(-0.25)
    set_paragraph_spacing(p, after=4, line=1.167)
    r = p.add_run(text)
    set_font(r, color=INK)


def callout(doc: Document, title: str, body: str, fill: str = CALLOUT) -> None:
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [6.5])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    set_paragraph_spacing(p, after=3, line=1.10)
    r = p.add_run(title)
    set_font(r, size=11, bold=True, color=DARK_BLUE)
    p = cell.add_paragraph()
    set_paragraph_spacing(p, after=0, line=1.10)
    r = p.add_run(body)
    set_font(r, size=10.5, color=INK)
    doc.add_paragraph()


def matrix(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float], header_fill: str = LIGHT_GRAY) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_geometry(table, widths)
    for i, header in enumerate(headers):
        cell = table.cell(0, i)
        cell.text = header
        set_cell_shading(cell, header_fill)
        for p in cell.paragraphs:
            set_paragraph_spacing(p, after=0, line=1.10)
            for r in p.runs:
                set_font(r, size=9.5, bold=True, color=INK)
    set_repeat_table_header(table.rows[0])

    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value
            for p in cells[i].paragraphs:
                set_paragraph_spacing(p, after=0, line=1.08)
                for r in p.runs:
                    set_font(r, size=9.2, color=INK)
    doc.add_paragraph()


def module_table(doc: Document) -> None:
    rows = [
        ["Auth & RBAC", "Đăng nhập local, session, role manager/resident, đổi mật khẩu lần đầu.", "Đủ demo", "Cần seed thống nhất mật khẩu demo."],
        ["Dashboard", "Tổng quan học viên, phòng, công tác, trạng thái vận hành.", "Đủ demo", "Cần DB demo có dữ liệu đủ dày."],
        ["Học viên", "Danh sách, hồ sơ, tài khoản portal, liên hệ, phòng, trạng thái.", "Đủ demo", "Có thể polish thêm form chi tiết sau demo."],
        ["Tổ chức", "Nhiệm kỳ, Tổ/Ban, chức vụ, bổ nhiệm, liên kết portal vai trò.", "Đủ demo", "Cần dữ liệu demo nhất quán nhiệm kỳ active."],
        ["Tài chính", "Kỳ thu, khoản phải thu, ghi nhận thu, thu chi ngoài học viên.", "Đủ demo", "Cần kịch bản demo 1 kỳ thu rõ ràng."],
        ["Sinh hoạt", "Lịch ngày, mẫu sinh hoạt, công tác gắn ngày.", "Đủ demo", "Cần seed công tác hôm nay để màn hình đẹp."],
        ["Công tác", "Mẫu công tác, phân công, cập nhật trạng thái, portal công tác.", "Đủ demo", "Cần test thêm phân công từ DB sạch."],
        ["Portal học viên", "Hôm nay, hồ sơ, công tác, tài chính, thông báo, hoạt động.", "Có nền", "Cần tài khoản resident demo chuẩn."],
        ["Cửa hàng", "Sổ hàng hóa, nhập/bán, trực ca, chốt sổ.", "Đang hoàn thiện", "Không đặt làm trọng tâm demo hệ thống chung nếu chưa xong seed/test."],
    ]
    matrix(doc, ["Module", "Nghiệp vụ chính", "Trạng thái", "Ghi chú demo"], rows, [1.35, 2.55, 1.05, 1.55])


def process_table(doc: Document, rows: list[list[str]]) -> None:
    matrix(doc, ["Bước", "Người thực hiện", "Thao tác", "Kết quả / kiểm soát"], rows, [0.55, 1.35, 2.55, 2.05])


def build_doc() -> None:
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    style_document(doc)
    add_header_footer(doc)
    add_title(doc)

    h1(doc, "1. Mục tiêu và phạm vi")
    para(doc, "Tài liệu này mô tả các quy trình nghiệp vụ chính của ResidenceCore để chuẩn bị deploy demo. Nội dung tập trung vào luồng vận hành lưu xá, không đi sâu vào phần kỹ thuật triển khai mã nguồn.")
    bullet(doc, "Chuẩn hóa cách trình bày business process theo module và vai trò.")
    bullet(doc, "Xác định chức năng đã đủ demo, chức năng còn thiếu và rủi ro trước khi deploy.")
    bullet(doc, "Làm căn cứ cho seed demo, test kịch bản và hướng dẫn người demo.")

    h2(doc, "Phạm vi P0 cho demo")
    module_table(doc)

    h2(doc, "Ngoài phạm vi demo P0")
    para(doc, "Các module học tập nâng cao, kỹ năng, phụng vụ, báo cáo, phụ huynh và smart assignment đang được xem là P1/P2. Trong demo simple mode, các module này nên được ẩn hoặc giữ disabled để tránh làm loãng luồng chính.")

    h1(doc, "2. Vai trò người dùng")
    matrix(
        doc,
        ["Vai trò", "Mục tiêu sử dụng", "Quyền nghiệp vụ demo"],
        [
            ["Manager", "Quản lý toàn bộ vận hành lưu xá.", "Xem dashboard, quản lý học viên, tổ chức, tài chính, sinh hoạt, công tác, người dùng."],
            ["Resident", "Theo dõi việc cá nhân trong portal.", "Xem hôm nay, hồ sơ, công tác, tài chính, thông báo, hoạt động; thao tác cửa hàng nếu được phân công."],
            ["System", "Cấu hình nền và kiểm soát truy cập.", "Xác thực, phân quyền, display mode, session, seed demo."],
        ],
        [1.3, 2.45, 2.75],
    )

    h1(doc, "3. Quy trình nghiệp vụ theo module")

    h2(doc, "3.1 Đăng nhập, phân quyền và chế độ giao diện")
    process_table(doc, [
        ["1", "Manager", "Mở trang login và đăng nhập bằng tài khoản demo.", "Hệ thống tạo session, nạp role manager và quyền tương ứng."],
        ["2", "System", "Nếu tài khoản có mustChangePassword, hiển thị modal đổi mật khẩu.", "Không cho dùng màn hình nghiệp vụ trước khi đổi mật khẩu."],
        ["3", "Manager", "Đổi mật khẩu và đăng nhập lại.", "Session mới hoạt động, vào dashboard."],
        ["4", "System", "Nạp display mode simple hoặc detailed.", "Simple mode chỉ hiển thị main flow đã chốt."],
    ])
    callout(doc, "Kiểm soát demo", "Seed demo phải thống nhất tài khoản admin / Admin@123 và role chính manager. Không dùng lại script cũ đưa admin về role admin hoặc password 123456.", fill=AMBER)

    h2(doc, "3.2 Dashboard vận hành")
    process_table(doc, [
        ["1", "Manager", "Vào Dashboard sau khi đăng nhập.", "Thấy tổng học viên, phòng, trạng thái phòng, công tác hôm nay."],
        ["2", "System", "Tổng hợp số liệu từ residents, rooms, duty assignments và các module liên quan.", "Nếu DB demo đủ dữ liệu, dashboard có nội dung sống thay vì trống."],
        ["3", "Manager", "Dùng dashboard để định hướng demo sang học viên, tổ chức, tài chính, sinh hoạt.", "Dashboard là điểm mở đầu main flow."],
    ])

    h2(doc, "3.3 Quản lý học viên")
    process_table(doc, [
        ["1", "Manager", "Xem danh sách học viên theo thẻ hoặc list.", "Hiển thị trạng thái lưu trú, phòng, tài khoản, thông tin thiếu."],
        ["2", "Manager", "Mở hồ sơ học viên.", "Xem thông tin cá nhân, phòng, liên hệ, học tập, vai trò liên quan."],
        ["3", "Manager", "Tạo hoặc liên kết tài khoản portal cho học viên.", "Resident có thể đăng nhập portal để xem dữ liệu cá nhân."],
        ["4", "System", "Kiểm soát trạng thái inactive, transferred_out, left.", "Không dùng học viên đã rời lưu xá cho phân công hoặc nghiệp vụ active."],
    ])

    h2(doc, "3.4 Tổ chức lưu xá")
    process_table(doc, [
        ["1", "Manager", "Xem nhiệm kỳ active và cơ cấu Tổ/Ban.", "Có sơ đồ tổ chức và danh sách phân công đang hiệu lực."],
        ["2", "Manager", "Bổ nhiệm học viên vào chức vụ hoặc đơn vị.", "Tạo assignment có start date, status và phạm vi."],
        ["3", "System", "Đồng bộ portal vai trò của học viên.", "Resident thấy vai trò, thành viên Tổ/Ban và công tác theo phạm vi."],
        ["4", "Manager", "Kết thúc hoặc chỉnh sửa bổ nhiệm.", "Lịch sử bổ nhiệm được giữ để truy vết."],
    ])

    h2(doc, "3.5 Tài chính lưu xá")
    process_table(doc, [
        ["1", "Manager", "Tạo hoặc chọn kỳ thu.", "Kỳ thu có tháng áp dụng, loại khoản thu và danh sách học viên."],
        ["2", "Manager", "Ghi nhận khoản phải thu / thanh toán.", "Số tiền thu phải lớn hơn 0, không vượt số còn phải thu."],
        ["3", "Manager", "Ghi thu chi ngoài học viên nếu có.", "Dòng tiền được phân loại theo mục đích và nguồn."],
        ["4", "Resident", "Xem tình trạng tài chính cá nhân trong portal.", "Thấy khoản cần đóng, đã đóng, tạm ứng hoặc chi thực tế liên quan."],
    ])
    callout(doc, "Kịch bản demo gợi ý", "Dùng một kỳ thu năm 2026, 3 đến 5 học viên có trạng thái khác nhau: đã thu đủ, còn nợ, thu một phần. Cách này giúp màn hình tài chính có chiều sâu nhưng vẫn dễ giải thích.", fill=GREEN)

    h2(doc, "3.6 Sinh hoạt hằng ngày và công tác")
    process_table(doc, [
        ["1", "Manager", "Xem lịch sinh hoạt theo ngày, tuần, tháng.", "Hiển thị khung giờ sinh hoạt và công tác trong ngày."],
        ["2", "Manager", "Tạo mẫu công tác hoặc dùng mẫu đã seed.", "Mẫu có loại công tác, thời gian, nơi làm, số người."],
        ["3", "Manager", "Phân công học viên, Tổ/Ban hoặc phạm vi phù hợp.", "Assignment được tạo, có trạng thái pending/completed/absent/cancelled."],
        ["4", "Resident", "Vào portal hôm nay hoặc công tác của tôi.", "Thấy công việc được giao và cập nhật trạng thái khi được phép."],
    ])

    h2(doc, "3.7 Portal học viên")
    process_table(doc, [
        ["1", "Resident", "Đăng nhập tài khoản portal đã liên kết hồ sơ.", "Hệ thống xác định resident profile từ userId."],
        ["2", "Resident", "Xem trang Hôm nay.", "Tóm tắt công tác, lịch sinh hoạt, thông báo cần chú ý."],
        ["3", "Resident", "Xem hồ sơ, tài chính, thông báo, hoạt động.", "Mỗi màn hình chỉ hiển thị dữ liệu liên quan đến học viên."],
        ["4", "System", "Chặn nếu user resident chưa có hồ sơ hoặc đã rời lưu xá.", "Trả lỗi rõ ràng, tránh lộ dữ liệu người khác."],
    ])

    h2(doc, "3.8 Cửa hàng lưu xá, trạng thái tích hợp")
    para(doc, "Module cửa hàng đang có nền nghiệp vụ mạnh hơn mức demo hệ thống chung, gồm hàng hóa, nhập bán, ca trực, bàn giao và chốt sổ. Tuy nhiên, nếu mục tiêu là deploy demo main flow, phần này nên để ở detailed mode hoặc demo sau khi seed resident portal và kịch bản 2 ca đã hoàn tất.")
    matrix(
        doc,
        ["Luồng", "Đã có", "Còn thiếu trước demo riêng"],
        [
            ["Manager store ledger", "Danh sách hàng hóa, nhập/bán, tổng hợp thu chi.", "Test lại từ DB sạch và xác nhận dữ liệu mẫu."],
            ["Resident store access", "Chọn ngày + ca, không cần token; non-current chỉ đọc/chốt.", "Tạo tài khoản resident demo và smoke browser thật."],
            ["Chốt sổ", "Logic chốt/review/đẩy Finance đã có nền.", "Kịch bản end-to-end 1 ngày, 2 ca."],
        ],
        [1.6, 2.45, 2.45],
    )

    h1(doc, "4. Luồng demo đề xuất")
    numbered(doc, "Đăng nhập manager bằng tài khoản demo, xử lý đổi mật khẩu nếu hệ thống yêu cầu.")
    numbered(doc, "Mở Dashboard để giới thiệu bức tranh tổng quan lưu xá.")
    numbered(doc, "Vào Học viên, xem danh sách dạng thẻ, mở một hồ sơ tiêu biểu.")
    numbered(doc, "Vào Tổ chức, trình bày nhiệm kỳ, Tổ/Ban, chức vụ và bổ nhiệm.")
    numbered(doc, "Vào Tài chính, trình bày kỳ thu, khoản phải thu và một giao dịch thanh toán.")
    numbered(doc, "Vào Sinh hoạt hằng ngày, xem lịch ngày và các công tác.")
    numbered(doc, "Vào Công tác / Trực nhật, tạo hoặc cập nhật một phân công đơn giản.")
    numbered(doc, "Đăng nhập resident demo để xem portal: Hôm nay, Công tác, Tài chính, Thông báo.")

    h1(doc, "5. Checklist deploy demo")
    matrix(
        doc,
        ["Hạng mục", "Yêu cầu hoàn thành", "Mức ưu tiên"],
        [
            ["Database migration", "Chạy toàn bộ migration cần thiết trên DB demo.", "P0"],
            ["Seed demo", "Có command repeatable tạo manager, học viên, phòng, tổ chức, tài chính, sinh hoạt, công tác.", "P0"],
            ["Login demo", "admin / Admin@123, role manager, simple mode mặc định.", "P0"],
            ["Smoke test browser", "Mở 7 route P0 không 404, không runtime crash, console sạch.", "P0"],
            ["Resident demo", "Có ít nhất 1 tài khoản resident liên kết hồ sơ để demo portal.", "P0"],
            ["Ẩn module chưa chốt", "Simple mode không hiển thị module P1/P2.", "P0"],
            ["Env deploy", "DB, JWT, STORE_ACCESS_SECRET, PORT, analytics optional.", "P0"],
            ["Storage", "Quyết định dùng ảnh base64/url hoặc cấu hình Forge/S3 nếu cần.", "P1"],
            ["Bundle size", "Tối ưu chunk lớn nếu demo public rộng.", "P1"],
        ],
        [1.55, 3.8, 1.15],
    )

    h1(doc, "6. Rủi ro và quyết định cần chốt")
    matrix(
        doc,
        ["Rủi ro", "Tác động", "Khuyến nghị"],
        [
            ["Seed demo cũ lệch tài khoản", "Deploy xong không đăng nhập đúng như UI hướng dẫn.", "Đồng bộ SQL demo với seedDefaultManager hoặc thay bằng TypeScript seed mới."],
            ["DB mới thiếu dữ liệu", "Màn hình render được nhưng trống, demo kém thuyết phục.", "Seed tối thiểu đủ cho 7 route P0."],
            ["Module chưa hoàn thiện xuất hiện trong Simple", "Người xem bấm vào luồng chưa chốt.", "Giữ Simple menu tinh gọn; detailed mode chỉ dùng khi cần."],
            ["Docs setup còn nhắc OAuth/Manus cũ", "Người deploy cấu hình sai auth.", "Viết DEMO_DEPLOYMENT.md ngắn, theo local auth hiện tại."],
            ["Cửa hàng lẫn vào demo main flow", "Demo bị kéo sang kịch bản dài và nhiều trạng thái.", "Để cửa hàng là demo phụ sau khi seed 2 ca đã sẵn sàng."],
        ],
        [1.65, 2.15, 2.7],
        header_fill=AMBER,
    )

    h1(doc, "7. Phụ lục: route và router đang dùng")
    matrix(
        doc,
        ["Nhóm", "Route / router chính"],
        [
            ["Manager routes", "/dashboard, /members, /organization, /finance, /daily-routine, /duties, /settings/users"],
            ["Resident routes", "/resident/today, /my-profile, /my-duties, /resident/finance, /resident/notifications, /resident/activities"],
            ["Store routes", "/store-ledger, /store-products, /store-purchase, /store-sales, /store-cashflow, /resident/store"],
            ["Server routers", "auth, dashboard, members, rooms, duties, organization, roles, users, dailyRoutine, residentPortal, activities, finance, storeLedger"],
        ],
        [1.65, 4.85],
    )

    h1(doc, "8. Kết luận")
    para(doc, "ResidenceCore đã có nền nghiệp vụ đủ để chuẩn bị deploy demo main flow quản lý. Việc quan trọng nhất trước deploy không phải mở thêm module mới, mà là đóng gói dữ liệu demo sạch, đồng bộ tài khoản demo, giữ menu simple gọn và smoke test lại từ DB mới tinh.")

    doc.core_properties.title = "ResidenceCore Business Process Document"
    doc.core_properties.subject = "Business process and demo readiness"
    doc.core_properties.author = "Codex"
    doc.core_properties.comments = "Generated from current ResidenceCore project audit."
    doc.save(OUTPUT)


if __name__ == "__main__":
    build_doc()
    print(OUTPUT)
