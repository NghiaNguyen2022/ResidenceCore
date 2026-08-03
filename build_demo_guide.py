from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = Path(r"D:\Source\Git\ResidenceCore")
ASSETS = ROOT / "demo-guide-assets"
OUT = ROOT / "Huong-dan-su-dung-va-du-lieu-demo-Luu-xa-Nu-Buon-Ma-Thuot.docx"
BLUE, DARK, LIGHT, PALE, GOLD, GRAY = "2E74B5","17365D","E8EEF5","F4F6F9","B7791F","666666"

doc = Document()
sec = doc.sections[0]
sec.page_width, sec.page_height = Inches(8.5), Inches(11)
sec.top_margin = sec.bottom_margin = Inches(.78)
sec.left_margin = sec.right_margin = Inches(.78)
sec.header_distance = sec.footer_distance = Inches(.38)

normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
normal.font.size = Pt(10.5)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.18
for nm,sz,col,bef,aft in [("Heading 1",16,BLUE,16,8),("Heading 2",13,BLUE,12,6),("Heading 3",11.5,DARK,9,4)]:
    s=doc.styles[nm]; s.font.name="Calibri"; s.font.size=Pt(sz); s.font.bold=True
    s.font.color.rgb=RGBColor.from_string(col); s.paragraph_format.space_before=Pt(bef); s.paragraph_format.space_after=Pt(aft)

def shade(cell, fill):
    pr=cell._tc.get_or_add_tcPr(); sh=OxmlElement("w:shd"); sh.set(qn("w:fill"),fill); pr.append(sh)

def margins(cell):
    pr=cell._tc.get_or_add_tcPr(); mar=OxmlElement("w:tcMar"); pr.append(mar)
    for k,v in [("top",90),("start",110),("bottom",90),("end",110)]:
        e=OxmlElement("w:"+k); e.set(qn("w:w"),str(v)); e.set(qn("w:type"),"dxa"); mar.append(e)

def add_table(headers, rows, widths):
    t=doc.add_table(rows=1,cols=len(headers)); t.style="Table Grid"; t.alignment=WD_TABLE_ALIGNMENT.CENTER; t.autofit=False
    for i,h in enumerate(headers):
        c=t.rows[0].cells[i]; c.text=h; shade(c,LIGHT); margins(c); c.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for r in c.paragraphs[0].runs: r.bold=True; r.font.size=Pt(9.3)
    hdr=OxmlElement("w:tblHeader"); hdr.set(qn("w:val"),"true"); t.rows[0]._tr.get_or_add_trPr().append(hdr)
    for row in rows:
        cells=t.add_row().cells
        for i,val in enumerate(row):
            cells[i].text=str(val); margins(cells[i]); cells[i].vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER
            cells[i].width=Inches(widths[i])
            for p in cells[i].paragraphs:
                p.paragraph_format.space_after=Pt(0)
                for r in p.runs: r.font.size=Pt(9)
    doc.add_paragraph()

def bullet(text):
    p=doc.add_paragraph(style="List Bullet"); p.add_run(text)
    p.paragraph_format.left_indent=Inches(.36); p.paragraph_format.first_line_indent=Inches(-.18); p.paragraph_format.space_after=Pt(3)

def step(n,title,body):
    p=doc.add_paragraph(); p.paragraph_format.space_before=Pt(5); p.paragraph_format.space_after=Pt(2)
    r=p.add_run(f"Bước {n}. {title}"); r.bold=True; r.font.color.rgb=RGBColor.from_string(DARK)
    q=doc.add_paragraph(body); q.paragraph_format.left_indent=Inches(.23)

def note(text,label="Lưu ý"):
    t=doc.add_table(rows=1,cols=1); t.alignment=WD_TABLE_ALIGNMENT.CENTER
    c=t.cell(0,0); shade(c,PALE); margins(c)
    p=c.paragraphs[0]; r=p.add_run(label+": "); r.bold=True; r.font.color.rgb=RGBColor.from_string(GOLD); p.add_run(text)
    doc.add_paragraph()

def fig(name,caption):
    f=ASSETS/name
    if not f.exists(): return
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.keep_with_next=True
    p.add_run().add_picture(str(f),width=Inches(6.55))
    c=doc.add_paragraph(caption); c.alignment=WD_ALIGN_PARAGRAPH.CENTER; c.paragraph_format.space_after=Pt(8)
    for r in c.runs: r.italic=True; r.font.size=Pt(9); r.font.color.rgb=RGBColor.from_string(GRAY)

hp=sec.header.paragraphs[0]; hp.text="RESIDENCECARE | SỔ TAY VẬN HÀNH LƯU XÁ NỮ"; hp.alignment=WD_ALIGN_PARAGRAPH.RIGHT
for r in hp.runs: r.font.size=Pt(8.5); r.font.bold=True; r.font.color.rgb=RGBColor.from_string(GRAY)
fp=sec.footer.paragraphs[0]; fp.text="Lưu xá nữ tại Buôn Ma Thuột - dữ liệu demo 08/2026"; fp.alignment=WD_ALIGN_PARAGRAPH.CENTER
for r in fp.runs: r.font.size=Pt(8); r.font.color.rgb=RGBColor.from_string(GRAY)

for _ in range(5): doc.add_paragraph()
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run("SỔ TAY VẬN HÀNH & DỮ LIỆU DEMO"); r.bold=True; r.font.size=Pt(12); r.font.color.rgb=RGBColor.from_string(GOLD)
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run("ResidenceCare"); r.bold=True; r.font.size=Pt(30); r.font.color.rgb=RGBColor.from_string(DARK)
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run("Quản lý lưu xá nữ của các sơ tại Buôn Ma Thuột"); r.font.size=Pt(15); r.font.color.rgb=RGBColor.from_string(BLUE)
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
p.add_run("Hướng dẫn theo menu, màn hình và tab - kèm ảnh dữ liệu thực").italic=True
for _ in range(7): doc.add_paragraph()
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
p.add_run("https://vireon.vn/app-portal/residence-management\n").bold=True
p.add_run("Cập nhật: 03/08/2026")
doc.add_page_break()

doc.add_heading("1. Mục tiêu và phạm vi",1)
doc.add_paragraph("Tài liệu hướng dẫn quản lý hằng ngày cho lưu xá nữ do các sơ phụ trách tại Buôn Ma Thuột. Bộ dữ liệu demo thể hiện quy trình tiếp nhận học viên, liên hệ gia đình, phân phòng, lịch học, tài chính, sinh hoạt, công tác và cửa hàng.")
note("CCCD, số điện thoại và email trong bộ demo là dữ liệu giả lập; không dùng để liên lạc thật.","An toàn dữ liệu")

doc.add_heading("2. Bộ dữ liệu demo đã thiết lập",1)
add_table(["Nhóm","Quy mô","Nội dung"],[
["Học viên","8","Nữ sinh 2005-2006, tên thánh, địa chỉ Đắk Lắk, có portal"],
["Phòng","5","Mân Côi, Nữ Vương, Fatima, La Vang, Têrêsa"],
["Phân phòng","8","Mỗi học viên có phòng; còn phòng dự phòng"],
["Liên hệ","1 mẫu","Liên hệ mẹ của Maria Nguyễn Thị Hồng Anh"],
["Học tập","1 + 2 lịch","Đại học Tây Nguyên, ngành Điều dưỡng"],
["Tài chính","2 nghiệp vụ","Tài trợ 5.000.000đ; chi tiền chợ 1.850.000đ"],
["Portal","8","Username riêng, mật khẩu tạm demo"]],[1.2,.8,4.4])

doc.add_heading("3. Đăng nhập và nguyên tắc",1)
step(1,"Mở hệ thống","Truy cập URL trên và chọn Đăng nhập.")
step(2,"Đăng nhập quản trị","Nhập tài khoản quản trị. Kiểm tra vai trò Quản lý lưu xá - Toàn hệ thống.")
step(3,"Đi theo menu","Các nhóm chính: Quản lý lưu trú, Sinh hoạt, Quản lý cửa hàng, Tài chính, Báo cáo & Thiết lập.")
note("Tạo dữ liệu theo thứ tự phòng -> học viên -> liên hệ/học tập -> lịch học -> công tác để giảm lỗi liên kết.")
doc.add_page_break()

doc.add_heading("4. Quản lý phòng ở",1)
doc.add_paragraph("Menu: Quản lý lưu trú -> Phòng ở.")
step(1,"Chọn Thêm phòng","Nhập mã, sức chứa, tổ (nếu đã tạo) và ghi chú.")
step(2,"Kiểm tra sức chứa","Hệ thống tự tính số chỗ còn trống theo học viên đang ở.")
step(3,"Lưu và đối chiếu","Kiểm tra mã, ghi chú, sức chứa và số chỗ còn lại.")
fig("01-them-phong.png","Hình 1 - Biểu mẫu thêm phòng MC-A101.")
fig("02-danh-sach-phong.png","Hình 2 - Danh sách 5 phòng demo.")
add_table(["Mã","Tên","Sức chứa","Ghi chú"],[
["MC-A101","Mân Côi","6","Tầng 1, khu A"],["NV-A102","Nữ Vương","6","Tầng 1, khu A"],
["FT-B201","Fatima","8","Tầng 2, khu B"],["LV-B202","La Vang","8","Tầng 2, khu B"],
["TR-C301","Têrêsa","4","Phòng dự phòng"]],[1.1,1.25,1,3.05])
doc.add_page_break()

doc.add_heading("5. Tiếp nhận học viên và cấp portal",1)
doc.add_paragraph("Menu: Quản lý lưu trú -> Học viên.")
step(1,"Chọn Thêm học viên","Nhập tên thánh, họ tên, ngày sinh, giới tính Nữ, CCCD, điện thoại, ngày vào và địa chỉ.")
step(2,"Cấp portal khi nhập học","Đánh dấu Tạo tài khoản đăng nhập; đặt username không dấu và mật khẩu tạm.")
step(3,"Lưu hồ sơ","Ghi rõ trường/ngành hoặc lưu ý vận hành trong phần Ghi chú.")
fig("03-them-hoc-vien-va-portal.png","Hình 3 - Hồ sơ học viên và tùy chọn tạo portal.")
fig("04-danh-sach-hoc-vien.png","Hình 4 - Danh sách học viên demo.")
add_table(["Học viên","Username","Phòng"],[
["Maria Nguyễn Thị Hồng Anh","hv.honganh","MC-A101"],["Anna Trần Ngọc Bảo Châu","hv.baochau","MC-A101"],
["Têrêsa Lê Thị Minh Thư","hv.minhthu","NV-A102"],["Cecilia Phạm Gia Hân","hv.giahan","NV-A102"],
["Lucia Võ Khánh Linh","hv.khanhlinh","FT-B201"],["Rosa Huỳnh Ngọc Mai","hv.ngocmai","FT-B201"],
["Agnes Nguyễn Thảo Vy","hv.thaovy","LV-B202"],["Monica Trần Thanh Trúc","hv.thanhtruc","LV-B202"]],[3.2,1.8,1.4])
note("Mật khẩu tạm demo: LuuXa@2026. Khi vận hành thật phải đổi ở lần đăng nhập đầu tiên.","Portal")
doc.add_page_break()

doc.add_heading("6. Hồ sơ học viên - từng tab",1)
doc.add_paragraph("Tại danh sách Học viên chọn Xem. Hồ sơ gồm Tổng quan, Liên hệ, Phòng ở, Học tập, Tổ chức và Tài khoản.")
doc.add_heading("6.1 Tổng quan",2)
bullet("Kiểm tra thông tin cá nhân, trạng thái lưu trú, phòng hiện tại, ngày vào và liên hệ chính.")
bullet("Dùng Chuyển / Trả phòng khi thay đổi chỗ ở; luôn ghi lý do và ngày hiệu lực.")
doc.add_heading("6.2 Liên hệ",2)
step(1,"Chọn tab Liên hệ","Bấm Thêm liên hệ.")
step(2,"Nhập quan hệ","Chọn Mẹ, Cha hoặc Người giám hộ; nhập họ tên, số điện thoại, nghề nghiệp, địa chỉ.")
step(3,"Kiểm tra liên hệ chính","Một học viên tối đa một Cha và một Mẹ.")
fig("05-them-lien-he-gia-dinh.png","Hình 5 - Biểu mẫu thêm liên hệ.")
fig("06-tab-lien-he.png","Hình 6 - Tab Liên hệ sau khi lưu.")
doc.add_heading("6.3 Phòng ở",2)
bullet("Chọn Gán phòng, chọn phòng còn chỗ, nhập ngày hiệu lực và lý do.")
bullet("Khi chuyển phòng, dùng thao tác Chuyển phòng để giữ lịch sử.")
doc.add_heading("6.4 Học tập",2)
step(1,"Thêm thông tin học tập","Nhập trường, bậc học, lớp/ngành, khóa và ghi chú.")
step(2,"Thêm lịch học","Chọn các ngày có cùng giờ; nhập giờ, môn học và địa điểm.")
step(3,"Xem Ngày/Tuần/Tháng","Lịch là dữ liệu cảnh báo khi phân công công tác.")
fig("07-them-thong-tin-hoc-tap.png","Hình 7 - Thông tin học tập tại Đại học Tây Nguyên.")
fig("08-them-lich-hoc.png","Hình 8 - Thêm lịch học nhiều ngày.")
fig("09-tab-hoc-tap-va-lich-hoc.png","Hình 9 - Tab Học tập và lịch học.")
note("Hồng Anh: Thứ 2-4-6 07:30-11:30; Thứ 3-5 13:30-17:00. Công tác phải nằm ngoài các khung này.")
doc.add_heading("6.5 Tổ chức",2)
bullet("Hiển thị Tổ/Ban và chức vụ đang hiệu lực. Chỉ bổ nhiệm sau khi tạo nhiệm kỳ.")
doc.add_heading("6.6 Tài khoản",2)
bullet("Kiểm tra username, trạng thái và yêu cầu đổi mật khẩu. Khi thôi lưu trú, khóa thay vì xóa tài khoản.")
doc.add_page_break()

doc.add_heading("7. Tổ chức lưu xá",1)
doc.add_paragraph("Menu: Quản lý lưu trú -> Tổ chức lưu xá. Các tab: Cơ cấu hiện tại, Lịch sử bổ nhiệm, Tổ/Ban, Nhiệm kỳ.")
step(1,"Tạo nhiệm kỳ","Mã NK-2026-2027; tên Năm học 2026-2027; trạng thái Đang áp dụng.")
step(2,"Tạo Tổ/Ban","Tổ Phụng vụ, Tổ Học tập, Tổ Hậu cần - Cửa hàng và Ban Đời sống.")
step(3,"Bổ nhiệm","Chọn học viên, vai trò, đơn vị, nhiệm kỳ và ngày hiệu lực.")
fig("10-them-nhiem-ky.png","Hình 10 - Biểu mẫu nhiệm kỳ 2026-2027.")
note("Trường ngày của biểu mẫu Nhiệm kỳ đôi lúc không nhận giá trị trên bản deploy. Nhập trực tiếp và kiểm tra lại trước khi Lưu.","Điểm cần kiểm tra")

doc.add_heading("8. Tài chính lưu xá",1)
doc.add_paragraph("Menu: Tài chính -> Tài chính lưu xá. Gồm Thu học viên, Thu chi khác, Theo dõi và Sổ dòng tiền.")
doc.add_heading("8.1 Kỳ thu học viên",2)
step(1,"Tạo kỳ thu","Đặt tên, năm, khoảng tháng và các thành phần phí lưu trú, ăn uống, khoản khác.")
step(2,"Phân bổ học viên","Chọn kỳ và tạo khoản phải thu cho học viên đang lưu trú.")
step(3,"Thu theo học viên","Ghi số tiền, ngày, phương thức và nội dung.")
fig("11-tao-ky-thu.png","Hình 11 - Biểu mẫu kỳ thu.")
doc.add_heading("8.2 Thu chi khác",2)
bullet("Tài trợ: ghi đơn vị tài trợ và mục đích.")
bullet("Khoản chi: chọn Tiền đi chợ, Văn phòng phẩm, Hoa đèn, Vật dụng, Sửa chữa hoặc Hỗ trợ.")
bullet("Kinh doanh: dành cho nghiệp vụ cửa hàng, không trộn với phí lưu xá.")
fig("12-ghi-nhan-tai-tro.png","Hình 12 - Tài trợ 5.000.000đ.")
fig("13-ghi-nhan-khoan-chi.png","Hình 13 - Chi tiền chợ 1.850.000đ.")
fig("14-tong-quan-tai-chinh.png","Hình 14 - Tổng quan tài chính.")
note("Mọi khoản thu/chi phải có ngày, đối tượng và ghi chú đủ để đối chiếu.")
doc.add_page_break()

doc.add_heading("9. Công tác ngày, tuần, tháng",1)
doc.add_paragraph("Menu: Sinh hoạt -> Công tác / Trực nhật.")
step(1,"Tạo mẫu công tác","Nhập mã, tên, mô tả, loại lặp, giờ và số người.")
step(2,"Bật kiểm tra xung đột","Giữ chọn Kiểm tra xung đột với lịch học.")
step(3,"Tạo phân công","Chọn mẫu, ngày, học viên/phòng/tổ và người phụ trách.")
add_table(["Mã","Loại","Giờ","Nội dung"],[
["TRUC_CH_S","Ngày","05:45-07:00","Trực cửa hàng sáng"],["VS_BEP","Ngày","17:30-18:15","Vệ sinh bếp"],
["TONG_VS","Tuần","07:00-09:00 T7","Tổng vệ sinh"],["KT_PCCC","Tháng","Tuần đầu","Kiểm tra PCCC"]],[1.2,.8,1.5,2.9])
note("Không phân công Hồng Anh 07:30-11:30 Thứ 2-4-6 hoặc 13:30-17:00 Thứ 3-5.","Tránh trùng lịch")

doc.add_heading("10. Hoạt động và điểm danh",1)
doc.add_paragraph("Menu: Sinh hoạt -> Hoạt động / Sự kiện; Điểm danh; Lịch điểm danh.")
step(1,"Tạo hoạt động","Nhập tên, thời gian, địa điểm, người phụ trách và nội dung.")
step(2,"Chọn người tham gia","Chọn theo phòng hoặc Tổ/Ban; kiểm tra lịch học.")
step(3,"Điểm danh","Ghi Có mặt, Vắng phép, Vắng không phép và ghi chú.")
add_table(["Hoạt động","Thời gian gợi ý","Địa điểm"],[
["Tĩnh tâm đầu năm","Chủ nhật 09/08/2026","Nhà nguyện"],["Kỹ năng sơ cấp cứu","19:00-20:30 Thứ 7","Phòng sinh hoạt"],
["Thăm mái ấm","Chủ nhật cuối tháng","Buôn Ma Thuột"]],[2.5,2.1,1.7])

doc.add_heading("11. Quản lý cửa hàng",1)
doc.add_paragraph("Menu gồm Dữ liệu sản phẩm, Mua hàng/Nhập kho, Bán hàng và Tổng hợp thu chi.")
step(1,"Tạo sản phẩm","Nhập mã, tên, đơn vị, giá bán và tồn tối thiểu.")
step(2,"Nhập kho","Ghi nhà cung cấp, ngày nhập, số lượng, đơn giá và người nhận.")
step(3,"Bán hàng","Chọn học viên/khách, sản phẩm, số lượng; kiểm tra tổng tiền.")
step(4,"Đối chiếu","Cuối ca đối chiếu tiền, tồn kho và sổ dòng tiền.")
add_table(["Mã","Sản phẩm","ĐVT","Giá","Tồn"],[
["ST-VO200","Vở 200 trang","Quyển","15.000đ","30"],["ST-BUTX","Bút bi xanh","Cây","5.000đ","50"],
["TP-MIG","Mì gói","Gói","6.000đ","60"],["TP-SUA","Sữa hộp","Hộp","8.000đ","48"],
["GD-XAPHONG","Xà phòng","Bánh","12.000đ","24"]],[1.2,2,0.7,1.1,0.8])

doc.add_heading("12. Portal học viên",1)
doc.add_paragraph("Portal tập trung vào Hôm nay, Hồ sơ, Công tác, Cửa hàng, Tài chính, Thông báo và Hoạt động.")
for x in ["Hôm nay: xem công tác, lịch sinh hoạt và thông báo.","Công tác: xem việc được giao và cập nhật hoàn thành.","Cửa hàng: xem giao dịch cá nhân theo quyền.","Tài chính: xem phải thu, đã nộp và còn thiếu.","Hoạt động: xem lịch và trạng thái tham gia."]: bullet(x)
note("Mỗi học viên phải dùng tài khoản riêng để bảo đảm truy vết.")

doc.add_heading("13. Quy trình vận hành",1)
add_table(["Chu kỳ","Việc cần làm"],[
["Hằng ngày","Điểm danh; công tác; trực cửa hàng; ghi thu chi"],["Hằng tuần","Rà lịch học; phân công tuần; tổng vệ sinh; đối chiếu kho"],
["Hằng tháng","Tạo kỳ thu; chốt sổ; kiểm kê; kiểm tra phòng"],["Đầu năm","Nhập học viên; cấp portal; phân phòng; nhập lịch học"],
["Cuối năm","Chốt công nợ; trả phòng; khóa tài khoản; lưu báo cáo"]],[1.25,5.05])

doc.add_heading("14. Checklist dữ liệu demo",1)
for x in ["Dashboard có 8 học viên và 5 phòng.","Tất cả học viên có portal và phòng.","Hồng Anh có liên hệ, học tập và 2 lịch học.","Tài chính có thu 5.000.000đ và chi 1.850.000đ.","Mẫu công tác bật kiểm tra trùng lịch.","Cửa hàng có sản phẩm, nhập, bán và đối chiếu."]: bullet(x)

doc.add_heading("15. Điểm cần hoàn thiện trên bản deploy",1)
for x in ["Kiểm tra nhập ngày ở biểu mẫu Nhiệm kỳ.","Kiểm tra giữ trạng thái Kỳ thu khi đổi tháng/năm.","Không hiển thị Tổ A/B/C khi chưa có bản ghi tổ thật.","Thay thông báo câu SQL bằng lỗi thân thiện."]: bullet(x)

doc.core_properties.title="Sổ tay ResidenceCare - Lưu xá nữ Buôn Ma Thuột"
doc.core_properties.subject="Hướng dẫn sử dụng và dữ liệu demo"
doc.core_properties.author="Codex"
doc.save(OUT)
print(OUT)
