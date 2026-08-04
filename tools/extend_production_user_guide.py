from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "docs" / "HUONG_DAN_SU_DUNG_RESIDENCECORE_LUU_XA_NU_BMT.docx"
OUTPUT = ROOT / "docs" / "HUONG_DAN_SU_DUNG_RESIDENCECORE_LUU_XA_NU_BMT_CHI_TIET_HOST.docx"
LIVE_ASSETS = ROOT / "demo-guide-assets"

BLUE = RGBColor(46, 116, 181)
DARK = RGBColor(31, 41, 55)
MUTED = RGBColor(91, 105, 120)


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=100, start=120, bottom=100, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.autofit = False
    for i, text in enumerate(headers):
        cell = table.rows[0].cells[i]
        shade(cell, "2E74B5")
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(str(text))
        r.bold = True
        r.font.color.rgb = RGBColor(255, 255, 255)
        r.font.size = Pt(9)
        set_cell_margins(cell)
    for row_idx, row in enumerate(rows):
        cells = table.add_row().cells
        for i, text in enumerate(row):
            cell = cells[i]
            if row_idx % 2:
                shade(cell, "F4F6F9")
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            cell.text = str(text)
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)
                    r.font.color.rgb = DARK
            set_cell_margins(cell)
    if widths:
        for row in table.rows:
            for i, width in enumerate(widths):
                row.cells[i].width = Inches(width)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item)


def add_steps(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.add_run(item)


def add_note(doc, title, body, fill="FFF8E8"):
    table = doc.add_table(rows=1, cols=1)
    table.autofit = False
    table.columns[0].width = Inches(6.25)
    cell = table.cell(0, 0)
    shade(cell, fill)
    set_cell_margins(cell, 150, 180, 150, 180)
    p = cell.paragraphs[0]
    r = p.add_run(f"{title}: ")
    r.bold = True
    r.font.color.rgb = BLUE
    p.add_run(body)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_picture(doc, filename, caption, width=6.25):
    path = LIVE_ASSETS / filename
    if not path.exists():
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.keep_with_next = True
    p.add_run().add_picture(str(path), width=Inches(width))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(7)
    r = cap.add_run(caption)
    r.italic = True
    r.font.size = Pt(9)
    r.font.color.rgb = MUTED


doc = Document(SOURCE)
doc.add_section(WD_SECTION.NEW_PAGE)

doc.add_heading("Phụ lục A. Vận hành bản đã triển khai trên vireon.vn", level=1)
doc.add_paragraph(
    "Phần này bổ sung cho 34 quy trình và 79 hình minh họa ở phần chính. Nội dung được viết theo "
    "dữ liệu demo đã tạo thực tế trên host Vietnix tại Buôn Ma Thuột, giúp người quản lý biết chính xác "
    "vào menu nào, nhập dữ liệu gì, kiểm tra ở đâu và xử lý thế nào khi kết quả chưa đúng."
)
add_note(
    doc,
    "Địa chỉ chính thức",
    "https://vireon.vn/app-portal/residence-management. Luôn giữ nguyên tiền tố /app-portal/residence-management; "
    "không dùng các đường dẫn gốc như /login hoặc /dashboard vì LiteSpeed sẽ trả về 404.",
    "E8EEF5",
)

doc.add_heading("A.1. Thông tin đăng nhập và bảo mật", level=2)
add_table(
    doc,
    ["Đối tượng", "Tên đăng nhập", "Mật khẩu ban đầu", "Việc phải làm"],
    [
        ["Quản trị", "admin", "Abcd@1234", "Đổi mật khẩu sau khi nghiệm thu; không chia sẻ qua nhóm công khai."],
        ["Học viên demo", "Theo cột Portal username", "LuuXa@2026", "Yêu cầu đổi mật khẩu ở lần đăng nhập đầu tiên."],
    ],
    [1.0, 1.35, 1.25, 2.65],
)
add_steps(
    doc,
    [
        "Mở đúng URL chính thức ở trên; hệ thống tự chuyển đến trang đăng nhập bên trong thư mục ứng dụng.",
        "Nhập admin và mật khẩu quản trị, chọn Đăng nhập.",
        "Kiểm tra thanh địa chỉ sau đăng nhập vẫn bắt đầu bằng /app-portal/residence-management/.",
        "Nếu bị đưa về https://vireon.vn/login hoặc https://vireon.vn/dashboard, tải lại URL chính thức và báo kỹ thuật kiểm tra APP_BASE_PATH, router và cookie path.",
    ],
)

doc.add_heading("A.2. Bộ dữ liệu demo đã tạo thực tế", level=2)
doc.add_paragraph("Dữ liệu mẫu sử dụng tên nữ sinh, trường học và hoạt động phù hợp bối cảnh lưu xá nữ do các sơ quản lý tại Buôn Ma Thuột.")

doc.add_heading("A.2.1. Phòng ở", level=3)
add_table(
    doc,
    ["Mã phòng", "Tên phòng", "Khu", "Sức chứa", "Mục đích demo"],
    [
        ["MC-A101", "Mân Côi A101", "Khu A", "4", "Phòng nữ sinh năm nhất"],
        ["NV-A102", "Nữ Vương A102", "Khu A", "4", "Phòng nữ sinh năm hai"],
        ["FT-B201", "Fatima B201", "Khu B", "4", "Phòng nữ sinh năm ba"],
        ["LV-B202", "Lộ Đức B202", "Khu B", "4", "Phòng nữ sinh năm cuối"],
        ["TR-C301", "Têrêsa C301", "Khu C", "4", "Phòng dự phòng/tiếp nhận mới"],
    ],
    [1.0, 1.6, 0.75, 0.8, 2.05],
)
add_picture(doc, "01-them-phong.png", "Hình A1. Biểu mẫu thêm phòng trên bản triển khai; ảnh chụp theo khung chức năng, không dùng ảnh toàn trang.")
add_steps(
    doc,
    [
        "Vào Quản lý lưu trú → Phòng ở.",
        "Chọn Thêm phòng; nhập lần lượt Mã phòng, Tên phòng, Khu/Tầng, Sức chứa và ghi chú.",
        "Chọn Lưu; quay lại danh sách và tìm theo mã phòng vừa tạo.",
        "Kiểm tra số người đang ở không vượt sức chứa; chỉ gán học viên sau khi phòng xuất hiện trong danh sách.",
    ],
)

doc.add_heading("A.2.2. Học viên và tài khoản portal", level=3)
students = [
    ["LX26-001", "Nguyễn Thị Hồng Anh", "honganh.nguyen", "MC-A101"],
    ["LX26-002", "Trần Thị Minh Châu", "minhchau.tran", "MC-A101"],
    ["LX26-003", "Lê Thị Thanh Hà", "thanhha.le", "NV-A102"],
    ["LX26-004", "Phạm Thị Bảo Ngọc", "baongoc.pham", "NV-A102"],
    ["LX26-005", "Vũ Thị Khánh Linh", "khanhlinh.vu", "FT-B201"],
    ["LX26-006", "Hoàng Thị Thu Trang", "thutrang.hoang", "FT-B201"],
    ["LX26-007", "Đỗ Thị Ngọc Mai", "ngocmai.do", "LV-B202"],
    ["LX26-008", "Bùi Thị Kim Oanh", "kimoanh.bui", "LV-B202"],
]
add_table(doc, ["Mã", "Họ và tên", "Portal username", "Phòng"], students, [0.85, 2.1, 1.7, 1.0])
add_picture(doc, "03-them-hoc-vien-va-portal.png", "Hình A2. Màn hình tạo học viên đồng thời cấp tài khoản portal.")
add_steps(
    doc,
    [
        "Vào Học viên → Danh sách học viên → Thêm học viên.",
        "Nhập Mã học viên, Họ tên, ngày sinh, giới tính Nữ, số điện thoại, địa chỉ gia đình và ngày nhập lưu xá.",
        "Bật tùy chọn tạo tài khoản portal; nhập username không dấu, viết thường và không có khoảng trắng.",
        "Đặt mật khẩu tạm LuuXa@2026 và bật yêu cầu đổi mật khẩu nếu biểu mẫu có tùy chọn này.",
        "Chọn Lưu; tìm lại theo mã học viên và mở từng tab để bổ sung Liên hệ, Phòng ở, Học tập, Tổ chức và Tài khoản.",
    ],
)
add_note(doc, "Quy tắc dữ liệu", "Mỗi học viên có một mã và một username duy nhất. Không dùng chung tài khoản portal. Trước khi cấp tài khoản, kiểm tra đúng số điện thoại người thân và tình trạng đang lưu trú.")

doc.add_heading("A.3. Hướng dẫn chi tiết từng tab hồ sơ học viên", level=2)
tab_rows = [
    ["Tổng quan", "Thông tin định danh, ngày nhập học/lưu xá, trạng thái", "Họ tên, mã, giới tính, điện thoại, địa chỉ", "Thẻ hồ sơ hiển thị đúng trạng thái Đang lưu trú"],
    ["Liên hệ", "Người thân và liên hệ khẩn cấp", "Họ tên, quan hệ, điện thoại, địa chỉ, liên hệ chính", "Có ít nhất một liên hệ chính"],
    ["Phòng ở", "Lịch sử phân phòng", "Phòng, ngày bắt đầu, lý do, ghi chú", "Không vượt sức chứa và không có hai phòng hiệu lực"],
    ["Học tập", "Trường, ngành, lớp, lịch học", "Trường, ngành, khóa, môn, thứ, giờ bắt đầu/kết thúc", "Lịch học được dùng để cảnh báo trùng công tác"],
    ["Tổ chức", "Tổ, Ban và chức vụ", "Đơn vị, vai trò, nhiệm kỳ, ngày hiệu lực", "Học viên xuất hiện đúng phạm vi phụ trách"],
    ["Tài khoản", "Quyền truy cập portal", "Username, trạng thái, đổi mật khẩu", "Đăng nhập được và chỉ thấy dữ liệu của chính mình"],
]
add_table(doc, ["Tab", "Mục đích", "Thông tin phải nhập", "Kết quả phải kiểm tra"], tab_rows, [0.85, 1.55, 2.15, 2.15])

doc.add_heading("A.3.1. Tab Liên hệ - mẫu Nguyễn Thị Hồng Anh", level=3)
add_steps(doc, [
    "Mở Học viên → Nguyễn Thị Hồng Anh → tab Liên hệ.",
    "Chọn Thêm liên hệ; nhập người mẹ/người giám hộ, quan hệ Mẹ, số điện thoại và địa chỉ tại Đắk Lắk.",
    "Đánh dấu Liên hệ chính và Liên hệ khẩn cấp khi phù hợp; chọn Lưu.",
    "Tải lại tab và gọi thử số liên hệ trong kịch bản nghiệm thu để xác nhận dữ liệu không bị sai ký tự hoặc mất số 0 đầu.",
])

doc.add_heading("A.3.2. Tab Học tập và lịch học", level=3)
add_steps(doc, [
    "Trong hồ sơ Hồng Anh, mở tab Học tập → Thêm thông tin học tập.",
    "Nhập tên trường tại Buôn Ma Thuột, ngành học, lớp/khóa, năm bắt đầu và trạng thái Đang học.",
    "Chọn Thêm lịch học; tạo tối thiểu hai lịch trong tuần, ghi rõ môn, thứ, giờ bắt đầu, giờ kết thúc và phòng học.",
    "Lưu từng lịch; kiểm tra thứ tự thời gian, không để giờ kết thúc nhỏ hơn giờ bắt đầu.",
    "Khi tạo công tác, chọn học viên này và đặt thời gian giao nhau với lịch học để xác nhận hệ thống cảnh báo; sau đó đổi giờ công tác sang khung không trùng trước khi lưu chính thức.",
])

doc.add_heading("A.4. Tổ chức lưu xá, nhiệm kỳ và phân bổ học viên", level=2)
add_picture(doc, "10-them-nhiem-ky.png", "Hình A3. Biểu mẫu tạo nhiệm kỳ dùng cho Tổ, Ban và chức vụ học viên.")
add_steps(doc, [
    "Vào Tổ chức lưu xá → Nhiệm kỳ → Thêm nhiệm kỳ.",
    "Nhập tên Nhiệm kỳ 2026-2027, ngày bắt đầu, ngày kết thúc và trạng thái Hoạt động.",
    "Tạo các Tổ theo phòng/khu và các Ban phù hợp sinh hoạt: Ban Phụng vụ, Ban Học tập, Ban Đời sống, Ban Văn nghệ.",
    "Mở từng đơn vị → Thêm thành viên; chọn học viên, vai trò, ngày hiệu lực và nhiệm kỳ.",
    "Với chức vụ Tổ trưởng/Trưởng ban, kiểm tra lịch sử bổ nhiệm để tránh hai người cùng giữ một vị trí trong cùng thời gian nếu quy định không cho phép.",
    "Quay lại hồ sơ học viên → tab Tổ chức để đối chiếu; dữ liệu phải xuất hiện ở cả hai phía.",
])

doc.add_heading("A.5. Công tác ngày, tuần, tháng và tránh trùng lịch học", level=2)
add_table(doc, ["Chu kỳ", "Ví dụ phù hợp lưu xá nữ", "Thời điểm", "Phạm vi"], [
    ["Ngày", "Trực cổng, vệ sinh nhà nguyện, rửa chén", "Theo ca sáng/chiều/tối", "Cá nhân hoặc phòng"],
    ["Tuần", "Tổng vệ sinh, phụng vụ Chúa nhật, kiểm kho", "Một ngày cố định mỗi tuần", "Tổ hoặc Ban"],
    ["Tháng", "Họp cộng đoàn, kiểm tra phòng, đối chiếu thu chi", "Cuối/tháng đầu", "Toàn lưu xá"],
], [0.8, 2.65, 1.5, 1.3])
add_steps(doc, [
    "Vào Công tác → Tạo công tác; chọn loại Ngày, Tuần hoặc Tháng.",
    "Nhập tiêu đề có động từ rõ ràng, mô tả kết quả cần đạt, ngày giờ bắt đầu/kết thúc, địa điểm và người phụ trách.",
    "Chọn phạm vi Cá nhân, Phòng, Tổ, Ban hoặc Toàn lưu xá; sau đó chọn đúng người/đơn vị nhận việc.",
    "Quan sát cảnh báo lịch học. Nếu có xung đột, mở lịch học của học viên để kiểm tra rồi đổi ca hoặc giao người khác; không bỏ qua cảnh báo chỉ để hoàn tất biểu mẫu.",
    "Lưu công tác; mở Điều hành để kiểm tra trạng thái Chưa làm/Đang làm/Hoàn thành và phạm vi hiển thị.",
    "Sau khi hoàn thành, người thực hiện phản hồi trên portal; quản lý kiểm tra nội dung, thời gian và xác nhận nếu quy trình yêu cầu.",
])
add_note(doc, "Nguyên tắc tránh trùng", "Lịch học là ưu tiên cố định. Công tác trực cổng, cửa hàng, vệ sinh và hoạt động chung phải được xếp ngoài giờ học; trường hợp khẩn cấp phải có người thay thế và ghi rõ lý do.")

doc.add_heading("A.6. Thu chi, tiền lưu xá và dữ liệu tài chính đã tạo", level=2)
add_table(doc, ["Nhóm", "Dữ liệu demo", "Số tiền", "Nơi kiểm tra"], [
    ["Khoản thu khác", "Tài trợ cho lưu xá nữ", "5.000.000 đ", "Tài chính → Thu chi khác/Sổ dòng tiền"],
    ["Khoản chi", "Đi chợ và nhu yếu phẩm", "1.850.000 đ", "Tài chính → Thu chi khác/Sổ dòng tiền"],
    ["Tiền lưu xá", "Kỳ thu theo tháng", "Theo chính sách", "Tài chính → Kỳ thu → Khoản phải thu"],
], [1.15, 2.35, 1.25, 1.75])
add_steps(doc, [
    "Vào Tài chính → Kỳ thu → Tạo kỳ thu; nhập tên kỳ, thời gian áp dụng, hạn đóng và các khoản tiền lưu xá.",
    "Chọn danh sách học viên áp dụng; sinh khoản phải thu và kiểm tra tổng số học viên × mức thu bằng tổng kỳ thu.",
    "Khi nhận tiền, mở khoản phải thu của từng học viên → Ghi nhận thanh toán; nhập ngày, số tiền, phương thức, người thu và mã chứng từ.",
    "Đối với tài trợ 5.000.000 đ, vào Thu chi khác → Ghi nhận khoản thu; chọn đúng nguồn tài trợ và ghi chú mục đích sử dụng.",
    "Đối với chi chợ 1.850.000 đ, vào Thu chi khác → Ghi nhận khoản chi; chọn nhóm nhu yếu phẩm, người chi, ngày chứng từ và đính kèm hóa đơn nếu có.",
    "Cuối ngày/tháng, đối chiếu Sổ dòng tiền: số dư đầu + tổng thu - tổng chi = số dư cuối; không sửa trực tiếp giao dịch đã khóa mà phải lập phiếu điều chỉnh.",
])

doc.add_heading("A.7. Hoạt động chung và điểm danh", level=2)
add_steps(doc, [
    "Vào Hoạt động → Thêm hoạt động; nhập tên, loại hoạt động, thời gian, địa điểm, người phụ trách và mô tả.",
    "Chọn phạm vi Toàn lưu xá hoặc nhóm tham gia; công khai lên portal nếu học viên cần xem và đăng ký.",
    "Trước giờ diễn ra, mở danh sách người tham gia; sau hoạt động cập nhật Có mặt, Đi trễ, Vắng có phép hoặc Vắng không phép.",
    "Kiểm tra tổng bốn trạng thái phải bằng tổng người tham gia; trường hợp vắng phải có ghi chú hoặc lý do.",
    "Với hoạt động lặp lại, tạo mẫu hoặc sao chép nhưng phải rà lại ngày giờ để tránh dùng nhầm lịch cũ.",
])

doc.add_heading("A.8. Cửa hàng lưu xá", level=2)
add_steps(doc, [
    "Vào Cửa hàng → Sản phẩm; tạo mã, tên hàng, đơn vị tính, giá bán, mức tồn tối thiểu và trạng thái.",
    "Vào Nhập hàng; chọn nhà cung cấp, ngày chứng từ, từng sản phẩm, số lượng và đơn giá. Lưu rồi kiểm tra tồn kho và khoản chi phát sinh.",
    "Vào Bán hàng; chọn sản phẩm, số lượng, người bán/ca trực và phương thức thu. Kiểm tra doanh thu và tồn kho giảm đúng.",
    "Cuối ca, kiểm đếm tiền mặt và hàng tồn; lập phiếu chốt ngày, review các chênh lệch trước khi xác nhận.",
    "Chỉ đẩy dữ liệu sang sổ tài chính chung sau khi phiếu bán, phiếu nhập, tiền mặt và tồn kho đã khớp.",
    "Phân bổ học viên trực cửa hàng qua Công tác; kiểm tra lịch học trước khi xếp ca và chỉ cấp quyền cửa hàng trong ca được phân công.",
])

doc.add_heading("A.9. Portal học viên - quy trình bàn giao tài khoản", level=2)
add_steps(doc, [
    "Quản lý mở hồ sơ học viên → tab Tài khoản, xác nhận username, trạng thái Hoạt động và yêu cầu đổi mật khẩu.",
    "Bàn giao riêng username và mật khẩu tạm; yêu cầu học viên đăng nhập đúng URL có tiền tố ứng dụng.",
    "Học viên đổi mật khẩu, sau đó kiểm tra các mục Hôm nay, Hồ sơ, Lịch học, Công tác, Hoạt động, Tài chính và Cửa hàng.",
    "Xác nhận học viên chỉ xem được hồ sơ/công nợ/công tác thuộc phạm vi của mình, không thấy dữ liệu riêng của học viên khác.",
    "Học viên phản hồi công tác và theo dõi thông báo; quản lý đối chiếu ở màn hình Điều hành.",
    "Khi học viên rời lưu xá, khóa tài khoản portal sau khi hoàn tất đối chiếu công nợ và bàn giao tài sản.",
])

doc.add_heading("A.10. Checklist nghiệm thu theo menu", level=2)
checks = [
    ["Đăng nhập/Router", "URL luôn giữ tiền tố ứng dụng; refresh trang con không 404", "□"],
    ["Học viên", "8 hồ sơ demo, tiếng Việt đúng dấu, tìm kiếm được", "□"],
    ["Liên hệ", "Hồng Anh có liên hệ gia đình chính", "□"],
    ["Phòng", "5 phòng; 8 học viên được gán; không vượt sức chứa", "□"],
    ["Học tập", "Có trường/ngành và tối thiểu 2 lịch học mẫu", "□"],
    ["Tổ chức", "Nhiệm kỳ, Tổ/Ban, thành viên và chức vụ đối chiếu hai chiều", "□"],
    ["Công tác", "Tạo được ngày/tuần/tháng và cảnh báo trùng lịch học", "□"],
    ["Hoạt động", "Tạo, công khai, điểm danh và tổng trạng thái khớp", "□"],
    ["Tài chính", "Thu 5.000.000 đ, chi 1.850.000 đ và kỳ thu học viên hiển thị đúng", "□"],
    ["Cửa hàng", "Nhập, bán, tồn, chốt ngày và sổ chung đối chiếu", "□"],
    ["Portal", "Đăng nhập, đổi mật khẩu, dữ liệu đúng phạm vi", "□"],
]
add_table(doc, ["Phân hệ", "Tiêu chí đạt", "Xác nhận"], checks, [1.25, 4.55, 0.65])

doc.add_heading("A.11. Lỗi thường gặp trên host Vietnix", level=2)
add_table(doc, ["Hiện tượng", "Nguyên nhân thường gặp", "Cách xử lý"], [
    ["404 tại /login hoặc /dashboard", "Router tạo đường dẫn tuyệt đối, thiếu base path", "Mở lại URL chính thức; kỹ thuật kiểm tra APP_BASE_PATH và route frontend/backend."],
    ["Đăng nhập xong quay lại login", "COOKIE_PATH hoặc cookie phiên không đúng thư mục", "Đặt COOKIE_PATH=/app-portal/residence-management; restart Node.js app."],
    ["Failed query với createdAt/isActive", "Tên cột MySQL khác chữ hoa-thường hoặc schema chưa đồng bộ", "Đối chiếu migration/schema với bảng thật; không tự đổi cột khi chưa backup."],
    ["Tiếng Việt thành Qu?n l?", "File/DB/connection không dùng UTF-8 đầy đủ", "Dùng utf8mb4 cho database, bảng, cột và connection; import file UTF-8 rồi kiểm tra dữ liệu đã hỏng."],
    ["Ảnh trong Word rất nhỏ", "Chèn ảnh full-page cao và ép vừa chiều ngang", "Chụp theo viewport/khung chức năng 16:9 hoặc chia thành các ảnh riêng; không kéo méo ảnh."],
    ["Không thấy menu", "Vai trò hoặc chế độ menu đơn giản", "Kiểm tra quyền tài khoản, đổi sang chế độ chi tiết hoặc mở đường dẫn phân hệ đã được quản trị cấp."],
], [1.55, 2.25, 2.65])

doc.add_heading("A.12. Quy trình vận hành đề xuất", level=2)
add_table(doc, ["Tần suất", "Người thực hiện", "Công việc"], [
    ["Hằng ngày", "Sơ quản lý/Trực ban", "Điểm danh; theo dõi công tác; thu chi; kiểm tra ca cửa hàng; xử lý thông báo."],
    ["Hằng tuần", "Sơ phụ trách/Tổ trưởng", "Xếp công tác tuần kế tiếp; đối chiếu lịch học; kiểm tra phòng; tổng vệ sinh; rà công nợ."],
    ["Hằng tháng", "Quản lý/Kế toán", "Tạo kỳ thu; chốt thu chi; kiểm kê cửa hàng; báo cáo hoạt động; rà tài khoản và quyền."],
    ["Đầu năm học", "Ban quản lý", "Tiếp nhận học viên; cấp portal; phân phòng/Tổ/Ban; nhập lịch học; ban hành nội quy."],
    ["Khi rời lưu xá", "Ban quản lý", "Chốt công nợ; thu hồi tài sản; kết thúc phân phòng/chức vụ; khóa portal; lưu hồ sơ."],
], [1.1, 1.65, 3.7])

doc.core_properties.title = "Hướng dẫn sử dụng ResidenceCore - Lưu xá nữ Buôn Ma Thuột - Bản chi tiết trên host"
doc.core_properties.subject = "Sổ tay thao tác, dữ liệu demo thực tế và checklist vận hành bản Vietnix"
doc.core_properties.comments = "Bản mở rộng ngày 03/08/2026; kế thừa 79 hình minh họa rõ từ tài liệu tham khảo và bổ sung dữ liệu host."
OUTPUT.parent.mkdir(parents=True, exist_ok=True)
doc.save(OUTPUT)
print(OUTPUT)
