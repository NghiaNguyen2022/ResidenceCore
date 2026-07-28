from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
ASSETS = ROOT / "docs" / "user-guide-assets" / "20260727"
OUTPUT = ROOT / "docs" / "HUONG_DAN_SU_DUNG_RESIDENCECORE_LUU_XA_NU_BMT.docx"

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
INK = RGBColor(31, 41, 55)
MUTED = RGBColor(91, 105, 120)
LIGHT_BLUE = "E8EEF5"
LIGHT_YELLOW = "FFF8E8"

FEATURE_GROUPS = [
    (1, "Khởi động và kiểm tra hệ thống", 1, 3, [
        "Đăng nhập hệ thống",
        "Kiểm tra Dashboard sau khi reset",
        "Kiểm tra cơ cấu tổ chức",
    ]),
    (2, "Hồ sơ học viên và lưu trú", 4, 10, [
        "Tạo hồ sơ nữ sinh",
        "Thêm liên hệ gia đình",
        "Tạo phòng nữ sinh",
        "Gán nữ sinh vào phòng",
        "Khai báo thông tin học tập",
        "Tạo lịch học",
        "Kiểm tra tài khoản học viên",
    ]),
    (3, "Tổ chức, công tác và sinh hoạt", 11, 18, [
        "Phân học viên vào Tổ và Ban",
        "Bổ nhiệm chức vụ nhiệm kỳ",
        "Kiểm tra cơ cấu và lịch sử bổ nhiệm",
        "Tạo và theo dõi công tác",
        "Thiết lập và thực hiện điểm danh",
        "Bổ sung khung giờ sinh hoạt hằng ngày",
        "Tạo hoạt động và sự kiện",
        "Thiết lập nội quy và nhắc nhở",
    ]),
    (4, "Quản lý cửa hàng", 19, 22, [
        "Tạo dữ liệu sản phẩm",
        "Mua hàng và nhập kho",
        "Bán hàng và cập nhật tồn kho",
        "Chốt ngày và đẩy dữ liệu sang sổ chung",
    ]),
    (5, "Tài chính lưu xá", 23, 25, [
        "Tạo kỳ thu và khoản phải thu",
        "Thu tiền theo học viên",
        "Ghi nhận khoản chi ngoài học viên",
    ]),
    (6, "Thông báo và thiết lập", 26, 26, [
        "Thiết lập phương thức nhận thông báo",
    ]),
    (7, "Portal học viên", 27, 31, [
        "Đăng nhập và kiểm tra portal",
        "Kiểm tra thông báo công tác và kỳ thu",
        "Kiểm tra hoạt động và tài chính",
        "Phản hồi công tác cá nhân, Tổ và Ban",
        "Kiểm tra quyền vào cửa hàng",
    ]),
    (8, "Vai trò và phạm vi phụ trách", 32, 34, [
        "Kiểm tra vai trò và cơ cấu phụ trách",
        "Kiểm tra thành viên Tổ và Ban",
        "Theo dõi công tác theo vai trò phụ trách",
    ]),
]

_active_group_number = None


def set_font(run, size=11, bold=False, italic=False, color=INK):
    run.font.name = "Calibri"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = color


def set_cell_fill(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def add_page_field(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run()
    set_font(run, 9, color=MUTED)
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char1, instr, fld_char2])


def add_callout(doc, title, body, fill=LIGHT_YELLOW):
    table = doc.add_table(rows=1, cols=1)
    table.autofit = False
    table.columns[0].width = Inches(6.5)
    cell = table.cell(0, 0)
    cell.width = Inches(6.5)
    set_cell_fill(cell, fill)
    set_cell_margins(cell, top=130, bottom=130, start=160, end=160)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(title)
    set_font(r, 11, bold=True, color=DARK_BLUE)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    r2 = p2.add_run(body)
    set_font(r2, 10.5)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def create_restarted_numbering(doc):
    numbering = doc.part.numbering_part.element
    style = doc.styles["List Number"]
    style_num_id = style._element.pPr.numPr.numId.val
    style_num = next(
        node
        for node in numbering.findall(qn("w:num"))
        if int(node.get(qn("w:numId"))) == int(style_num_id)
    )
    abstract_num_id = style_num.find(qn("w:abstractNumId")).get(qn("w:val"))

    existing_ids = [
        int(node.get(qn("w:numId")))
        for node in numbering.findall(qn("w:num"))
    ]
    new_num_id = max(existing_ids, default=0) + 1

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(new_num_id))
    abstract = OxmlElement("w:abstractNumId")
    abstract.set(qn("w:val"), abstract_num_id)
    num.append(abstract)
    override = OxmlElement("w:lvlOverride")
    override.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:startOverride")
    start.set(qn("w:val"), "1")
    override.append(start)
    num.append(override)
    numbering.append(num)

    return new_num_id


def add_numbered_steps(doc, steps):
    num_id = create_restarted_numbering(doc)
    for text in steps:
        p = doc.add_paragraph()
        p_pr = p._p.get_or_add_pPr()
        num_pr = p_pr.get_or_add_numPr()
        num_pr.get_or_add_ilvl().set(qn("w:val"), "0")
        num_pr.get_or_add_numId().set(qn("w:val"), str(num_id))
        p.paragraph_format.left_indent = Inches(0.375)
        p.paragraph_format.first_line_indent = Inches(-0.188)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.25
        r = p.add_run(text)
        set_font(r, 11)


def add_bullets(doc, items):
    for text in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.left_indent = Inches(0.375)
        p.paragraph_format.first_line_indent = Inches(-0.188)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.25
        r = p.add_run(text)
        set_font(r, 11)


def add_figure(doc, filename, caption, width=6.3):
    path = ASSETS / filename
    if not path.exists():
        raise FileNotFoundError(path)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(3)
    shape = p.add_run().add_picture(str(path), width=Inches(width))
    shape._inline.docPr.set("descr", caption)
    shape._inline.docPr.set("title", caption)
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(8)
    r = cap.add_run(caption)
    set_font(r, 9.5, italic=True, color=MUTED)


def get_feature_group(step_number):
    for group in FEATURE_GROUPS:
        if group[2] <= step_number <= group[3]:
            return group
    raise ValueError(f"Không tìm thấy nhóm cho bước {step_number}")


def add_step_section(doc, number, title, purpose, actions, figures, checks):
    global _active_group_number

    group_number, group_title, group_start, _, _ = get_feature_group(number)
    is_first_in_group = _active_group_number != group_number
    if is_first_in_group:
        group_heading = doc.add_heading(
            f"Nhóm {group_number}. {group_title}",
            level=1,
        )
        group_heading.paragraph_format.page_break_before = True
        _active_group_number = group_number

    local_step_number = number - group_start + 1
    h = doc.add_heading(
        f"{group_number}.{local_step_number}. {title}",
        level=2,
    )
    if not is_first_in_group:
        h.paragraph_format.page_break_before = True
    p = doc.add_paragraph()
    r = p.add_run("Mục đích: ")
    set_font(r, 11, bold=True, color=DARK_BLUE)
    r = p.add_run(purpose)
    set_font(r, 11)
    doc.add_heading("Cách thực hiện", level=3)
    add_numbered_steps(doc, actions)
    for filename, caption in figures:
        add_figure(doc, filename, caption)
    doc.add_heading("Kiểm tra sau thao tác", level=3)
    add_bullets(doc, checks)


doc = Document()
section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.top_margin = Inches(1)
section.bottom_margin = Inches(1)
section.left_margin = Inches(1)
section.right_margin = Inches(1)
section.header_distance = Inches(0.492)
section.footer_distance = Inches(0.492)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Calibri"
normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
normal.font.size = Pt(11)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.25

for name, size, color, before, after in (
    ("Heading 1", 16, BLUE, 18, 10),
    ("Heading 2", 13, BLUE, 14, 7),
    ("Heading 3", 12, DARK_BLUE, 10, 5),
):
    style = styles[name]
    style.font.name = "Calibri"
    style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    style.font.size = Pt(size)
    style.font.bold = True
    style.font.color.rgb = color
    style.paragraph_format.space_before = Pt(before)
    style.paragraph_format.space_after = Pt(after)
    style.paragraph_format.keep_with_next = True

header = section.header
hp = header.paragraphs[0]
hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
hr = hp.add_run("RESIDENCECORE  |  SỔ TAY VẬN HÀNH LƯU XÁ NỮ")
set_font(hr, 9, bold=True, color=MUTED)

footer = section.footer
fp = footer.paragraphs[0]
fr = fp.add_run("Hướng dẫn nội bộ  •  ")
set_font(fr, 9, color=MUTED)
add_page_field(fp)

# Editorial-cover opening.
doc.add_paragraph().paragraph_format.space_after = Pt(82)
kicker = doc.add_paragraph()
kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
kr = kicker.add_run("SỔ TAY VẬN HÀNH")
set_font(kr, 11, bold=True, color=BLUE)

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.paragraph_format.space_after = Pt(10)
tr = title.add_run("HƯỚNG DẪN SỬ DỤNG\nRESIDENCECORE")
set_font(tr, 28, bold=True, color=DARK_BLUE)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle.paragraph_format.space_after = Pt(28)
sr = subtitle.add_run("Quản lý lưu xá nữ tại Buôn Ma Thuột, Đắk Lắk")
set_font(sr, 15, color=BLUE)

context = doc.add_paragraph()
context.alignment = WD_ALIGN_PARAGRAPH.CENTER
context.paragraph_format.space_after = Pt(75)
cr = context.add_run(
    "Dành cho lưu xá do các sơ phụ trách, tiếp nhận nữ học sinh và sinh viên "
    "từ các địa phương lân cận đến Buôn Ma Thuột học tập."
)
set_font(cr, 11.5, italic=True, color=MUTED)

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
mr = meta.add_run("Phiên bản hướng dẫn: 28/07/2026  |  Môi trường kiểm thử nội bộ")
set_font(mr, 10, bold=True, color=MUTED)

doc.add_page_break()

doc.add_heading("Danh mục nội dung theo nhóm tính năng", level=1)
doc.add_paragraph(
    "Tài liệu được tổ chức theo nhóm nghiệp vụ. Trong mỗi nhóm, số mục bắt đầu lại "
    "từ 1 và được trình bày theo dạng Nhóm.Mục để dễ tra cứu."
)
for group_number, group_title, _, _, feature_titles in FEATURE_GROUPS:
    group_paragraph = doc.add_paragraph()
    group_paragraph.paragraph_format.space_before = Pt(8)
    group_paragraph.paragraph_format.space_after = Pt(3)
    group_run = group_paragraph.add_run(f"Nhóm {group_number}. {group_title}")
    set_font(group_run, 11.5, bold=True, color=DARK_BLUE)
    add_bullets(
        doc,
        [
            f"{group_number}.{feature_number}. {feature_title}"
            for feature_number, feature_title in enumerate(feature_titles, start=1)
        ],
    )

support_paragraph = doc.add_paragraph()
support_paragraph.paragraph_format.space_before = Pt(8)
support_run = support_paragraph.add_run("Nhóm 9. Hỗ trợ vận hành và kiểm thử")
set_font(support_run, 11.5, bold=True, color=DARK_BLUE)
add_bullets(
    doc,
    [
        "9.1. Xử lý khi không thấy menu Người dùng",
        "9.2. Checklist vận hành tiếp theo",
        "9.3. Nhật ký kiểm thử của tài liệu",
    ],
)

doc.add_page_break()
doc.add_heading("Giới thiệu, phạm vi và nguyên tắc sử dụng", level=1)
doc.add_paragraph(
    "Tài liệu hướng dẫn quy trình thiết lập dữ liệu nền và tiếp nhận một nữ sinh "
    "vào lưu xá từ trạng thái hệ thống sạch. Các màn hình được chụp trực tiếp "
    "trên ResidenceCore trong lúc thực hiện lại kịch bản."
)
add_callout(
    doc,
    "Bảo mật tài khoản",
    "Tài liệu chỉ nêu username quản trị là “admin”; mật khẩu phải được bàn giao "
    "riêng cho sơ hoặc người quản lý có thẩm quyền. Không gửi mật khẩu trong nhóm chat.",
)
doc.add_heading("Ngữ cảnh dữ liệu mẫu", level=2)
add_bullets(
    doc,
    [
        "Lưu xá nữ đặt tại Buôn Ma Thuột, tỉnh Đắk Lắk.",
        "Đơn vị vận hành: các sơ phụ trách lưu xá.",
        "Đối tượng: nữ học sinh, sinh viên từ Buôn Hồ và các khu vực lân cận về Buôn Ma Thuột học.",
        "Hồ sơ minh họa là dữ liệu kiểm thử, không phải thông tin người thật.",
    ],
)
doc.add_heading("Chuỗi nghiệp vụ trong hướng dẫn", level=2)
add_numbered_steps(
    doc,
    [
        "Đăng nhập bằng tài khoản quản lý lưu xá.",
        "Kiểm tra Dashboard sau khi làm sạch dữ liệu.",
        "Kiểm tra cơ cấu tổ chức, nhiệm kỳ và các vị trí trống.",
        "Tạo hồ sơ nữ sinh và tài khoản đăng nhập.",
        "Thêm liên hệ gia đình.",
        "Tạo phòng nữ sinh và gán học viên vào phòng.",
        "Khai báo trường, ngành/lớp và năm học.",
        "Tạo lịch học để hỗ trợ tránh xung đột khi phân công công tác.",
        "Kiểm tra tài khoản học viên tại Người dùng & quyền truy cập.",
    ],
)

add_step_section(
    doc,
    1,
    "Đăng nhập hệ thống",
    "Xác thực người quản lý trước khi truy cập dữ liệu lưu xá.",
    [
        "Mở ResidenceCore trên trình duyệt.",
        "Nhập username “admin”.",
        "Nhập mật khẩu được bàn giao riêng.",
        "Chọn “Đăng nhập”.",
        "Xác nhận góc trái hiển thị tên người dùng và vai trò “Quản lý lưu xá”.",
    ],
    [("00-dang-nhap.png", "Hình 1. Màn hình đăng nhập ResidenceCore.")],
    [
        "Đăng nhập thành công và chuyển tới Dashboard.",
        "Vai trò hiển thị là “Quản lý lưu xá · Toàn hệ thống”.",
        "Nếu sai mật khẩu, hệ thống hiển thị thông báo và không cấp quyền truy cập.",
    ],
)

add_step_section(
    doc,
    2,
    "Kiểm tra Dashboard sau khi reset",
    "Xác nhận hệ thống đang ở trạng thái sạch trước khi nhập dữ liệu thật.",
    [
        "Mở mục “Dashboard”.",
        "Kiểm tra tổng học viên, tổng phòng, sức chứa và tỷ lệ chiếm dụng.",
        "Đối chiếu khu vực phân bổ công tác hôm nay.",
    ],
    [("01-dashboard-sau-reset.png", "Hình 2. Dashboard sạch: 0 học viên và 0 phòng.")],
    [
        "Tổng học viên lưu trú bằng 0.",
        "Phòng ở hiển thị 0/0 và tỷ lệ chiếm dụng 0%.",
        "Không có phân bổ công tác phát sinh.",
    ],
)

add_step_section(
    doc,
    3,
    "Kiểm tra cơ cấu tổ chức",
    "Đảm bảo nhiệm kỳ, Tổ/Ban và chức vụ mẫu còn nguyên nhưng chưa gán học viên.",
    [
        "Mở “Quản lý lưu trú”, chọn “Tổ chức lưu xá”.",
        "Kiểm tra nhiệm kỳ hiện tại.",
        "Kiểm tra bốn Tổ và bốn Ban.",
        "Xác nhận các vị trí hiển thị “Đang trống”.",
    ],
    [("02-co-cau-to-chuc.png", "Hình 3. Cơ cấu tổ chức được giữ lại sau reset.")],
    [
        "Nhiệm kỳ hiện tại là 2026-2027.",
        "Có 8 đơn vị hoạt động.",
        "Tổng số phân công đang phụ trách bằng 0.",
    ],
)

add_step_section(
    doc,
    4,
    "Tạo hồ sơ nữ sinh",
    "Tiếp nhận một nữ sinh từ vùng lân cận đến Buôn Ma Thuột học tập và lưu trú.",
    [
        "Mở “Quản lý lưu trú”, chọn “Học viên”.",
        "Chọn “Thêm học viên”.",
        "Nhập tên thánh “Maria” và họ tên “Nguyễn Thị An”.",
        "Giữ giới tính “Nữ”; nhập CCCD, điện thoại và địa chỉ thường trú.",
        "Bật “Tạo tài khoản đăng nhập cho học viên”.",
        "Nhập ghi chú tiếp nhận và chọn “Thêm học viên”.",
    ],
    [
        ("03-them-hoc-vien-bieu-mau.png", "Hình 4. Biểu mẫu thêm học viên."),
        ("04-hoc-vien-da-tao.png", "Hình 5. Hồ sơ nữ sinh đã được tạo và có tài khoản."),
    ],
    [
        "Tổng học viên tăng lên 1.",
        "Trạng thái là “Đang lưu trú”.",
        "Hệ thống tự sinh mã lưu trú.",
        "Thẻ hồ sơ hiển thị “Đã có tài khoản”.",
    ],
)

add_step_section(
    doc,
    5,
    "Thêm liên hệ gia đình",
    "Lưu thông tin người thân để các sơ có thể liên lạc khi cần.",
    [
        "Tại danh sách học viên, chọn “Xem”.",
        "Mở tab “Liên hệ”.",
        "Chọn “Thêm liên hệ”.",
        "Chọn quan hệ “Mẹ” và nhập họ tên, điện thoại, email, nghề nghiệp, địa chỉ.",
        "Chọn “Thêm liên hệ” để lưu.",
    ],
    [
        ("05-them-lien-he-bieu-mau.png", "Hình 6. Biểu mẫu liên hệ gia đình."),
        ("06-lien-he-da-tao.png", "Hình 7. Liên hệ chính xuất hiện trên thẻ học viên."),
    ],
    [
        "Thẻ học viên hiển thị mẹ và số điện thoại liên hệ.",
        "Chỉ số “Thiếu liên hệ” giảm về 0.",
        "Một học viên chỉ nên có tối đa một liên hệ Cha và một liên hệ Mẹ.",
    ],
)

add_step_section(
    doc,
    6,
    "Tạo phòng nữ sinh",
    "Thiết lập phòng ở trước khi tiếp nhận học viên vào phòng.",
    [
        "Mở mục “Quản lý phòng”.",
        "Chọn “+ Thêm phòng”.",
        "Nhập mã phòng “BMT-N101”.",
        "Chọn sức chứa 4 chỗ.",
        "Nhập ghi chú “Phòng nữ sinh - Khu nhà chính”.",
        "Chọn “Lưu” và đóng thông báo trình duyệt.",
    ],
    [
        ("07-them-phong-bieu-mau.png", "Hình 8. Biểu mẫu tạo phòng nữ sinh."),
        ("08-phong-da-tao.png", "Hình 9. Phòng BMT-N101 sau khi tạo."),
    ],
    [
        "Tổng phòng tăng lên 1.",
        "Tổng sức chứa bằng 4 và còn trống 4 chỗ.",
        "Mã phòng không được trùng với phòng đã có.",
    ],
)

add_step_section(
    doc,
    7,
    "Gán nữ sinh vào phòng",
    "Ghi nhận chính xác nơi ở hiện tại của học viên.",
    [
        "Mở hồ sơ “Maria Nguyễn Thị An”.",
        "Chọn “Gán phòng”.",
        "Giữ loại thao tác “Gán phòng mới”.",
        "Chọn “BMT-N101 - còn 4/4 chỗ”.",
        "Nhập lý do tiếp nhận đầu năm học.",
        "Chọn “Lưu thao tác phòng”.",
    ],
    [
        ("09-gan-phong-bieu-mau.png", "Hình 10. Biểu mẫu gán phòng."),
        ("10-hoc-vien-da-gan-phong.png", "Hình 11. Hồ sơ ổn định sau khi gán phòng."),
    ],
    [
        "Thẻ học viên hiển thị phòng BMT-N101.",
        "Chỉ số “Chưa có phòng” giảm về 0.",
        "Hồ sơ chuyển sang trạng thái “Hồ sơ ổn định”.",
    ],
)

add_step_section(
    doc,
    8,
    "Khai báo thông tin học tập",
    "Theo dõi trường, ngành/lớp và năm học của nữ sinh đang lưu trú.",
    [
        "Mở đường dẫn “Thông tin học hành”.",
        "Chọn học viên “Maria Nguyễn Thị An - BMT-N101”.",
        "Chọn “Thêm”.",
        "Nhập trường “Trường Đại học Tây Nguyên”.",
        "Chọn bậc “Đại học”; nhập ngành/lớp và năm học 2026-2027.",
        "Chọn “Lưu”.",
    ],
    [
        ("11-thong-tin-hoc-tap-bieu-mau.png", "Hình 12. Biểu mẫu thông tin học tập."),
        ("12-thong-tin-hoc-tap-da-luu.png", "Hình 13. Thông tin Đại học Tây Nguyên đã lưu."),
    ],
    [
        "Tên trường, bậc học, ngành/lớp và năm học hiển thị đúng.",
        "Hệ thống hiển thị thông báo “Đã lưu thông tin học hành”.",
    ],
)

add_step_section(
    doc,
    9,
    "Tạo lịch học",
    "Cung cấp dữ liệu để hệ thống cảnh báo hoặc tránh phân công công tác trùng giờ học.",
    [
        "Mở mục “Lịch học”.",
        "Chọn học viên “Nguyễn Thị An - BMT-N101”.",
        "Chọn “Thêm lịch học”.",
        "Chọn Thứ Hai, nhập môn “Ngữ văn Việt Nam”.",
        "Giữ giờ 07:30-11:00 và nhập địa điểm tại Trường Đại học Tây Nguyên.",
        "Chọn “Thêm”.",
    ],
    [
        ("13-them-lich-hoc-bieu-mau.png", "Hình 14. Biểu mẫu thêm lịch học."),
        ("14-lich-hoc-da-tao.png", "Hình 15. Lịch học đã được thêm vào tuần."),
    ],
    [
        "Số môn học và số buổi/tuần cùng bằng 1.",
        "Bảng hiển thị đúng Thứ Hai, giờ học, môn học và địa điểm.",
        "Dữ liệu có thể dùng ở bước phân công công tác để kiểm tra xung đột.",
    ],
)

add_step_section(
    doc,
    10,
    "Kiểm tra tài khoản học viên",
    "Xác nhận tài khoản đăng nhập được tạo đúng và nhận biết trạng thái cần đổi mật khẩu.",
    [
        "Mở nhóm “Báo cáo & Thiết lập” trên menu bên trái.",
        "Chọn “Người dùng & quyền truy cập”.",
        "Tìm theo tên học viên hoặc username.",
        "Kiểm tra tên, vai trò, trạng thái hoạt động, hồ sơ liên kết và yêu cầu đổi mật khẩu.",
        "Việc khóa/mở lại tài khoản học viên phải thực hiện từ hồ sơ học viên; không chỉnh role trực tiếp tại màn hình này.",
    ],
    [("15-quan-ly-nguoi-dung.png", "Hình 16. Tài khoản học viên an.nguyen trong Người dùng & quyền truy cập.")],
    [
        "Tổng tài khoản bằng 2: một quản lý và một học viên.",
        "Tài khoản học viên là “an.nguyen”, vai trò “Học viên”, trạng thái “Đang hoạt động”.",
        "Hồ sơ liên kết là Nguyễn Thị An.",
        "Trạng thái “Cần đổi” nghĩa là học viên phải đổi mật khẩu ở lần đăng nhập đầu tiên.",
    ],
)

add_step_section(
    doc,
    11,
    "Phân học viên vào Tổ và Ban",
    "Tạo dữ liệu nền tổ chức để quản lý sinh hoạt và phân công công tác theo nhóm.",
    [
        "Mở “Cơ cấu tổ chức” và chọn tab “Tổ / Ban”.",
        "Ở từng Tổ, chọn “Thành viên”, chọn học viên rồi bấm “Thêm thành viên”.",
        "Phân đủ 19 học viên theo tải: Tổ 1 = 5, Tổ 2 = 5, Tổ 3 = 5, Tổ 4 = 4.",
        "Tiếp tục phân mỗi học viên vào một Ban phục vụ; kết quả bốn Ban là 5, 5, 5 và 4 thành viên.",
        "Khi cần đổi Tổ, dùng chức năng “Đổi tổ”; một học viên chỉ được thuộc một Tổ đang hoạt động tại một thời điểm.",
    ],
    [
        ("16-phan-thanh-vien-to-bieu-mau.png", "Hình 17. Biểu mẫu chọn học viên để thêm vào Tổ."),
        ("17-thanh-vien-to-1-da-them.png", "Hình 18. Thành viên đầu tiên đã được thêm vào Tổ 1."),
        ("19-to-1-du-5-thanh-vien.png", "Hình 19. Tổ 1 đã đủ 5 học viên sau khi hoàn thiện dữ liệu nền."),
    ],
    [
        "Tổng số học viên đang lưu trú là 19 và tất cả đều có một Tổ.",
        "Số thành viên các Tổ là 5–5–5–4; số thành viên các Ban là 5–5–5–4.",
        "Danh sách chi tiết của Tổ 1 hiển thị đủ 5 người và đúng phòng đang ở.",
    ],
)

add_step_section(
    doc,
    12,
    "Bổ nhiệm chức vụ nhiệm kỳ 2026–2027",
    "Thiết lập ban điều hành, Tổ trưởng và Trưởng ban để hoàn chỉnh cơ cấu trách nhiệm.",
    [
        "Tại “Cơ cấu tổ chức”, chọn “Bổ nhiệm”.",
        "Chọn nhiệm kỳ “Nhiệm kỳ 2026 - 2027”, vai trò, học viên và Tổ/Ban nếu vai trò yêu cầu.",
        "Bổ nhiệm ban điều hành: Trưởng, Phó, Thư ký và Thủ quỹ.",
        "Bổ nhiệm một Tổ trưởng cho mỗi Tổ và một Trưởng ban cho mỗi Ban.",
        "Không bổ nhiệm trùng một vai trò cho cùng học viên; mỗi Tổ/Ban chỉ có tối đa một người đứng đầu.",
    ],
    [("18-co-cau-12-phan-cong.png", "Hình 20. Cơ cấu hiện tại sau 12 lượt bổ nhiệm.")],
    [
        "Có 12 phân công đang hiệu lực: 4 chức vụ cấp lưu xá, 4 Tổ trưởng và 4 Trưởng ban.",
        "Nhiệm kỳ đang áp dụng từ 01/06/2026 đến 31/05/2027.",
        "Các vị trí Trưởng, Phó, Thư ký và Thủ quỹ đều có người phụ trách.",
    ],
)

add_step_section(
    doc,
    13,
    "Kiểm tra cơ cấu và lịch sử bổ nhiệm",
    "Đối chiếu dữ liệu hiện tại với lịch sử để bảo đảm trách nhiệm được truy vết.",
    [
        "Chọn tab “Cơ cấu hiện tại” để xem ban điều hành, Tổ trưởng và Trưởng ban.",
        "Mở “Thành viên” của từng đơn vị để kiểm tra danh sách thực tế.",
        "Chọn tab “Lịch sử bổ nhiệm”.",
        "Mở cây nhiệm kỳ và đối chiếu từng vai trò, người phụ trách, đơn vị và ngày bắt đầu.",
        "Khi kết thúc một nhiệm vụ, dùng nút “Kết thúc”; không xóa lịch sử đã phát sinh.",
    ],
    [("20-lich-su-12-bo-nhiem.png", "Hình 21. Lịch sử 12 lượt bổ nhiệm trong nhiệm kỳ hiện tại.")],
    [
        "Lịch sử hiển thị đủ 12 dòng và tất cả có trạng thái “Đang giữ vị trí”.",
        "Ngày bắt đầu là 27/07/2026 và ngày kết thúc là “Hiện tại”.",
        "Lưu ý UAT: số “người” trên thẻ Tổ/Ban hiện đang đếm người được bổ nhiệm, không phải toàn bộ thành viên; cần xem danh sách “Thành viên” để có số thực tế.",
    ],
)

add_step_section(
    doc,
    14,
    "Tạo và theo dõi công tác",
    "Kiểm tra bốn kiểu giao việc và vòng đời từ chưa làm đến hoàn thành.",
    [
        "Mở “Công tác & Phân công”, chọn “Mẫu công tác” để kiểm tra các mẫu đã giữ lại sau reset.",
        "Chọn “Tạo phân công”, nhập ngày 27/07/2026 và chọn mẫu.",
        "Tạo bốn bản ghi đại diện: giao cho Học viên, Tổ 1, phòng BMT-N102 và Ban Thanh nhạc.",
        "Với học viên Nguyễn Thị An, chọn “Dọn phòng chung ca sáng” 07:30–08:30 để kiểm thử trùng lịch học Thứ Hai 07:30–11:00.",
        "Sau khi lưu, mở “Công tác hôm nay”, kiểm tra tên công tác, đối tượng, ghi chú và trạng thái.",
        "Chọn “Hoàn thành” cho công tác đầu tiên để kiểm thử chuyển trạng thái.",
    ],
    [
        ("21-danh-sach-mau-cong-tac.png", "Hình 22. Danh sách mẫu công tác được giữ lại sau reset."),
        ("22-form-cong-tac-trung-lich-hoc.png", "Hình 23. Phân công cố ý trùng lịch học để kiểm thử cảnh báo."),
        ("24-bon-kieu-phan-cong-cong-tac.png", "Hình 24. Bốn kiểu giao công tác: học viên, Tổ, phòng và Ban."),
        ("25-cong-tac-hoan-thanh.png", "Hình 25. Công tác học viên đã chuyển sang trạng thái Hoàn thành."),
    ],
    [
        "Tạo và lưu được cả bốn kiểu đối tượng giao việc.",
        "Chuyển trạng thái “Hoàn thành” hoạt động đúng.",
        "CHƯA ĐẠT: hệ thống mới hiển thị thông báo hook kiểm tra lịch học, chưa cảnh báo hoặc chặn phân công trùng lịch.",
        "CHƯA ĐẠT: giờ hiển thị sau lưu lệch +7 giờ (07:30–08:30 thành 14:30–15:30); cần sửa xử lý múi giờ trước khi vận hành chính thức.",
    ],
)

add_callout(
    doc,
    "Kết luận kiểm thử công tác",
    "Luồng tạo và cập nhật trạng thái hoạt động, nhưng chưa được nghiệm thu để vận hành thật "
    "cho tới khi hoàn thiện kiểm tra xung đột lịch học và sửa lỗi hiển thị múi giờ.",
    fill="FDECEC",
)

add_step_section(
    doc,
    15,
    "Thiết lập và thực hiện điểm danh",
    "Tạo lịch điểm danh hằng ngày, ghi nhận tình trạng của 19 học viên và kiểm tra dữ liệu sau khi lưu.",
    [
        "Mở “Sinh hoạt & Đời sống” > “Lịch điểm danh”.",
        "Chọn “Thêm lịch”, nhập tên “Điểm danh tối”, loại “Điểm danh vào” và giờ 21:00.",
        "Chọn “Thêm mới”, sau đó kiểm tra dòng lịch phải hiển thị đúng 21:00 và tần suất “Hàng ngày”.",
        "Mở “Điểm danh”; chọn ngày 27/07/2026 và lịch “Điểm danh tối · 21:00”.",
        "Chọn “Tất cả có mặt” để tạo trạng thái nền nhanh cho toàn bộ danh sách.",
        "Đổi Bùi Thị Quỳnh Như thành “Đi trễ”, Dương Thị Hải Yến thành “Vắng phép” và Đặng Thị Phương Thảo thành “Vắng”.",
        "Mở “Ghi chú” của người vắng phép, nhập lý do và xác nhận đã báo sơ phụ trách.",
        "Mở “Ghi chú” của người vắng không phép, ghi yêu cầu liên hệ xác minh.",
        "Kiểm tra các thẻ thống kê phải là 19 tổng, 16 có mặt, 1 đi trễ, 1 vắng phép và 1 vắng.",
        "Chọn “Lưu điểm danh”; tải lại trang để xác nhận trạng thái và ghi chú vẫn còn.",
        "Dùng bộ lọc “Vắng phép” để lập danh sách cần theo dõi.",
    ],
    [
        ("26-diem-danh-chua-co-lich.png", "Hình 26. Màn hình điểm danh khi chưa có lịch."),
        ("28-tao-lich-diem-danh-toi.png", "Hình 27. Biểu mẫu tạo lịch Điểm danh tối."),
        ("30-lich-diem-danh-toi-21h.png", "Hình 28. Lịch Điểm danh tối đã hiển thị đúng 21:00."),
        ("31-diem-danh-16-1-1-1.png", "Hình 29. Kết quả điểm danh 16 có mặt, 1 trễ, 1 vắng phép và 1 vắng."),
        ("32-loc-hoc-vien-vang-phep.png", "Hình 30. Bộ lọc chỉ hiển thị học viên vắng phép và lý do."),
    ],
    [
        "Danh sách lấy đủ 19 học viên và đúng phòng hiện tại.",
        "Lưu hàng loạt thành công; tải lại trang vẫn giữ đúng 16–1–1–1 và hai ghi chú.",
        "Bộ lọc trạng thái hoạt động đúng.",
        "CẦN THEO DÕI UAT: lần tạo đầu tiên đã lưu giờ mặc định 07:00 dù biểu mẫu được nhập 21:00; phải kiểm tra lại dòng lịch sau khi lưu. Dữ liệu kiểm thử đã được hiệu chỉnh về 21:00 trước khi điểm danh.",
    ],
)

add_callout(
    doc,
    "Xử lý học viên vắng",
    "Vắng phép phải có lý do và người tiếp nhận thông tin. Vắng không phép cần được liên hệ xác minh "
    "ngay sau giờ điểm danh; không tự chuyển thành vắng phép khi chưa có xác nhận.",
    fill=LIGHT_BLUE,
)

add_step_section(
    doc,
    16,
    "Bổ sung khung giờ sinh hoạt hằng ngày",
    "Hoàn thiện lịch ngày thường bằng một khung giờ điểm danh tối liên kết với quy trình điểm danh.",
    [
        "Mở nhóm “Sinh hoạt” và chọn “Sinh hoạt hằng ngày”.",
        "Chọn tab “Sinh hoạt” để xem các khung giờ đang áp dụng; dữ liệu ban đầu có 11 khung giờ.",
        "Chọn “Thêm khung giờ”.",
        "Giữ mẫu “Lịch ngày thường”.",
        "Chọn giờ bắt đầu 21:00 và giờ kết thúc 21:15 từ danh sách giờ.",
        "Nhập tên hoạt động “Điểm danh tối” và địa điểm “Phòng sinh hoạt chung”.",
        "Nhập thứ tự 11, giữ trạng thái “Đang áp dụng”.",
        "Nhập ghi chú: “Kiểm tra hiện diện, ghi nhận vắng phép và trường hợp cần liên hệ.”",
        "Chọn “Lưu”.",
        "Tải lại trang, mở lại tab “Sinh hoạt” và kiểm tra lịch tăng lên 12 khung giờ.",
    ],
    [
        ("36-lich-sinh-hoat-11-khung-gio.png", "Hình 31. Lịch ngày thường trước khi bổ sung, gồm 11 khung giờ."),
        ("37-them-khung-gio-diem-danh-toi.png", "Hình 32. Biểu mẫu thêm khung giờ Điểm danh tối."),
        ("38-lich-sinh-hoat-co-diem-danh-toi.png", "Hình 33. Lịch sinh hoạt sau khi bổ sung khung giờ 21:00–21:15."),
    ],
    [
        "Lịch ngày thường có 12 khung giờ đang áp dụng.",
        "Khung “Điểm danh tối” hiển thị đúng 21:00–21:15.",
        "Dữ liệu vẫn tồn tại sau khi tải lại trang.",
        "Khung giờ kết nối hợp lý với lịch điểm danh 21:00 đã tạo ở mục 3.5.",
    ],
)

add_callout(
    doc,
    "Lưu ý chọn giờ",
    "Đối với biểu mẫu khung giờ sinh hoạt, nên chọn giờ từ danh sách 15 phút thay vì chỉ nhập bằng bàn phím; "
    "sau khi lưu phải kiểm tra lại thời gian hiển thị trên lịch.",
    fill=LIGHT_BLUE,
)

add_step_section(
    doc,
    17,
    "Tạo hoạt động và sự kiện",
    "Tạo một hoạt động cộng đoàn, công bố trên portal và chuẩn bị dữ liệu người tham gia để theo dõi.",
    [
        "Mở nhóm “Sinh hoạt” và chọn “Hoạt động / Sự kiện”.",
        "Chọn “Tạo hoạt động”.",
        "Giữ tùy chọn “Portal” để học viên có thể xem hoạt động.",
        "Nhập mã “SHCD-20260727” và tên “Sinh hoạt cộng đoàn đầu tuần”.",
        "Chọn loại “Sinh hoạt chung”, ngày 27/07/2026, giờ 19:00–20:30.",
        "Nhập địa điểm “Phòng sinh hoạt chung” và phụ trách “Ban sinh hoạt”.",
        "Nhập mô tả: cầu nguyện, phổ biến lịch tuần, nhắc công tác và lắng nghe nhu cầu của nữ sinh.",
        "Giữ trạng thái “Dự kiến” và chọn “Tạo hoạt động”.",
        "Kiểm tra các thẻ thống kê: Tổng = 1, Dự kiến = 1.",
        "Lọc trạng thái “Dự kiến” và tìm mã “SHCD-20260727” để xác nhận hoạt động hiển thị đúng.",
    ],
    [
        ("39-hoat-dong-tong-quan-desktop.png", "Hình 34. Tổng quan hoạt động sau khi tạo dữ liệu, chụp toàn màn hình desktop."),
        ("40-form-hoat-dong-desktop.png", "Hình 35. Biểu mẫu hoạt động hiển thị đầy đủ trên màn hình desktop."),
    ],
    [
        "Hoạt động hiển thị đúng ngày, khung giờ, địa điểm, ban phụ trách và trạng thái.",
        "Hoạt động được bật hiển thị trên portal.",
        "Dữ liệu nền có 19 học viên tham gia; 17 người được đánh dấu tham dự và 2 người vắng theo kết quả điểm danh.",
        "Tìm kiếm theo mã và lọc trạng thái Dự kiến hoạt động đúng.",
        "CHƯA ĐẠT: màn hình Hoạt động chưa có giao diện thêm/xóa người tham gia hoặc ghi nhận tham dự; dữ liệu nền hiện phải được đồng bộ bằng lớp dữ liệu.",
    ],
)

add_callout(
    doc,
    "Giới hạn quản lý người tham gia",
    "Database đã hỗ trợ người tham gia và trạng thái tham dự, nhưng màn hình quản trị chưa hiển thị danh sách này. "
    "Không xem bước người tham gia là hoàn tất cho vận hành thật cho tới khi bổ sung giao diện và API tương ứng.",
    fill="FDECEC",
)

add_step_section(
    doc,
    18,
    "Thiết lập nội quy và nhắc nhở",
    "Bổ sung quy định giờ trở về phù hợp với lưu xá nữ tại Buôn Ma Thuột và kiểm tra khả năng tra cứu.",
    [
        "Mở nhóm “Sinh hoạt” và chọn “Nội quy & Nhắc nhở”.",
        "Kiểm tra ba nội quy nền đang áp dụng, sau đó chọn “Thêm nội quy”.",
        "Nhập mã “CURFEW_BMT” và tên “Giờ trở về lưu xá buổi tối”.",
        "Chọn nhóm “Nề nếp / Tác phong”, mức độ “Nghiêm trọng”, ngày áp dụng 27/07/2026.",
        "Nhập nội dung: học viên trở về trước 21:00; trường hợp đặc biệt cần báo trước cho sơ phụ trách.",
        "Nhập hành vi mong đợi: có mặt đúng giờ, chủ động báo lý do và thời gian dự kiến khi về trễ.",
        "Nhập hướng xử lý: nhắc nhở lần đầu; trao đổi riêng và ghi nhận nếu tái diễn hoặc không báo trước.",
        "Giữ trạng thái “Đang áp dụng” và chọn “Thêm nội quy”.",
        "Tìm theo mã “CURFEW_BMT” để xác nhận card và dòng dữ liệu hiển thị đúng.",
    ],
    [
        ("43-them-noi-quy-desktop.png", "Hình 36. Biểu mẫu nội quy CURFEW_BMT trên màn hình desktop."),
        ("44-noi-quy-da-tao-desktop.png", "Hình 37. Nội quy đã tạo và được tìm thấy theo mã CURFEW_BMT."),
    ],
    [
        "Tổng nội quy tăng từ 3 lên 4; cả 4 nội quy đang áp dụng.",
        "Nhóm Nề nếp / Tác phong và mức Nghiêm trọng hiển thị đúng.",
        "Ngày áp dụng 27/07/2026, mô tả và hướng xử lý được lưu đúng.",
        "Tìm kiếm theo mã CURFEW_BMT trả về đúng một nội quy.",
    ],
)

add_callout(
    doc,
    "Cách xử lý trường hợp về trễ",
    "Sơ phụ trách nên ưu tiên xác minh an toàn của học viên, ghi nhận lý do và thời gian thực tế. "
    "Chỉ chuyển sang xử lý kỷ luật khi đã có đủ thông tin và xem xét hoàn cảnh cụ thể.",
    fill=LIGHT_BLUE,
)

add_step_section(
    doc,
    19,
    "Tạo dữ liệu sản phẩm cho cửa hàng",
    "Khai báo một mặt hàng thiết yếu, thiết lập giá bán ban đầu và chuẩn bị dữ liệu cho quy trình nhập kho.",
    [
        "Mở nhóm “Quản lý cửa hàng” và chọn “Dữ liệu sản phẩm”.",
        "Kiểm tra danh sách ban đầu chưa có mặt hàng, sau đó chọn “Thêm hàng hóa”.",
        "Nhập mã “NUOC_SUOI_500” và tên “Nước suối 500ml”.",
        "Chọn nhóm “Đồ uống”, đơn vị tính “Chai” và tồn tối thiểu 12.",
        "Nhập ghi chú: “Nước uống đóng chai phục vụ học viên và khách tại lưu xá.”",
        "Chọn “Thêm hàng hóa”. Sản phẩm mới có tồn kho 0 và chưa có giá vốn.",
        "Tại sản phẩm vừa tạo, chọn “Thông tin giá”, sau đó chọn “Cập nhật giá bán”.",
        "Nhập giá bán mới 5.000đ, ngày áp dụng 27/07/2026 và lý do “Điều chỉnh theo thực tế”.",
        "Nhập ghi chú “Giá bán khởi tạo cho cửa hàng lưu xá.” và chọn “Lưu giá bán”.",
        "Tìm theo mã “NUOC_SUOI_500”, tải lại trang và kiểm tra sản phẩm cùng giá bán vẫn hiển thị.",
    ],
    [
        ("45-them-san-pham-desktop.png", "Hình 38. Biểu mẫu thêm sản phẩm Nước suối 500ml trên màn hình desktop."),
        ("46-cap-nhat-gia-ban-desktop.png", "Hình 39. Biểu mẫu thiết lập giá bán ban đầu cho sản phẩm."),
        ("47-san-pham-da-tao-desktop.png", "Hình 40. Sản phẩm đã tạo, có giá bán 5.000đ và được tìm thấy theo mã."),
    ],
    [
        "Tổng mặt hàng tăng từ 0 lên 1.",
        "Sản phẩm thuộc nhóm Đồ uống, đơn vị chai và tồn tối thiểu 12.",
        "Giá bán 5.000đ được lưu đúng và còn hiển thị sau khi tải lại.",
        "Tìm kiếm theo mã NUOC_SUOI_500 trả về đúng sản phẩm.",
        "Cảnh báo Hết hàng hiển thị đúng vì tồn kho hiện tại bằng 0; giá vốn sẽ được hình thành ở bước nhập kho.",
    ],
)

add_callout(
    doc,
    "Phân biệt giá bán và giá vốn",
    "Giá bán được thiết lập trong Thông tin giá. Giá vốn không nhập trực tiếp ở màn hình sản phẩm mà được tính từ "
    "các lần mua hàng/nhập kho, giúp lưu lại lịch sử và phản ánh đúng chi phí thực tế.",
    fill=LIGHT_BLUE,
)

add_step_section(
    doc,
    20,
    "Mua hàng và nhập kho cửa hàng",
    "Tạo phiếu mua hàng nhiều mặt hàng, cập nhật tồn kho, hình thành giá vốn và tự động ghi nhận khoản chi cửa hàng.",
    [
        "Bảo đảm hệ thống đã có sổ “Cửa hàng lưu xá nữ BMT”. Nếu chưa có, liên hệ quản trị kỹ thuật để khởi tạo trước khi nhập kho.",
        "Mở nhóm “Quản lý cửa hàng” và chọn “Mua hàng / Nhập kho”.",
        "Chọn “Tạo phiếu nhập”, giữ nguồn nhập “Mua hàng” và ngày nhập 27/07/2026.",
        "Nhập nhà cung cấp “Đại lý hàng tiêu dùng Buôn Ma Thuột” và chọn thanh toán “Chuyển khoản”.",
        "Nhập Nước suối 500ml: 48 chai, giá mua 3.500đ/chai.",
        "Thêm dòng Bút bi xanh: 30 cái, giá mua 3.000đ/cái.",
        "Thêm dòng Cà phê hạt 500g: 10 gói, giá mua 85.000đ/gói.",
        "Thêm dòng Cà phê rang xay 500g: 10 gói, giá mua 70.000đ/gói.",
        "Thêm dòng Túi thổ cẩm: 6 cái, giá mua 125.000đ/cái.",
        "Kiểm tra tổng số lượng 104 và tổng giá trị nhập 2.558.000đ.",
        "Nhập ghi chú “Nhập kho ban đầu cho cửa hàng lưu xá nữ tại Buôn Ma Thuột.” rồi chọn “Lưu phiếu nhập”.",
        "Kiểm tra phiếu xem trước, lịch sử phiếu nhập, tồn kho từng sản phẩm và khoản chi mua hàng trong “Tổng hợp thu chi”.",
    ],
    [
        ("48-phieu-nhap-thong-tin-desktop.png", "Hình 41. Thông tin nhà cung cấp và năm mặt hàng trong phiếu nhập."),
        ("49-phieu-nhap-tong-cong-desktop.png", "Hình 42. Tổng số lượng, tổng giá trị và ghi chú trước khi lưu phiếu."),
        ("50-xem-truoc-phieu-nhap-desktop.png", "Hình 43. Bản xem trước phiếu mua hàng/nhập kho sau khi lưu."),
        ("51-lich-su-nhap-kho-desktop.png", "Hình 44. Lịch sử phiếu nhập và số liệu tổng hợp trên màn hình desktop."),
        ("52-ton-kho-sau-khi-nhap-desktop.png", "Hình 45. Tồn kho và giá vốn sản phẩm sau khi nhập hàng."),
    ],
    [
        "Tạo thành công một phiếu mua hàng gồm 5 mặt hàng và 104 đơn vị.",
        "Tổng giá trị nhập là 2.558.000đ; lịch sử phiếu hiển thị đúng nhà cung cấp và ngày nhập.",
        "Tồn kho sau nhập: 48 chai nước suối, 30 bút bi, 10 gói cà phê hạt, 10 gói cà phê rang xay và 6 túi thổ cẩm.",
        "Giá vốn từng sản phẩm bằng giá mua của lần nhập đầu tiên.",
        "Không còn sản phẩm thuộc nhóm sắp hết/hết hàng.",
        "Hệ thống tự động ghi một khoản chi “Mua hàng nhập kho” trị giá 2.558.000đ vào dòng tiền cửa hàng.",
        "CHƯA ĐẠT: giao diện hiện chưa có nút khả dụng để khởi tạo sổ cửa hàng khi dữ liệu sổ bị xóa; quản trị kỹ thuật phải tạo sổ nền trước.",
    ],
)

add_callout(
    doc,
    "Điều kiện trước khi nhập kho",
    "Nếu bấm Lưu phiếu nhập và nhận thông báo “Vui lòng khởi tạo cửa hàng trước”, không nhập lại nhiều lần. "
    "Hãy yêu cầu quản trị kỹ thuật khởi tạo sổ cửa hàng, tải lại trang rồi tạo phiếu. Sau khi phiếu được lưu, "
    "sản phẩm đã có giao dịch sẽ không thể xóa trực tiếp nhằm bảo toàn lịch sử tồn kho.",
    fill="FDECEC",
)

add_step_section(
    doc,
    21,
    "Bán hàng và cập nhật tồn kho",
    "Tạo phiếu bán nhiều mặt hàng, giảm tồn kho và tự động ghi nhận khoản thu bán hàng vào dòng tiền cửa hàng.",
    [
        "Mở nhóm “Quản lý cửa hàng” và chọn “Bán hàng”.",
        "Chọn “Tạo phiếu bán”, giữ ngày bán 27/07/2026 và phương thức “Tiền mặt”.",
        "Nhập khách hàng “Khách vãng lai”.",
        "Chọn Nước suối 500ml, số lượng 4; hệ thống lấy giá bán hiện hành 5.000đ/chai.",
        "Thêm Bút bi xanh, số lượng 2, giá bán 5.000đ/cái.",
        "Thêm Cà phê hạt 500g, số lượng 1, giá bán 120.000đ/gói.",
        "Thêm Cà phê rang xay 500g, số lượng 1, giá bán 95.000đ/gói.",
        "Thêm Túi thổ cẩm, số lượng 1, giá bán 180.000đ/cái.",
        "Kiểm tra tổng số lượng 9 và tổng giá trị bán 425.000đ.",
        "Nhập ghi chú “Phiếu bán thử nghiệm cho khách vãng lai tại lưu xá.” rồi chọn “Lưu phiếu bán”.",
        "Kiểm tra phiếu xem trước, lịch sử bán hàng, tồn kho sản phẩm và khoản thu trong “Tổng hợp thu chi”.",
    ],
    [
        ("53-phieu-ban-thong-tin-desktop.png", "Hình 46. Thông tin khách hàng và các dòng sản phẩm trong phiếu bán."),
        ("54-phieu-ban-tong-cong-desktop.png", "Hình 47. Tổng số lượng, tổng giá trị và ghi chú trước khi lưu phiếu bán."),
        ("55-xem-truoc-phieu-ban-desktop.png", "Hình 48. Bản xem trước phiếu bán hàng sau khi lưu."),
        ("56-lich-su-phieu-ban-desktop.png", "Hình 49. Lịch sử phiếu bán và doanh thu trên màn hình desktop."),
        ("57-ton-kho-sau-ban-desktop.png", "Hình 50. Tồn kho sản phẩm sau giao dịch bán hàng."),
        ("58-dong-tien-sau-ban-desktop.png", "Hình 51. Dòng tiền cửa hàng sau khi ghi nhận khoản thu bán hàng."),
    ],
    [
        "Tạo thành công một phiếu bán gồm 5 mặt hàng, tổng số lượng 9 và doanh thu 425.000đ.",
        "Tồn kho giảm đúng: nước suối còn 44 chai, bút bi còn 28 cái, hai loại cà phê còn 9 gói mỗi loại và túi thổ cẩm còn 5 cái.",
        "Hệ thống tự động ghi khoản thu bán hàng 425.000đ.",
        "Dòng tiền có 2 phát sinh: thu 425.000đ, chi 2.558.000đ, chênh lệch -2.133.000đ.",
        "Đối chiếu thủ công: giá vốn hàng đã bán là 300.000đ và lãi gộp của phiếu là 125.000đ.",
        "Dữ liệu vẫn hiển thị đúng sau khi chuyển giữa các màn hình Bán hàng, Dữ liệu sản phẩm và Tổng hợp thu chi.",
    ],
)

add_callout(
    doc,
    "Kiểm tra trước khi lưu phiếu bán",
    "Luôn kiểm tra tồn hiện có, số lượng bán và giá bán hiện hành. Phiếu bán sau khi lưu sẽ giảm tồn kho và tạo khoản thu; "
    "không dùng thao tác xóa trực tiếp sản phẩm để sửa sai giao dịch.",
    fill=LIGHT_BLUE,
)

add_step_section(
    doc,
    22,
    "Chốt ngày và đẩy dữ liệu sang sổ chung",
    "Kiểm tra toàn bộ thu chi trong ngày, thực hiện chốt của người lập và xác nhận để khóa dữ liệu, đưa số liệu sang sổ tài chính chung.",
    [
        "Mở “Quản lý cửa hàng” và chọn “Tổng hợp thu chi”.",
        "Kiểm tra ngày chốt là 27/07/2026, tổng thu 425.000đ, tổng chi 2.558.000đ và chênh lệch -2.133.000đ.",
        "Chọn “Chốt ngày”.",
        "Trong cửa sổ “Xem trước chốt ngày”, đối chiếu 2 phát sinh: một khoản thu bán hàng và một khoản chi mua hàng nhập kho.",
        "Chọn “Xác nhận chốt ngày”. Đây là thao tác của người lập, tạo lịch sử chốt ở trạng thái “Đã chốt · Chờ xác nhận”.",
        "Tại “Lịch sử chốt ngày”, chọn “Review”.",
        "Kiểm tra lại số lượng phát sinh, tổng thu, tổng chi và dòng tiền.",
        "Nếu số liệu sai, chọn “Bỏ chốt để bổ sung”. Nếu đúng, chọn “Xác nhận & đẩy sổ chung”.",
        "Kiểm tra trạng thái cuối cùng “Đã xác nhận” và “Đã đẩy sổ chung”.",
    ],
    [
        ("59-xem-truoc-chot-ngay-desktop.png", "Hình 52. Bản xem trước số liệu trước khi chốt ngày."),
        ("60-lich-su-da-chot-desktop.png", "Hình 53. Ngày đã chốt và đang chờ người có quyền xác nhận."),
        ("61-review-chot-ngay-desktop.png", "Hình 54. Màn hình Review trước khi xác nhận và đẩy sổ chung."),
        ("62-da-xac-nhan-day-so-chung-desktop.png", "Hình 55. Trạng thái đã xác nhận và đã đẩy sang sổ chung."),
    ],
    [
        "Tạo thành công mã chốt ngày CHOT-20260727-4.",
        "Bản chốt chứa đúng 2 phát sinh với tổng thu 425.000đ và tổng chi 2.558.000đ.",
        "Chênh lệch dòng tiền được tính đúng là -2.133.000đ.",
        "Sau thao tác đầu tiên, trạng thái chuyển thành “Đã chốt · Chờ xác nhận” và chưa đẩy sổ chung.",
        "Sau Review và xác nhận, trạng thái chuyển thành “Đã xác nhận · Đã đẩy sổ chung”.",
        "Dữ liệu ngày 27/07/2026 đã được khóa chính thức theo quy trình hai cấp.",
    ],
)

add_callout(
    doc,
    "Phân tách trách nhiệm chốt sổ",
    "Người lập thực hiện Chốt ngày; người có quyền xác nhận phải Review trước khi đẩy sang sổ chung. "
    "Không xác nhận nếu chưa đối chiếu chứng từ mua hàng, phiếu bán và tiền thực tế.",
    fill=LIGHT_BLUE,
)

add_step_section(
    doc,
    23,
    "Tạo kỳ thu và khoản phải thu học viên",
    "Thiết lập kỳ thu học kỳ I, sau đó sinh các khoản phí lưu trú và ăn uống cho toàn bộ nữ sinh trong tháng 08/2026.",
    [
        "Mở “Quản lý lưu trú” và chọn “Tài chính lưu xá”.",
        "Chọn “Tạo kỳ thu”.",
        "Nhập tên “Phí lưu xá học kỳ I năm 2026”, năm 2026, từ tháng 08 đến tháng 12.",
        "Nhập phí lưu trú 1.200.000đ, ăn uống sinh hoạt 1.800.000đ và khoản khác 200.000đ.",
        "Nhập ghi chú “Kỳ thu mẫu cho nữ sinh lưu xá tại Buôn Ma Thuột.” và chọn “Lưu kỳ thu”.",
        "Chọn tháng 08/2026 và mở “Tạo khoản phải thu”.",
        "Giữ chọn Phí lưu trú và Ăn uống sinh hoạt cho 19 học viên; không chọn Khoản thu khác.",
        "Kiểm tra dự kiến tạo 38 khoản cho 19 học viên, tổng 57.000.000đ.",
        "Chọn “Tạo khoản phải thu” và kiểm tra sổ phải thu tháng 08/2026.",
    ],
    [
        ("63-tao-ky-thu-desktop.png", "Hình 56. Biểu mẫu tạo kỳ thu học kỳ I năm 2026."),
        ("64-tao-khoan-phai-thu-desktop.png", "Hình 57. Chọn khoản phí cho các học viên trong tháng 08/2026."),
        ("65-so-phai-thu-thang-08-desktop.png", "Hình 58. Sổ phải thu tháng 08 sau khi sinh khoản."),
    ],
    [
        "Tạo thành công kỳ thu tháng 08–12/2026.",
        "Tháng 08 có 19 học viên và 38 khoản phải thu.",
        "Mỗi học viên phải thu 3.000.000đ gồm 1.200.000đ phí lưu trú và 1.800.000đ ăn uống.",
        "Tổng phải thu đúng của tháng 08 là 57.000.000đ.",
        "CHƯA ĐẠT: thẻ tổng quan “Còn phải thu kỳ” hiển thị 171.000.000đ, gấp ba số đúng trong sổ tháng.",
    ],
)

add_callout(
    doc,
    "Đối chiếu tổng phải thu",
    "Trước khi gửi thông báo cho phụ huynh, phải đối chiếu tổng theo sổ tháng: số học viên × tổng phí mỗi học viên. "
    "Trong lần kiểm thử này số đúng là 19 × 3.000.000đ = 57.000.000đ; không dùng số 171.000.000đ trên thẻ tổng quan.",
    fill="FDECEC",
)

add_step_section(
    doc,
    24,
    "Thu tiền theo học viên",
    "Ghi nhận một học viên thanh toán đủ hai khoản của tháng 08/2026 và kiểm tra trạng thái công nợ.",
    [
        "Trong sổ phải thu tháng 08/2026, chọn “Thu tiền” tại học viên Vũ Thị Bảo Ngọc.",
        "Kiểm tra kỳ thu, tháng và học viên đã được chọn đúng.",
        "Giữ chọn Phí lưu trú 1.200.000đ và Ăn uống sinh hoạt 1.800.000đ.",
        "Kiểm tra tổng thu lần này 3.000.000đ và còn lại sau thu 0đ.",
        "Giữ ngày thu 27/07/2026, chọn phương thức “Chuyển khoản”.",
        "Nhập ghi chú “Phụ huynh Vũ Thị Bảo Ngọc chuyển khoản đủ phí tháng 08/2026.”",
        "Chọn “Lưu thanh toán” và kiểm tra trạng thái “Đã thu đủ”.",
    ],
    [
        ("66-thu-tien-hoc-vien-desktop.png", "Hình 59. Biểu mẫu thu đủ hai khoản của một học viên."),
        ("67-hoc-vien-da-thu-du-desktop.png", "Hình 60. Học viên đã thu đủ và công nợ còn lại bằng 0."),
    ],
    [
        "Vũ Thị Bảo Ngọc được ghi nhận đã thu đủ 3.000.000đ.",
        "Hai khoản thành phần đều có trạng thái Đã thu và còn lại 0đ.",
        "Sổ tháng hiển thị đã thu 3.000.000đ, còn phải thu 54.000.000đ.",
        "CHƯA ĐẠT: thẻ tổng quan hiển thị đã thu 9.000.000đ và còn phải thu 162.000.000đ, tiếp tục sai theo hệ số ba.",
    ],
)

add_step_section(
    doc,
    25,
    "Ghi nhận khoản chi ngoài học viên",
    "Ghi nhận chi phí mua vật dụng vệ sinh trong kỳ và kiểm tra sổ dòng tiền lưu xá.",
    [
        "Trong “Tài chính lưu xá”, chọn “Thu chi khác” hoặc mở tab “Sổ dòng tiền” rồi chọn “Chi một lần”.",
        "Chọn loại nghiệp vụ “Khoản chi” và nhóm gợi ý “Vật dụng”.",
        "Nhập số tiền 450.000đ và ngày ghi nhận 28/08/2026.",
        "Nhập đơn vị nhận “Cửa hàng gia dụng Buôn Ma Thuột”.",
        "Nhập nội dung “Mua nước lau sàn, túi rác và dụng cụ vệ sinh khu sinh hoạt chung.”",
        "Chọn “Lưu nghiệp vụ”.",
        "Mở “Sổ dòng tiền” và kiểm tra khoản chi cùng số cân đối.",
    ],
    [
        ("68-ghi-nhan-khoan-chi-desktop.png", "Hình 61. Biểu mẫu ghi nhận khoản chi vật dụng."),
        ("69-so-dong-tien-co-khoan-chi-desktop.png", "Hình 62. Sổ dòng tiền hiển thị khoản chi 450.000đ."),
    ],
    [
        "Sổ dòng tiền có 1 dòng chi thật trị giá 450.000đ.",
        "Cân đối thu chi ngoài học viên trong kỳ là -450.000đ.",
        "Khoản chi hiển thị đúng đơn vị nhận, ngày 28/08/2026 và nội dung mua vật dụng vệ sinh.",
        "Khoản thu học viên 3.000.000đ vẫn được tổng hợp riêng, không trộn vào thẻ thu khác.",
    ],
)

add_step_section(
    doc,
    26,
    "Thiết lập phương thức nhận thông báo",
    "Kiểm tra các nhóm thông báo quản lý và bật thêm kênh email cho tài khoản đang đăng nhập.",
    [
        "Mở đường dẫn “Cài đặt thông báo” trong phần thiết lập tài khoản.",
        "Kiểm tra bốn nhóm thông báo công nợ, thanh toán, công việc và điểm danh đang bật.",
        "Giữ “Thông báo trong ứng dụng” và bật thêm “Thông báo qua email”.",
        "Chọn “Lưu cài đặt”.",
        "Tải lại trang và kiểm tra tùy chọn email vẫn ở trạng thái bật.",
    ],
    [
        ("70-cai-dat-thong-bao-desktop.png", "Hình 63. Bật kênh thông báo qua email trên màn hình desktop."),
        ("71-thong-bao-da-luu-desktop.png", "Hình 64. Thiết lập thông báo được giữ nguyên sau khi tải lại."),
    ],
    [
        "Cả bốn loại thông báo nghiệp vụ đều đang bật.",
        "Thông báo trong ứng dụng đang bật.",
        "Kênh email được bật, lưu thành công và còn hiệu lực sau khi tải lại.",
        "CHƯA ĐẠT VỀ ĐIỀU HƯỚNG: menu đơn giản hiện không có liên kết trực tiếp đến trang Cài đặt thông báo.",
    ],
)

add_step_section(
    doc,
    27,
    "Đăng nhập và kiểm tra portal học viên",
    "Đăng nhập bằng tài khoản học viên Nguyễn Thị An, hoàn tất đổi mật khẩu lần đầu và kiểm tra menu theo các vai trò được bổ nhiệm.",
    [
        "Đăng xuất tài khoản quản lý và đăng nhập bằng tài khoản học viên được bàn giao.",
        "Ở lần đăng nhập đầu tiên, nhập mật khẩu tạm thời và đặt mật khẩu mới theo quy định bảo mật.",
        "Đăng nhập lại sau khi đổi mật khẩu.",
        "Kiểm tra menu “Lưu xá của tôi” gồm Hồ sơ, Công tác, Cửa hàng, Tài chính, Thông báo và Hoạt động.",
        "Kiểm tra các vai trò Nguyễn Thị An đang giữ: Trưởng, Tổ trưởng Tổ 1 và Trưởng ban Ban Thanh nhạc.",
    ],
    [
        ("72-portal-thong-bao-ky-thu-desktop.png", "Hình 65. Portal học viên hiển thị thông báo kỳ thu mới."),
    ],
    [
        "Tài khoản học viên đăng nhập thành công sau khi đổi mật khẩu lần đầu.",
        "Portal nhận đúng hồ sơ Nguyễn Thị An và ba vai trò nghiệp vụ.",
        "Thông báo nổi về khoản thu tháng 08/2026 hiển thị ngay sau đăng nhập.",
    ],
)

add_step_section(
    doc,
    28,
    "Kiểm tra thông báo công tác và kỳ thu",
    "Xác nhận học viên nhận đủ thông báo phát sinh từ kỳ thu và các công tác được giao theo ba phạm vi.",
    [
        "Trong “Lưu xá của tôi”, chọn “Thông báo”.",
        "Kiểm tra số thông báo chưa đọc và phân loại theo Công tác, Tài chính, Hệ thống.",
        "Mở thông báo “Khoản thu mới” và đối chiếu số tiền 3.000.000đ của tháng 08/2026.",
        "Kiểm tra thông báo công tác cá nhân “Dọn phòng chung ca sáng”.",
        "Kiểm tra thông báo công tác theo Tổ 1 “Dọn phòng chung ca chiều”.",
        "Kiểm tra thông báo công tác theo Ban Thanh nhạc “Tập hát”.",
        "Chọn “Đánh dấu” tại thông báo kỳ thu, tải lại trang và kiểm tra số chưa đọc giảm từ 4 xuống 3.",
    ],
    [
        ("73-portal-danh-sach-thong-bao-desktop.png", "Hình 66. Danh sách bốn thông báo tài chính và công tác của học viên."),
        ("79-portal-thong-bao-da-doc-desktop.png", "Hình 67. Thông báo kỳ thu đã đọc và số chưa đọc giảm còn 3."),
    ],
    [
        "Portal có 4 thông báo chưa đọc: 1 tài chính và 3 công tác.",
        "Nội dung thông báo kỳ thu, số tiền và tháng thu hiển thị đúng.",
        "Thông báo công tác phân biệt đúng phạm vi cá nhân, Tổ và Ban.",
        "Trạng thái đã đọc được lưu đúng sau khi tải lại.",
    ],
)

add_step_section(
    doc,
    29,
    "Kiểm tra hoạt động và tài chính học viên",
    "Đối chiếu hoạt động được công khai với khoản học phí cá nhân trên portal.",
    [
        "Chọn “Hoạt động” và kiểm tra hoạt động “Sinh hoạt cộng đoàn đầu tuần”.",
        "Đối chiếu ngày 27/07/2026, thời gian 19:00–20:30 và địa điểm Phòng sinh hoạt chung.",
        "Chọn “Tài chính”.",
        "Kiểm tra tháng 08/2026 có hai khoản: phí lưu trú 1.200.000đ và ăn uống 1.800.000đ.",
        "Đối chiếu tổng cần thanh toán là 3.000.000đ.",
    ],
    [
        ("74-portal-hoat-dong-cong-khai-desktop.png", "Hình 68. Hoạt động công khai trên portal học viên."),
        ("75-portal-tai-chinh-hoc-vien-desktop.png", "Hình 69. Hai khoản phí tháng 08/2026 của học viên."),
    ],
    [
        "Hoạt động công khai hiển thị đúng nội dung do quản lý tạo.",
        "Tài chính cá nhân hiển thị đúng 2 khoản với tổng 3.000.000đ.",
        "Nguyễn Thị An chưa có phiếu thu nên cả hai khoản đang ở trạng thái Chưa đóng.",
    ],
)

add_step_section(
    doc,
    30,
    "Phản hồi công tác cá nhân, Tổ và Ban",
    "Kiểm tra danh sách ba công tác của học viên và thử phản hồi hoàn thành đối với công tác chung của Tổ, Ban.",
    [
        "Chọn “Công tác”.",
        "Kiểm tra công tác cá nhân “Dọn phòng chung ca sáng” đã hoàn thành.",
        "Kiểm tra công tác “Dọn phòng chung ca chiều” được giao cho Tổ 1.",
        "Kiểm tra công tác “Tập hát” được giao cho Ban Thanh nhạc.",
        "Chọn “Hoàn thành” cho công tác theo Tổ, chờ thông báo thành công rồi tải lại trang.",
        "Thực hiện tương tự cho công tác theo Ban và tải lại trang.",
    ],
    [
        ("76-portal-cong-tac-ca-nhan-to-ban-desktop.png", "Hình 70. Ba công tác theo cá nhân, Tổ và Ban trên portal."),
        ("77-portal-phan-hoi-cong-tac-khong-luu-desktop.png", "Hình 71. Trạng thái công tác Tổ/Ban vẫn chưa làm sau khi tải lại."),
    ],
    [
        "Danh sách công tác và phạm vi giao việc hiển thị đúng.",
        "Công tác cá nhân đã hoàn thành hiển thị đúng trạng thái.",
        "CHƯA ĐẠT: thao tác Hoàn thành cho công tác Tổ/Ban báo thành công nhưng không được lưu; sau tải lại vẫn còn 2 công tác Chưa làm.",
        "CHƯA ĐẠT: portal chưa có trường phản hồi nội dung, ghi chú hoặc minh chứng; học viên chỉ có nút Hoàn thành.",
    ],
)

add_step_section(
    doc,
    31,
    "Kiểm tra quyền vào cửa hàng của học viên",
    "Xác nhận portal chỉ cho học viên vào cửa hàng khi có phân công ca trực phù hợp.",
    [
        "Chọn “Cửa hàng” trên portal.",
        "Kiểm tra ngày trực 27/07/2026 và hai lựa chọn Ca sáng, Ca chiều.",
        "Đối chiếu phân công ca trực của tài khoản đang đăng nhập.",
        "Nếu không có phân công, hệ thống phải khóa nút “Vào Cửa hàng”.",
    ],
    [
        ("78-portal-cua-hang-chua-co-ca-truc-desktop.png", "Hình 72. Portal khóa quyền vào cửa hàng khi học viên chưa có ca trực."),
    ],
    [
        "Hệ thống không tìm thấy phân công ca trực phù hợp cho Nguyễn Thị An.",
        "Nút “Vào Cửa hàng” bị vô hiệu hóa đúng theo điều kiện truy cập.",
        "Luồng thao tác bán hàng trong ca trực đã được kiểm thử ở đợt trước; lần này chỉ xác nhận lại lớp kiểm soát quyền vào ca.",
    ],
)

add_step_section(
    doc,
    32,
    "Kiểm tra vai trò và cơ cấu phụ trách",
    "Xác nhận các chức vụ hiện tại của học viên và phạm vi tổ chức tương ứng trong nhiệm kỳ 2026–2027.",
    [
        "Mở nhóm “Phụ trách” và chọn “Tổng quan”.",
        "Kiểm tra ba vai trò: Trưởng toàn lưu xá, Tổ trưởng Tổ 1 và Trưởng ban Ban Thanh nhạc.",
        "Mở “Xem cơ cấu lưu xá”.",
        "Đối chiếu nhóm điều hành gồm Trưởng, Phó, Thư ký và Thủ quỹ.",
        "Đối chiếu Tổ 1 và Ban Thanh nhạc đều có Nguyễn Thị An ở vai trò phụ trách.",
    ],
    [
        ("80-portal-tong-quan-vai-tro-desktop.png", "Hình 73. Tổng quan ba vai trò đang hoạt động của học viên."),
        ("81-portal-co-cau-phu-trach-desktop.png", "Hình 74. Cơ cấu điều hành, Tổ và Ban thuộc phạm vi phụ trách."),
    ],
    [
        "Ba vai trò và nhiệm kỳ hiển thị đúng dữ liệu bổ nhiệm.",
        "Cơ cấu toàn lưu xá có đủ bốn vị trí điều hành.",
        "Tổ 1 và Ban Thanh nhạc đều hiển thị đúng người phụ trách và thành viên.",
    ],
)

add_step_section(
    doc,
    33,
    "Kiểm tra thành viên Tổ và Ban phụ trách",
    "Đối chiếu danh sách thành viên mà Tổ trưởng và Trưởng ban được phép xem.",
    [
        "Chọn “Tổ phụ trách” và mở Tổ 1.",
        "Kiểm tra 5 thành viên cùng mã học viên, tên thánh, phòng và số điện thoại.",
        "Chọn “Ban phụ trách” và mở Ban Thanh nhạc.",
        "Kiểm tra 5 lượt thành viên cùng thông tin liên hệ.",
        "Đối chiếu Nguyễn Thị An là người phụ trách ở cả hai danh sách.",
    ],
    [
        ("82-portal-thanh-vien-to-desktop.png", "Hình 75. Năm thành viên thuộc Tổ 1."),
        ("83-portal-thanh-vien-ban-desktop.png", "Hình 76. Năm thành viên thuộc Ban Thanh nhạc."),
    ],
    [
        "Tổ 1 hiển thị đúng 5 thành viên đang hoạt động.",
        "Ban Thanh nhạc hiển thị đúng 5 thành viên đang hoạt động.",
        "Thông tin phòng của thành viên hiển thị đúng dữ liệu đã gán.",
        "CHƯA ĐẠT: Dashboard sau đăng nhập vai trò Trưởng hiển thị phòng 0/0, không khớp dữ liệu phòng đang xuất hiện tại danh sách thành viên.",
    ],
)

add_step_section(
    doc,
    34,
    "Theo dõi công tác theo vai trò phụ trách",
    "Kiểm tra công tác toàn lưu xá, công tác của Tổ 1 và công tác của Ban Thanh nhạc.",
    [
        "Mở “Điều hành” để xem công tác toàn lưu xá trong khoảng 27/07–03/08/2026.",
        "Đối chiếu bốn công tác hiển thị trong danh sách điều hành.",
        "Mở đường dẫn “Công tác tổ” và kiểm tra công tác Dọn phòng chung ca chiều của Tổ 1.",
        "Mở đường dẫn “Công tác ban” và kiểm tra công tác Tập hát của Ban Thanh nhạc.",
        "Đối chiếu trạng thái từng công tác với màn hình Công tác của học viên.",
    ],
    [
        ("84-portal-cong-tac-dieu-hanh-desktop.png", "Hình 77. Công tác điều hành trong phạm vi toàn lưu xá."),
        ("85-portal-cong-tac-to-desktop.png", "Hình 78. Công tác được giao cho Tổ 1."),
        ("86-portal-cong-tac-ban-desktop.png", "Hình 79. Công tác được giao cho Ban Thanh nhạc."),
    ],
    [
        "Danh sách điều hành hiển thị bốn công tác, gồm một hoàn thành và ba chưa làm.",
        "Trang Công tác tổ lọc đúng một công tác của Tổ 1.",
        "Trang Công tác ban lọc đúng một công tác của Ban Thanh nhạc.",
        "CHƯA ĐẠT: các thẻ Tổng/Chưa làm/Hoàn thành trên trang Tổ và Ban vẫn lấy số toàn hệ thống (6/5/1), không khớp danh sách phạm vi chỉ có 1 công tác.",
    ],
)

support_heading = doc.add_heading("Nhóm 9. Hỗ trợ vận hành và kiểm thử", level=1)
support_heading.paragraph_format.page_break_before = True

doc.add_heading("9.1. Nếu không thấy menu Người dùng", level=2)
doc.add_paragraph(
    "Trước đây mục “Người dùng & quyền truy cập” chỉ xuất hiện ở Chế độ chi tiết, "
    "nên người quản lý dùng Chế độ đơn giản không thấy tài khoản học viên vừa tạo. "
    "Menu đã được điều chỉnh để chức năng cốt lõi này luôn xuất hiện trong nhóm "
    "“Báo cáo & Thiết lập” của Chế độ đơn giản."
)
add_numbered_steps(
    doc,
    [
        "Tải lại trang ResidenceCore sau khi hệ thống được cập nhật.",
        "Mở nhóm “Báo cáo & Thiết lập”.",
        "Chọn “Người dùng & quyền truy cập”.",
        "Nếu vẫn chưa thấy, mở trực tiếp đường dẫn /settings/users và báo quản trị kỹ thuật kiểm tra phiên bản giao diện.",
    ],
)

doc.add_heading("9.2. Checklist vận hành tiếp theo", level=2)
doc.add_paragraph(
    "Sau quy trình tiếp nhận ban đầu, người quản lý tiếp tục hoàn thiện hồ sơ "
    "và vận hành các phân hệ sau:"
)
add_bullets(
    doc,
    [
        "Đã phân 19 nữ sinh vào Tổ/Ban và bổ nhiệm 12 vị trí.",
        "Đã tạo bốn kiểu phân công công tác từ mẫu được giữ lại.",
        "Cần sửa cảnh báo xung đột giữa công tác và lịch học.",
        "Cần sửa lỗi hiển thị giờ công tác theo múi giờ Asia/Bangkok.",
        "Đã thiết lập lịch Điểm danh tối và lưu kết quả điểm danh mẫu cho 19 học viên.",
        "Cần kiểm tra lại giá trị giờ sau mỗi lần tạo/sửa lịch điểm danh.",
        "Đã bổ sung khung Điểm danh tối vào lịch sinh hoạt ngày thường và kiểm tra sau tải lại.",
        "Đã tạo hoạt động cộng đoàn và dữ liệu nền 19 người tham gia.",
        "Cần bổ sung giao diện quản lý người tham gia/điểm danh theo hoạt động.",
        "Đã bổ sung nội quy giờ trở về lưu xá và kiểm tra tìm kiếm theo mã.",
        "Đã tạo sản phẩm Nước suối 500ml và thiết lập giá bán ban đầu 5.000đ.",
        "Tiếp theo cần nhập kho để hình thành tồn kho và giá vốn thực tế.",
        "Đã nhập 104 đơn vị của 5 sản phẩm, hình thành giá vốn 2.558.000đ và khoản chi mua hàng tương ứng.",
        "Cần bổ sung nút khởi tạo sổ cửa hàng trên giao diện khi chưa có sổ hoạt động.",
        "Đã tạo phiếu bán 5 mặt hàng, doanh thu 425.000đ và kiểm tra giảm tồn kho tương ứng.",
        "Tiếp theo cần kiểm tra chốt ngày và quy trình xác nhận đưa dữ liệu sang sổ tài chính chung.",
        "Đã chốt ngày 27/07/2026, Review và xác nhận đẩy 2 phát sinh sang sổ tài chính chung.",
        "Các giao dịch mới phải được ghi vào ngày làm việc tiếp theo vì ngày 27/07/2026 đã khóa.",
        "Đã tạo kỳ thu tháng 08–12/2026 và 38 khoản phải thu tháng 08 cho 19 học viên.",
        "Đã thu đủ 3.000.000đ của một học viên; cần sửa lỗi thẻ tổng quan tài chính nhân ba số liệu.",
        "Đã ghi nhận khoản chi vật dụng vệ sinh 450.000đ trong tháng 08/2026.",
        "Đã bật và kiểm tra lưu kênh thông báo email; cần bổ sung liên kết điều hướng trong menu.",
        "Đã kiểm tra portal nhận thông báo kỳ thu, hoạt động và ba công tác theo cá nhân/Tổ/Ban.",
        "Cần sửa lỗi phản hồi Hoàn thành công tác Tổ/Ban không được lưu và bổ sung nội dung phản hồi.",
        "Đã xác nhận portal khóa quyền vào cửa hàng khi học viên chưa có phân công ca trực.",
        "Đã kiểm tra ba vai trò phụ trách, thành viên Tổ/Ban và công tác theo từng phạm vi.",
        "Cần sửa thống kê công tác Tổ/Ban đang lấy số toàn hệ thống và lỗi Dashboard hiển thị phòng 0/0.",
        "Tiếp tục theo dõi thông báo và báo cáo.",
        "Thiết lập khoản thu, thanh toán và công nợ theo chính sách của lưu xá.",
        "Quản lý cửa hàng, hoạt động, câu lạc bộ và kỹ năng khi các module được sử dụng.",
        "Kiểm tra portal học viên và phân quyền trước khi bàn giao tài khoản.",
    ],
)

add_callout(
    doc,
    "Quy tắc dữ liệu thật",
    "Khi vận hành chính thức, chỉ nhập thông tin cá nhân tối thiểu cần thiết; "
    "kiểm tra số điện thoại người thân; phân quyền đúng vai trò; và sao lưu trước "
    "các đợt xóa/chuyển dữ liệu lớn.",
    fill=LIGHT_BLUE,
)

doc.add_heading("9.3. Nhật ký kiểm thử của tài liệu", level=2)
add_bullets(
    doc,
    [
        "Database được backup trước khi reset.",
        "Reset hoàn tất trong transaction và giữ đúng admin, cơ cấu, mẫu/khung công tác.",
        "Toàn bộ 34 mục thuộc 8 nhóm nghiệp vụ đã được thực hiện lại trên giao diện.",
        "Kết quả tạo học viên, liên hệ, phòng, gán phòng, thông tin học tập và lịch học đều đạt.",
        "Đã hoàn thiện dữ liệu nền: 19 học viên, 4 Tổ, 4 Ban, 12 lượt bổ nhiệm và 4 công tác đại diện.",
        "Menu Người dùng đã được bổ sung vào Chế độ đơn giản và kiểm tra trực tiếp tài khoản an.nguyen.",
        "Kiểm thử công tác phát hiện hai tồn tại: chưa kiểm tra xung đột lịch học và hiển thị lệch giờ +7.",
        "Kiểm thử điểm danh đạt: lưu và tải lại đúng 16 có mặt, 1 trễ, 1 vắng phép, 1 vắng không phép.",
        "Luồng tạo lịch điểm danh cần theo dõi lỗi giờ nhập 21:00 nhưng có lúc lưu về mặc định 07:00.",
        "Lịch sinh hoạt ngày thường đã tăng từ 11 lên 12 khung giờ với mục Điểm danh tối 21:00–21:15.",
        "Đã tạo hoạt động SHCD-20260727 và dữ liệu nền 19 người tham gia, gồm 17 tham dự và 2 vắng.",
        "Đã sửa ánh xạ cột vai trò người tham gia từ activityParticipantRole sang role cho đúng database.",
        "Đã tạo nội quy CURFEW_BMT, xác nhận tổng số nội quy tăng lên 4 và tìm kiếm theo mã hoạt động đúng.",
        "Đã tạo sản phẩm NUOC_SUOI_500, thiết lập giá bán 5.000đ và xác nhận dữ liệu còn tồn tại sau khi tải lại.",
        "Đã khởi tạo sổ CUA_HANG_LUU_XA để khắc phục điều kiện thiếu sổ sau reset dữ liệu.",
        "Đã tạo phiếu nhập 5 mặt hàng, 104 đơn vị, tổng 2.558.000đ; tồn kho, giá vốn và khoản chi tự động đều khớp.",
        "Đã tạo phiếu bán 9 đơn vị, doanh thu 425.000đ; tồn kho giảm đúng và khoản thu bán hàng được tự động ghi nhận.",
        "Đã chốt ngày CHOT-20260727-4, Review và xác nhận đẩy sổ chung với trạng thái cuối cùng đã khóa.",
        "Đã tạo kỳ thu học kỳ I, sinh 38 khoản phải thu trị giá đúng 57.000.000đ cho tháng 08/2026.",
        "Đã thu đủ 3.000.000đ của Vũ Thị Bảo Ngọc và phát hiện thẻ tổng quan tài chính nhân ba số liệu đúng.",
        "Đã ghi nhận khoản chi một lần 450.000đ và xác nhận hiển thị đúng trong sổ dòng tiền lưu xá.",
        "Đã bật thông báo email và xác nhận tùy chọn còn hiệu lực sau khi tải lại trang.",
        "Đã đăng nhập portal Nguyễn Thị An, kiểm tra 4 thông báo, hoạt động công khai và công nợ 3.000.000đ.",
        "Đã kiểm tra ba công tác theo cá nhân/Tổ/Ban; phát hiện phản hồi hoàn thành công tác tập thể không được lưu.",
        "Đã kiểm tra điều kiện vào cửa hàng và xác nhận nút truy cập bị khóa khi không có ca trực.",
        "Đã kiểm tra tổng quan vai trò, cơ cấu, 5 thành viên Tổ 1 và 5 thành viên Ban Thanh nhạc.",
        "Đã kiểm tra công tác điều hành/Tổ/Ban và phát hiện các thẻ thống kê phạm vi dùng sai số toàn hệ thống.",
        "Ảnh mục 3.7 và mục 3.8 được chụp lại ở viewport desktop 1600×1100, không dùng ảnh mobile hoặc ảnh full-page bị co layout.",
        "Ảnh minh họa được chụp trực tiếp trong cùng lượt kiểm thử ngày 27/07/2026.",
    ],
)

OUTPUT.parent.mkdir(parents=True, exist_ok=True)
doc.save(OUTPUT)
print(OUTPUT)
